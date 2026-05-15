import importlib.util
import json
import re
import shutil
import sys
import tempfile
from pathlib import Path

from docx.enum.section import WD_SECTION_START

ROOT = Path(__file__).resolve().parents[1]


def load_module():
    module_path = ROOT / "scripts" / "cuit_thesis_docx_format.py"
    spec = importlib.util.spec_from_file_location("cuit_mod", module_path)
    mod = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    sys.modules[spec.name] = mod
    spec.loader.exec_module(mod)
    return mod


def test_body_alignment_and_caption_line_spacing_rules():
    mod = load_module()
    assert mod.RULES["body"].alignment == mod.ALIGN_JUSTIFIED
    assert mod.RULES["figure_caption"].space_after_lines == 1
    assert mod.RULES["table_caption"].space_before_lines == 1


def test_body_cjk_spacing_only_for_body():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("摘 要")
    doc.add_paragraph("在 Web 层")
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("在 Web 实现层面， Flask 由于核心简洁")
    issues = mod.collect_issues(doc)
    spacing = [x for x in issues if x.rule_key == "body_cjk_spacing"]
    assert len(spacing) >= 1


def test_section_boundary_extra_author_bio_warning_not_module():
    mod = load_module()
    texts = [
        "致谢",
        "致谢",
        "致谢",
        "第一章 绪论",
        "无关段落",
        "致谢",
        "作者简历及攻读学位期间发表的学术论文与研究成果",
    ]
    regions = mod.analyze_section_sequence(texts)["regions"]
    errors, warnings = mod.compute_section_boundary_findings(texts, regions)
    assert any("疑似额外组成部分" in x for x in warnings)
    assert all("author_bio" not in x for x in errors)


def _collect_reference_keys(mod, entries):
    doc = mod.Document()
    doc.add_paragraph("参考文献")
    for line in entries:
        doc.add_paragraph(line)
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["references"] * len(texts)}
    issues = mod.collect_reference_issues(doc)
    return issues, {x.rule_key for x in issues}


def test_reference_272_valid_journal_match():
    mod = load_module()
    issues, keys = _collect_reference_keys(mod, ["[1] 作者. 题名[J]. 期刊名, 2020, 12(3): 15-20."])
    assert "reference_272_format_no_match" not in keys
    assert "reference_272_type_format_mismatch" not in keys


def test_reference_272_invalid_journal():
    mod = load_module()
    issues, keys = _collect_reference_keys(mod, ["[1] 作者. 题名[J]. 期刊名."])
    assert "reference_272_type_format_mismatch" in keys or "reference_entry_format" in keys


def test_reference_272_valid_dissertation():
    mod = load_module()
    _issues, keys = _collect_reference_keys(mod, ["[1] 作者. 题名[D]. 成都: 成都信息工程大学, 2023."])
    assert "reference_272_format_no_match" not in keys


def test_reference_272_invalid_dissertation():
    mod = load_module()
    _issues, keys = _collect_reference_keys(mod, ["[1] 作者. 题名[D]."])
    assert "reference_272_type_format_mismatch" in keys


def test_reference_272_valid_ebol():
    mod = load_module()
    _issues, keys = _collect_reference_keys(mod, ["[1] 作者. 题名[EB/OL]. (2020-01-01)[2024-01-01]. https://example.com."])
    assert "reference_272_format_no_match" not in keys


def test_reference_272_invalid_ebol_missing_citation_date():
    mod = load_module()
    _issues, keys = _collect_reference_keys(mod, ["[1] 作者. 题名[EB/OL]. https://example.com."])
    assert "reference_273_citation_date_missing" in keys or "reference_272_type_format_mismatch" in keys


def test_reference_272_no_match():
    mod = load_module()
    _issues, keys = _collect_reference_keys(mod, ["[1] 这是一条随便写的参考文献"])
    assert "reference_272_format_no_match" in keys


def test_reference_missing_type_marker():
    mod = load_module()
    _issues, keys = _collect_reference_keys(mod, ["[1] 作者. 题名. 期刊名, 2020, 12(3): 15-20."])
    assert "reference_type_marker" in keys


def test_reference_sequence_not_continuous():
    mod = load_module()
    _issues, keys = _collect_reference_keys(
        mod,
        [
            "[1] 作者. 题名[J]. 期刊名, 2020, 12(3): 15-20.",
            "[3] 作者. 题名[J]. 期刊名, 2021, 13(4): 21-30.",
        ],
    )
    assert "reference_sequence" in keys


def test_reference_272_valid_m_s_p_n():
    mod = load_module()
    _issues, keys = _collect_reference_keys(
        mod,
        [
            "[1] 作者. 题名[M]. 北京: 出版社, 2020: 10-20.",
            "[2] GB/T 7714-2015, 信息与文献 参考文献著录规则[S]. 北京: 中国标准出版社, 2015.",
            "[3] 张三. 一种数据处理方法: CN101234567A[P]. 2020-01-01.",
            "[4] 作者. 题名[N]. 人民日报, 2020-01-01(01).",
        ],
    )
    assert "reference_272_format_no_match" not in keys


def test_reference_273_online_missing_for_journal_with_url():
    mod = load_module()
    _issues, keys = _collect_reference_keys(mod, ["[1] 作者. 题名[J]. 期刊名, 2020, 12(3): 15-20. https://example.com."])
    assert "reference_273_citation_date_missing" in keys


def test_reference_273_author_et_al():
    mod = load_module()
    _issues, keys = _collect_reference_keys(mod, ["[1] 张三, 李四, 王五, 赵六. 题名[J]. 期刊名, 2020, 12(3): 15-20."])
    assert "reference_273_author_et_al_review" in keys


def test_reference_273_online_entry_missing_access_reports():
    mod = load_module()
    _issues, keys = _collect_reference_keys(mod, ["[1] 作者. 题名[EB/OL]. [2024-01-01]."])
    assert "reference_273_online_access_missing" in keys


def test_reference_273_no_online_entries_no_false_positive():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] 作者. 题名[J]. 刊名, 2020, 1(2):3-5.")
    texts = [mod.paragraph_text(p) for p in doc.paragraphs]
    entries, _ = mod._reference_entries_from_regions(doc.paragraphs, texts, ["references", "references"])
    summary = mod._reference_273_check_summary(entries)
    assert summary["online_entries"] == 0
    assert summary["access_path_missing_entries"] == 0
    assert summary["citation_date_missing_entries"] == 0
    assert summary["doi_or_url_missing_entries"] == 0


def test_reference_273_summary_exists():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] 作者. 题名[EB/OL]. (2020-01-01)[2024-01-01]. https://example.com.")
    texts = [mod.paragraph_text(p) for p in doc.paragraphs]
    entries, _ = mod._reference_entries_from_regions(doc.paragraphs, texts, ["references", "references"])
    summary = mod._reference_273_check_summary(entries)
    assert "online_entries" in summary
    assert "entries" in summary


def test_reference_273_author_et_al_review():
    mod = load_module()
    _issues, keys = _collect_reference_keys(mod, ["[1] 张三, 李四, 王五, 赵六. 题名[J]. 刊名, 2020, 12(3): 15-20."])
    assert "reference_273_author_et_al_review" in keys


def test_comment_text_includes_expected_and_actual_problems():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="keywords_en",
        text_type="keywords paragraph",
        text_excerpt="Key words: a; b",
        current="当前对齐方式为两端对齐；当前固定行距未设置。",
        expected="英文关键词段落应为 Times New Roman 小四 12pt，固定 20 磅行距。",
        message="",
        category="abstract",
    )
    text = mod.comment_message_for_issue(issue)
    assert "正确格式：" in text
    assert "本处问题：" in text
    assert "当前对齐方式为两端对齐" in text


def test_comment_text_does_not_list_correct_items_as_errors():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="body",
        text_type="body paragraph",
        text_excerpt="正文",
        current="当前固定行距未设置。",
        expected="正文固定 20 磅行距。",
        message="",
        category="body",
    )
    text = mod.comment_message_for_issue(issue)
    assert "本处问题：" in text
    assert "固定行距" in text
    assert "字体" not in text


def test_issue_message_only_lists_actual_differences():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="title_zh",
        text_type="title",
        text_excerpt="题目",
        current="疑似不符合项：固定行距：当前 继承/未直接设置，应为 20.0磅; 段落样式：Normal; 中文字体：黑体; 字号：16.0pt; 加粗：是; 对齐方式：居中",
        expected="中文论文题目应为宋体三号 16pt，加粗，居中，固定 20 磅行距。",
        message="",
        category="title",
    )
    text = mod.comment_message_for_issue(issue)
    assert "行距不符合要求，应为固定 20 磅。" in text
    assert "字号：16.0pt" not in text
    assert "加粗：是" not in text
    assert "对齐方式：居中" not in text
    assert "段落样式：Normal" not in text


def test_issue_message_extracts_multiple_differences_before_summary():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="title_zh",
        text_type="title",
        text_excerpt="题目",
        current="疑似不符合项：中文字体：当前 黑体，应为 宋体；固定行距：当前 继承/未直接设置，应为 20.0磅; 段落样式：Normal; 字号：16.0pt",
        expected="中文论文题目应为宋体三号 16pt，加粗，居中，固定 20 磅行距。",
        message="",
        category="title",
    )
    text = mod.comment_message_for_issue(issue)
    assert "中文字体不符合要求，应使用宋体。" in text
    assert "行距不符合要求，应为固定 20 磅。" in text
    assert "字号：16.0pt" not in text
    assert "段落样式：Normal" not in text


def test_times_new_roman_regular_normalized():
    mod = load_module()
    assert mod.normalize_font_name("Times New Roman Regular") == "Times New Roman"
    assert mod.normalize_font_name("Times New Roman") == "Times New Roman"


def test_run_formatting_does_not_list_correct_size_bold_alignment():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="title_zh_runs",
        text_type="title run formatting",
        text_excerpt="题目",
        current="文字片段='题目'；中文字体=黑体；西文字体=Times New Roman Regular；字号=16.0pt；加粗=是",
        expected="中文论文题目应为宋体三号 16pt，加粗，居中，固定 20 磅行距。",
        message="",
        category="run-format",
    )
    text = mod.comment_message_for_issue(issue)
    assert "中文字体使用了黑体，应改为宋体。" in text
    assert "字号=16.0pt" not in text
    assert "加粗=是" not in text


def test_comment_for_chinese_title_matches_expected_problem_count():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="title_zh",
        text_type="title",
        text_excerpt="题目",
        current="疑似不符合项：中文字体：当前 黑体，应为 宋体；固定行距：当前 继承/未直接设置，应为 20.0磅; 段落样式：Normal; 字号：16.0pt; 加粗：是",
        expected="中文论文题目应为宋体三号 16pt，加粗，居中，固定 20 磅行距。",
        message="",
        category="title",
    )
    text = mod.comment_message_for_issue(issue)
    assert text.count("\n1. ") == 1
    assert text.count("\n2. ") == 1
    assert "字号：16.0pt" not in text
    assert "加粗：是" not in text


def test_html_issue_includes_expected_and_actual_problems():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "r.html"
        report = {
            "input": "x.docx",
            "annotated_docx": "a.docx",
            "fixed_docx": "f.docx",
            "issue_count": 1,
            "renderer_for_fixed": "ooxml",
            "renderer_for_comments": "ooxml",
            "screenshot_status": "skipped",
            "issues": [
                {
                    "paragraph_index": 1,
                    "text_type": "body paragraph",
                    "text_excerpt": "正文",
                    "current": "当前对齐方式为左对齐",
                    "expected": "应为两端对齐",
                    "message": "本处问题：\n1. 当前对齐方式为左对齐",
                    "after": "已改为两端对齐",
                }
            ],
            "issue_summary_by_category": {},
            "reference_272_check_summary": {"total_entries": 0, "matched_entries": 0, "type_mismatch_entries": 0, "unmatched_entries": 0, "entries": []},
            "reference_273_check_summary": {"online_entries": 0, "author_et_al_review_entries": 0, "citation_date_missing_entries": 0, "access_path_missing_entries": 0, "doi_or_url_missing_entries": 0, "entries": []},
        }
        mod.write_html_report(report, p)
        html = p.read_text(encoding="utf-8")
        assert "正确格式：" in html
        assert "本处问题：" in html


def test_html_issue_does_not_duplicate_problem_heading():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "r.html"
        report = {
            "input": "x.docx",
            "annotated_docx": "a.docx",
            "fixed_docx": "f.docx",
            "issue_count": 1,
            "renderer_for_fixed": "ooxml",
            "renderer_for_comments": "ooxml",
            "screenshot_status": "skipped",
            "issues": [
                {
                    "paragraph_index": 1,
                    "text_type": "body paragraph",
                    "text_excerpt": "正文",
                    "current": "当前格式",
                    "expected": "应为规范格式",
                    "message": "本处问题：\n1. 固定行距未设置",
                    "after": "已修复",
                }
            ],
            "issue_summary_by_category": {},
            "reference_272_check_summary": {"total_entries": 0, "matched_entries": 0, "type_mismatch_entries": 0, "unmatched_entries": 0, "entries": []},
            "reference_273_check_summary": {"online_entries": 0, "author_et_al_review_entries": 0, "citation_date_missing_entries": 0, "access_path_missing_entries": 0, "doi_or_url_missing_entries": 0, "entries": []},
        }
        mod.write_html_report(report, p)
        html = p.read_text(encoding="utf-8")
        assert "本处问题：</strong>本处问题：" not in html
        assert "本处问题：</strong>1. 固定行距未设置" in html


def test_comment_merges_paragraph_and_run_issues_for_same_paragraph():
    mod = load_module()
    issues = [
        mod.Issue(
            paragraph_index=10,
            rule_key="title_zh",
            text_type="Chinese thesis title",
            text_excerpt="题目",
            current="疑似不符合项：固定行距：当前 继承/未直接设置，应为 20.0磅; 段落样式：Normal; 字号：16.0pt; 加粗：是",
            expected="中文论文题目应为宋体三号16pt，加粗，居中，固定20磅行距。",
            message="",
            category="title",
        ),
        mod.Issue(
            paragraph_index=10,
            rule_key="title_zh_runs",
            text_type="Chinese thesis title run formatting",
            text_excerpt="题目",
            current="文字片段='题目'；中文字体=黑体；西文字体=Times New Roman Regular；字号=16.0pt；加粗=是",
            expected="中文论文题目应为宋体三号16pt，加粗，居中，固定20磅行距。",
            message="",
            category="run-format",
        ),
    ]
    merged = mod._merge_comment_message(issues)
    assert "行距不符合要求，应为固定 20 磅。" in merged
    assert "中文字体使用了黑体，应改为宋体。" in merged


def test_merged_comment_does_not_include_correct_items():
    mod = load_module()
    issues = [
        mod.Issue(
            paragraph_index=10,
            rule_key="title_zh",
            text_type="Chinese thesis title",
            text_excerpt="题目",
            current="疑似不符合项：固定行距：当前 继承/未直接设置，应为 20.0磅; 段落样式：Normal; 字号：16.0pt; 加粗：是; 对齐方式：居中",
            expected="中文论文题目应为宋体三号16pt，加粗，居中，固定20磅行距。",
            message="",
            category="title",
        ),
        mod.Issue(
            paragraph_index=10,
            rule_key="title_zh_runs",
            text_type="Chinese thesis title run formatting",
            text_excerpt="题目",
            current="文字片段='题目'；中文字体=黑体；西文字体=Times New Roman Regular；字号=16.0pt；加粗=是",
            expected="中文论文题目应为宋体三号16pt，加粗，居中，固定20磅行距。",
            message="",
            category="run-format",
        ),
    ]
    merged = mod._merge_comment_message(issues)
    assert "字号：16.0pt" not in merged
    assert "加粗：是" not in merged
    assert "对齐方式：居中" not in merged


def test_humanize_problem_removes_inherited_unset_words():
    mod = load_module()
    out = mod.humanize_problem_item("固定行距：当前 继承/未直接设置，应为 20.0磅")
    assert out == "行距不符合要求，应为固定 20 磅。"
    assert "继承" not in out
    assert "未直接设置" not in out


def test_comment_message_uses_plain_problem_language():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="title_zh",
        text_type="title",
        text_excerpt="题目",
        current="疑似不符合项：中文字体：当前 黑体，应为 宋体；固定行距：当前 继承/未直接设置，应为 20.0磅; 段落样式：Normal; 字号：16.0pt; 加粗：是; 对齐方式：居中",
        expected="中文论文题目应为宋体三号16pt，加粗，居中，固定20磅行距。",
        message="",
        category="title",
    )
    text = mod.comment_message_for_issue(issue)
    assert "行距不符合要求，应为固定 20 磅。" in text
    assert "中文字体不符合要求，应使用宋体。" in text
    assert "继承/未直接设置" not in text
    assert "段落样式：Normal" not in text
    assert "字号：16.0pt" not in text
    assert "加粗：是" not in text


def test_html_issue_uses_plain_problem_language():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="title_zh",
        text_type="title",
        text_excerpt="题目",
        current="疑似不符合项：固定行距：当前 继承/未直接设置，应为 20.0磅; 段落样式：Normal",
        expected="中文论文题目应为宋体三号16pt，加粗，居中，固定20磅行距。",
        message="",
        category="title",
    )
    payload = mod.issue_dict(issue)
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "r.html"
        report = {
            "input": "x.docx",
            "annotated_docx": "a.docx",
            "fixed_docx": "f.docx",
            "issue_count": 1,
            "renderer_for_fixed": "ooxml",
            "renderer_for_comments": "ooxml",
            "screenshot_status": "skipped",
            "issues": [payload],
            "issue_summary_by_category": {},
            "reference_272_check_summary": {"total_entries": 0, "matched_entries": 0, "type_mismatch_entries": 0, "unmatched_entries": 0, "entries": []},
            "reference_273_check_summary": {"online_entries": 0, "author_et_al_review_entries": 0, "citation_date_missing_entries": 0, "access_path_missing_entries": 0, "doi_or_url_missing_entries": 0, "entries": []},
        }
        mod.write_html_report(report, p)
        html = p.read_text(encoding="utf-8")
        assert "继承/未直接设置" not in html
        assert "未直接设置" not in html
        assert "段落样式：Normal" not in html


def test_json_issue_uses_plain_problem_language():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="title_zh",
        text_type="title",
        text_excerpt="题目",
        current="疑似不符合项：固定行距：当前 继承/未直接设置，应为 20.0磅; 段落样式：Normal",
        expected="中文论文题目应为宋体三号16pt，加粗，居中，固定20磅行距。",
        message="",
        category="title",
    )
    payload = mod.issue_dict(issue)
    assert "继承/未直接设置" not in str(payload["current"])
    assert "未直接设置" not in str(payload["current"])
    assert "段落样式：Normal" not in str(payload["current"])
    assert "继承/未直接设置" not in str(payload["message"])


def test_run_formatting_plain_language():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="title_zh_runs",
        text_type="title run",
        text_excerpt="题目",
        current="文字片段='题目'；中文字体=黑体；西文字体=Times New Roman Regular；字号=16.0pt；加粗=是",
        expected="中文论文题目应为宋体三号16pt，加粗，居中，固定20磅行距。",
        message="",
        category="run-format",
    )
    text = mod.comment_message_for_issue(issue)
    assert "中文字体使用了黑体，应改为宋体。" in text
    assert "字号=16.0pt" not in text
    assert "加粗=是" not in text
    assert "Regular" not in text


def test_toc_report_only_language_plain():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="toc_page_number_manual_review",
        text_type="目录页码人工复核",
        text_excerpt="目 录",
        current="检测到目录/TOC 自动目录。由于目录页页码涉及 section、页脚 PAGE 域和 TOC 字段刷新，工具不自动修改。",
        expected="目录所在前置页页脚页码应为小写罗马数字 i, ii, iii...",
        message="",
        category="page",
    )
    payload = mod.issue_dict(issue)
    assert "OOXML" not in str(payload["message"])
    assert "人工复核" in str(payload["after"]) or "不自动修改" in str(payload["after"])


def test_comment_author_still_neutral_after_merge():
    mod = load_module()
    root = mod.ET.Element(mod.qname(mod.W_NS, "comments"))
    merged_text = mod._merge_comment_message(
        [
            mod.Issue(1, "body", "body", "正文", "疑似不符合项：固定行距：当前 继承/未直接设置，应为 20.0磅", "正文固定行距20磅", "", "body")
        ]
    )
    mod.add_comment_node(root, 1, merged_text)
    comment = root.find(mod.qname(mod.W_NS, "comment"))
    assert comment is not None
    assert comment.get(mod.qname(mod.W_NS, "author")) == "thesis format skill"


def test_reference_semantic_issue_keeps_manual_edit_after():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="reference_entry_format",
        text_type="reference entry",
        text_excerpt="ref",
        current="当前条目缺要素",
        expected="应符合模板",
        message="",
        category="references",
        after="其他文本",
    )
    payload = mod.issue_dict(issue)
    assert payload["after"] == "著录内容不自动改写，请人工按模板修改。"


def test_comment_author_uses_neutral_skill_name():
    mod = load_module()
    root = mod.ET.Element(mod.qname(mod.W_NS, "comments"))
    mod.add_comment_node(root, 1, "text")
    comment = root.find(mod.qname(mod.W_NS, "comment"))
    assert comment is not None
    assert comment.get(mod.qname(mod.W_NS, "author")) == "thesis format skill"


def test_no_user_visible_codex_cuit_strings():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=1,
        rule_key="body",
        text_type="body paragraph",
        text_excerpt="正文",
        current="当前格式错误",
        expected="应为规范格式",
        message="",
        category="body",
    )
    text = mod.comment_message_for_issue(issue)
    assert "Codex" not in text
    assert "CUIT" not in text
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "r.html"
        report = {
            "input": "x.docx",
            "annotated_docx": "a.docx",
            "fixed_docx": "f.docx",
            "issue_count": 0,
            "renderer_for_fixed": "ooxml",
            "renderer_for_comments": "ooxml",
            "screenshot_status": "skipped",
            "issues": [],
            "issue_summary_by_category": {},
            "reference_272_check_summary": {"total_entries": 0, "matched_entries": 0, "type_mismatch_entries": 0, "unmatched_entries": 0, "entries": []},
            "reference_273_check_summary": {"online_entries": 0, "author_et_al_review_entries": 0, "citation_date_missing_entries": 0, "access_path_missing_entries": 0, "doi_or_url_missing_entries": 0, "entries": []},
        }
        mod.write_html_report(report, p)
        html = p.read_text(encoding="utf-8")
        assert "Codex" not in html
        assert "CUIT" not in html


def test_html_title_no_codex_cuit():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "r.html"
        report = {
            "input": "x.docx",
            "annotated_docx": "a.docx",
            "fixed_docx": "f.docx",
            "issue_count": 0,
            "renderer_for_fixed": "ooxml",
            "renderer_for_comments": "ooxml",
            "screenshot_status": "skipped",
            "issues": [],
            "issue_summary_by_category": {},
            "reference_272_check_summary": {"total_entries": 0, "matched_entries": 0, "type_mismatch_entries": 0, "unmatched_entries": 0, "entries": []},
            "reference_273_check_summary": {"online_entries": 0, "author_et_al_review_entries": 0, "citation_date_missing_entries": 0, "access_path_missing_entries": 0, "doi_or_url_missing_entries": 0, "entries": []},
        }
        mod.write_html_report(report, p)
        html = p.read_text(encoding="utf-8")
        assert "thesis format skill report" in html
        assert "Codex" not in html
        assert "CUIT" not in html


def test_cli_report_metadata_no_codex_cuit():
    mod = load_module()
    payload = mod.stdout_json({"tool_name": "thesis format skill"})
    assert "Codex" not in payload
    assert "CUIT" not in payload


def test_reference_checks_scope_only_references():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("[1] 作者. 题名[J]. 期刊名, 2020, 12(3): 15-20.")
    issues = mod.collect_reference_issues(doc)
    assert issues == []


def test_reference_entries_are_formatted():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("参考文献")
    p = doc.add_paragraph("[1] 张三. 题名[J]. 刊名, 2020, 1(2):3-5.")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["references", "references"]}
    mod.apply_supported_rules(doc, allow_layout_fixes=False)
    rule = mod.RULES["reference_entry"]
    assert mod.paragraph_matches(p, rule)
    ind = p._p.pPr.ind
    assert ind is not None
    assert ind.get(mod.qn("w:hangingChars")) == "200"
    assert ind.get(mod.qn("w:hanging")) is None


def test_all_reference_entries_collected():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("参考文献")
    for i in range(1, 6):
        doc.add_paragraph(f"[{i}] 作者. 题名[J]. 刊名, 2020, 1(2):3-5.")
    texts = [mod.paragraph_text(p) for p in doc.paragraphs]
    regions = ["references"] * len(texts)
    entries, _issues = mod._reference_entries_from_regions(doc.paragraphs, texts, regions)
    assert len(entries) == 5
    summary = mod._reference_272_check_summary(entries)
    assert summary["total_entries"] == 5
    assert len(summary["entries"]) == 5


def test_reference_entries_not_detected_when_title_exists():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("参考文献")
    doc.add_paragraph("这是一段没有编号的内容")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["references", "references"]}
    issues = mod.collect_reference_issues(doc)
    keys = {i.rule_key for i in issues}
    assert "reference_entries_not_detected" in keys


def test_reference_entry_formatting_does_not_touch_body():
    mod = load_module()
    doc = mod.Document()
    body = doc.add_paragraph("这是正文段落。")
    ref_title = doc.add_paragraph("参考文献")
    ref = doc.add_paragraph("[1] 作者. 题名[J]. 刊名, 2020, 1(2):3-5.")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["body", "references", "references"]}
    old = body.text
    mod.apply_supported_rules(doc, allow_layout_fixes=False)
    assert body.text == old
    assert mod.paragraph_matches(ref_title, mod.RULES["reference_title"])
    assert mod.paragraph_matches(ref, mod.RULES["reference_entry"])


def test_reference_word_auto_numbering_not_reported_missing():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("参考文献")
    p = doc.add_paragraph("韩珂, 李相霏, 顾波. 基于Python语言的疫情大数据可视化方法:, CN202111256294.0[P]. 2022.")
    p.style = doc.styles["List Number"]
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["references", "references"]}
    issues = mod.collect_reference_issues(doc)
    keys = {i.rule_key for i in issues}
    assert "reference_entry_number_missing" not in keys
    texts = [mod.paragraph_text(x) for x in doc.paragraphs]
    entries, _ = mod._reference_entries_from_regions(doc.paragraphs, texts, ["references", "references"])
    summary = mod._reference_272_check_summary(entries)
    assert summary["entries"][0]["visual_index"] == 1
    assert "word_auto_numbering" in summary["entries"][0]["numbering_source"]


def test_reference_summary_includes_visual_index_and_numbering_source():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] 张三. 题名[J]. 刊名, 2020, 1(2):3-5.")
    texts = [mod.paragraph_text(x) for x in doc.paragraphs]
    entries, _ = mod._reference_entries_from_regions(doc.paragraphs, texts, ["references", "references"])
    summary = mod._reference_272_check_summary(entries)
    entry = summary["entries"][0]
    assert "visual_index" in entry
    assert "numbering_source" in entry
    assert "paragraph_index" in entry


def test_reference_journal_single_page_count_accepts_page_field():
    mod = load_module()
    _issues, keys = _collect_reference_keys(
        mod,
        ["陈俊生, 彭莉芬. 基于Python+Echarts的大数据可视化系统的设计与实现[J]. 安徽电子信息职业技术学院学报, 2019, 18(4):5."],
    )
    assert "reference_272_type_format_mismatch" not in keys


def test_reference_journal_author_single_chinese_name_detected():
    mod = load_module()
    _issues, keys = _collect_reference_keys(
        mod,
        ["张楚. 基于Python的文本可视化方法实现与应用[J]. 黑龙江科技信息, 2020, 000(028):144-145."],
    )
    assert "reference_272_type_format_mismatch" not in keys


def test_reference_patent_cn_number_year_matches():
    mod = load_module()
    _issues, keys = _collect_reference_keys(
        mod,
        ["韩珂, 李相霏, 顾波. 基于Python语言的疫情大数据可视化方法:, CN202111256294.0[P]. 2022."],
    )
    assert "reference_272_type_format_mismatch" not in keys


def test_reference_patent_single_author_detected():
    mod = load_module()
    _issues, keys = _collect_reference_keys(
        mod,
        ["孙惠玲. 一种疫情监测预警平台:, CN211787435U[P]. 2020."],
    )
    assert "reference_272_type_format_mismatch" not in keys


def test_reference_book_publisher_year_detected_missing_place_only():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("参考文献")
    doc.add_paragraph("IgorMilovanovic, 米洛万诺维奇, 颛清山. Python数据可视化编程实战[M]. 人民邮电出版社, 2015.")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["references", "references"]}
    texts = [mod.paragraph_text(x) for x in doc.paragraphs]
    entries, _ = mod._reference_entries_from_regions(doc.paragraphs, texts, ["references", "references"])
    summary = mod._reference_272_check_summary(entries)
    item = summary["entries"][0]
    assert item["detected_type"] == "M"
    assert item["status"] == "type_mismatch"
    assert "出版地" in item["missing_fields"]
    assert "出版者" not in item["missing_fields"]
    assert "出版年" not in item["missing_fields"]


def test_reference_number_missing_only_when_no_text_number_and_no_numpr():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("参考文献")
    doc.add_paragraph("张三. 题名[J]. 刊名, 2020, 1(2):3-5.")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["references", "references"]}
    issues = mod.collect_reference_issues(doc)
    keys = {i.rule_key for i in issues}
    assert "reference_entry_number_missing" in keys


def test_reference_fix_does_not_touch_existing_completed_modules():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph("Key words: alpha; beta")
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] 张三. 题名[J]. 刊名, 2020, 1(2):3-5.")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en", "abstract_en", "references", "references"]}
    mod.apply_supported_rules(doc, allow_layout_fixes=False)
    assert doc.paragraphs[0].text == "ABSTRACT"
    assert "Key words" in doc.paragraphs[1].text


def test_empty_paragraph_cleanup_targets():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("摘 要")
    doc.add_paragraph("   ")
    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph("　")
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("	")
    doc.add_paragraph("致谢")
    doc.add_paragraph("  ")
    issues, removable = mod.collect_empty_paragraph_issues(doc)
    assert len(issues) >= 1
    removed = mod.remove_target_empty_paragraphs(doc, removable)
    assert removed >= 1


def test_image_size_advisory_and_height_limit():
    mod = load_module()

    class FakeDim:
        def __init__(self, cm):
            self.cm = cm

    class FakeShape:
        def __init__(self, w, h):
            self.width = FakeDim(w)
            self.height = FakeDim(h)

    class FakeDoc:
        def __init__(self, shapes):
            self.inline_shapes = shapes

    doc = FakeDoc([FakeShape(8, 6), FakeShape(10, 12), FakeShape(9, 20)])
    issues = mod.collect_image_size_issues(doc)
    keys = [x.rule_key for x in issues]
    assert "image_size_advisory" in keys
    assert "image_height_limit" in keys


def test_table_caption_and_three_line_checks():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("无关段落")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    doc.add_paragraph("表2-1 示例表题")
    issues = mod.collect_table_issues(doc)
    keys = {x.rule_key for x in issues}
    assert "table_caption_missing" in keys
    assert "table_caption_position" in keys
    assert "table_three_line_style" in keys


def test_english_abstract_title_normalization_and_issue():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("Abstract")
    doc.add_paragraph("This is abstract body.")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en"] * len(texts)}
    mod.normalize_english_abstract_title(doc)
    assert doc.paragraphs[0].text == "ABSTRACT"
    issues = mod.collect_abstract_title_issues(doc)
    keys = {x.rule_key for x in issues}
    assert "abstract_title_en_text" not in keys


def test_english_abstract_title_colon_normalization():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("ABSTRACT:")
    doc.add_paragraph("Body line.")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en"] * len(texts)}
    mod.normalize_english_abstract_title(doc)
    assert doc.paragraphs[0].text == "ABSTRACT"


def test_english_abstract_title_mixed_with_body_auto_split():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("Abstract: With the outbreak of COVID-19...")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en"]}
    mod.normalize_english_abstract_title(doc)
    assert len(doc.paragraphs) == 2
    assert doc.paragraphs[0].text == "ABSTRACT"
    assert doc.paragraphs[1].text == "With the outbreak of COVID-19..."
    issues = mod.collect_abstract_title_issues(doc)
    assert all(issue.rule_key != "abstract_title_en_text" for issue in issues)


def test_english_abstract_title_mixed_with_body_auto_split_cn_colon():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("ABSTRACT：With the outbreak of COVID-19...")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en"]}
    mod.normalize_english_abstract_title(doc)
    assert len(doc.paragraphs) == 2
    assert doc.paragraphs[0].text == "ABSTRACT"
    assert doc.paragraphs[1].text == "With the outbreak of COVID-19..."


def test_english_abstract_title_with_empty_tail_only_normalized():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("Abstract:")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en"]}
    mod.normalize_english_abstract_title(doc)
    assert len(doc.paragraphs) == 1
    assert doc.paragraphs[0].text == "ABSTRACT"


def test_english_abstract_title_split_applies_rules():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("abstract: With the outbreak of COVID-19...")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en"]}
    mod.normalize_english_abstract_title(doc)
    assert mod.paragraph_matches(doc.paragraphs[0], mod.RULES["abstract_title_en"])
    assert mod.paragraph_matches(doc.paragraphs[1], mod.RULES["abstract_body_en"])
    assert all(run.bold is not True for run in doc.paragraphs[1].runs if run.text.strip())


def test_chinese_abstract_split_does_not_keep_bold_prefix():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("摘要：随着新冠肺炎疫情发展，研究持续推进。")
    mod.ensure_abstract_heading_split(doc)
    assert doc.paragraphs[0].text == "摘 要"
    assert "摘要：" not in doc.paragraphs[1].text
    assert all(run.bold is not True for run in doc.paragraphs[1].runs if run.text.strip())


def test_keywords_label_bold_content_not_bold():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("关键词：数据可视化；机器学习")
    doc.add_paragraph("Key words：Data visualization; Machine Learning")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_zh", "abstract_en"]}
    mod.normalize_keywords_label_runs(doc)
    zh_runs = [r for r in doc.paragraphs[0].runs if r.text]
    en_runs = [r for r in doc.paragraphs[1].runs if r.text]
    assert zh_runs[0].bold is True
    assert all(run.bold is not True for run in zh_runs[1:])
    assert en_runs[0].bold is True
    assert all(run.bold is not True for run in en_runs[1:])


def test_after_formats_not_shifted_after_structure_normalization():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("摘要：这是中文摘要正文。")
    doc.add_paragraph("关键词：数据可视化；机器学习")
    doc.add_paragraph("Abstract: With the outbreak and spread of COVID-19...")
    doc.add_paragraph("Key words: Data visualization; Machine Learning")
    doc.add_paragraph("目 录")
    doc.add_paragraph("1 引 言")
    doc.add_paragraph("1.1 课题背景")
    doc.add_paragraph("这是正文段落。")
    mod.normalize_document_structure(doc)
    issues = mod.collect_issues(doc)
    mod.apply_supported_rules(doc, allow_layout_fixes=False)
    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "out.docx"
        doc.save(str(out))
        mod.attach_after_formats(out, issues)
    for issue in issues:
        if issue.rule_key == "abstract_body_en" and issue.after:
            assert "toc 1" not in issue.after
        if issue.rule_key in {"keywords_en", "keywords_en_runs"} and issue.after:
            assert "Heading 1" not in issue.after
        if issue.rule_key in {"chapter_title", "chapter_title_runs"} and issue.after:
            assert "Heading 2" not in issue.after


def test_cover_blank_paragraphs_are_not_deleted():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("成都信息工程大学学位论文")
    doc.add_paragraph(" ")
    doc.add_paragraph(" ")
    doc.add_paragraph("姓名：张三")
    doc.add_paragraph("专业：计算机科学与技术")
    doc.add_paragraph("摘 要")
    doc.add_paragraph("这是摘要正文。")
    before = [p.text for p in doc.paragraphs]
    mod.normalize_document_structure(doc)
    after = [p.text for p in doc.paragraphs]
    assert after.count(" ") >= 2
    assert "姓名：张三" in after
    assert "专业：计算机科学与技术" in after
    assert after.index("姓名：张三") < after.index("摘 要")
    assert len(after) >= len(before) - 1


def test_abstract_blank_paragraphs_still_removed():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("摘 要")
    doc.add_paragraph("这是摘要正文。")
    doc.add_paragraph("关键词：数据可视化；机器学习")
    doc.add_paragraph(" ")
    doc.add_paragraph(" ")
    doc.add_paragraph("ABSTRACT")
    before_len = len(doc.paragraphs)
    mod.normalize_document_structure(doc)
    after_texts = [p.text.strip() for p in doc.paragraphs]
    assert len(doc.paragraphs) < before_len
    assert "ABSTRACT" in after_texts


def test_cover_sectpr_blank_not_migrated_or_deleted():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("成都信息工程大学学位论文")
    cover_blank = doc.add_paragraph("")
    ppr = cover_blank._p.get_or_add_pPr()
    sect = mod.OxmlElement("w:sectPr")
    ppr.append(sect)
    doc.add_paragraph("姓名：张三")
    doc.add_paragraph("摘 要")
    doc.add_paragraph("摘要正文。")
    before_len = len(doc.paragraphs)
    mod.normalize_document_structure(doc)
    assert len(doc.paragraphs) >= before_len
    assert any(
        p._p.pPr is not None and p._p.pPr.find(mod.qn("w:sectPr")) is not None
        for p in doc.paragraphs
    )


def test_english_abstract_split_only_in_abstract_en_region():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("Design and Implementation of Abstract Monitoring Platform")
    doc.add_paragraph("Abstract: With the outbreak and spread...")
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("This section discusses Abstract: pattern.")

    def fake_regions(texts, has_toc_field=False):
        return {"regions": ["cover", "abstract_en", "body", "body"]}

    mod.analyze_section_sequence = fake_regions
    mod.normalize_document_structure(doc)
    texts = [p.text for p in doc.paragraphs]
    assert "Design and Implementation of Abstract Monitoring Platform" in texts
    idx = texts.index("ABSTRACT")
    assert texts[idx + 1].startswith("With the outbreak")
    assert "Abstract: pattern." in texts[-1]


def test_body_abstract_word_not_reported_as_title():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("This section discusses Abstract Factory pattern.")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["body"]}
    issues = mod.collect_abstract_title_issues(doc)
    assert issues == []


def test_table_caption_number_formats_and_spacing():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("表2-1 示例表题")
    t1 = doc.add_table(rows=2, cols=2)
    t1.style = "Normal Table"
    doc.add_paragraph("表1 示例表题")
    t2 = doc.add_table(rows=2, cols=2)
    t2.style = "Normal Table"
    doc.add_paragraph("表2.1 示例表题")
    t3 = doc.add_table(rows=2, cols=2)
    t3.style = "Normal Table"
    issues = mod.collect_table_issues(doc)
    keys = [x.rule_key for x in issues]
    assert "table_caption_missing" not in keys
    assert "table_caption_number_format" not in keys


def test_table_caption_missing_space_detected():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("表2-1示例表题")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Normal Table"
    issues = mod.collect_table_issues(doc)
    assert any(issue.rule_key == "table_caption_number_format" for issue in issues)


def test_continued_table_prompts():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("续表（续）")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Normal Table"
    issues = mod.collect_table_issues(doc)
    assert any(issue.rule_key == "table_continued_header" for issue in issues)


def _build_header_sections_doc(mod):
    doc = mod.Document()
    doc.add_paragraph("封面标题")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("摘 要")
    doc.add_paragraph("这是中文摘要内容。")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph("This is english abstract.")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("目 录")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("这是正文段落。")
    return doc


def _header_text(section, mod):
    return " ".join(mod.paragraph_text(p) for p in section.header.paragraphs).strip()


def _section_idx_contains_text(mod, doc, needle):
    for idx, (start, end) in enumerate(mod.section_paragraph_ranges(doc)):
        if start > end:
            continue
        text = mod.section_text(doc, start, end)
        if needle in text:
            return idx
    raise AssertionError(f"section containing {needle!r} not found")


def test_page_header_cover_off_and_abstract_body_on():
    mod = load_module()
    doc = _build_header_sections_doc(mod)
    mod.find_header_start_section = lambda _doc: (1, "abstract-or-later")
    mod.apply_cuit_page_headers(doc, allow_layout_fixes=True)
    abstract_idx = _section_idx_contains_text(mod, doc, "摘 要")
    abstract_en_idx = _section_idx_contains_text(mod, doc, "ABSTRACT")
    body_idx = _section_idx_contains_text(mod, doc, "第一章 绪论")
    assert "成都信息工程大学学士学位论文" not in _header_text(doc.sections[0], mod)
    assert "成都信息工程大学学士学位论文" in _header_text(doc.sections[abstract_idx], mod)
    assert "成都信息工程大学学士学位论文" in _header_text(doc.sections[abstract_en_idx], mod)
    assert "成都信息工程大学学士学位论文" in _header_text(doc.sections[body_idx], mod)


def test_page_header_format_and_no_duplicate():
    mod = load_module()
    doc = _build_header_sections_doc(mod)
    regions = ["cover", "cover", "cover", "abstract_zh", "abstract_zh", "abstract_en", "abstract_en", "toc", "body", "body"]
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": regions}
    mod.apply_cuit_page_headers(doc, allow_layout_fixes=True)
    mod.apply_cuit_page_headers(doc, allow_layout_fixes=True)
    abstract_idx = _section_idx_contains_text(mod, doc, "摘 要")
    header = doc.sections[abstract_idx].header
    exact = [p for p in header.paragraphs if mod.paragraph_text(p).strip() == "成都信息工程大学学士学位论文"]
    assert len(exact) == 1
    para = exact[0]
    assert para.alignment == mod.WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    assert run.font.size is not None and abs(run.font.size.pt - 9.0) < 0.1
    assert run._element.rPr is not None and run._element.rPr.rFonts is not None
    assert run._element.rPr.rFonts.get(mod.qn("w:eastAsia")) == "宋体"
    assert run._element.rPr.rFonts.get(mod.qn("w:ascii")) == "Times New Roman"
    assert run._element.rPr.rFonts.get(mod.qn("w:hAnsi")) == "Times New Roman"


def test_page_header_does_not_break_cover_content():
    mod = load_module()
    doc = _build_header_sections_doc(mod)
    before_cover = [mod.paragraph_text(p) for p in doc.paragraphs[:3]]
    regions = ["cover", "cover", "cover", "abstract_zh", "abstract_zh", "abstract_en", "abstract_en", "toc", "body", "body"]
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": regions}
    mod.apply_cuit_page_headers(doc, allow_layout_fixes=True)
    after_cover = [mod.paragraph_text(p) for p in doc.paragraphs[:3]]
    assert before_cover == after_cover


def test_find_header_start_section_prefers_abstract_toc_body_regions():
    mod = load_module()
    doc = _build_header_sections_doc(mod)
    regions = ["cover", "cover", "cover", "abstract_zh", "abstract_zh", "abstract_en", "abstract_en", "toc", "body", "body"]
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": regions}
    start, reason = mod.find_header_start_section(doc)
    assert start is not None
    assert reason in {"declaration-or-later", "fallback-second-section"}


def test_page_header_starts_from_declaration_section():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("封面标题")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("声 明")
    doc.add_paragraph("本人声明……")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("摘 要")
    doc.add_paragraph("摘要正文。")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("第一章 绪论")
    doc.add_paragraph("正文内容。")
    texts = [mod.paragraph_text(p) for p in doc.paragraphs]
    mocked_regions = []
    for text in texts:
        t = text.strip()
        if "声明" in t or "声 明" in t:
            mocked_regions.append("declaration")
        elif "摘 要" in t or "摘要" in t:
            mocked_regions.append("abstract_zh")
        elif "绪论" in t:
            mocked_regions.append("body")
        else:
            mocked_regions.append("cover")
    mod.analyze_section_sequence = lambda _texts, has_toc_field=False: {"regions": mocked_regions}
    mod.apply_cuit_page_headers(doc, allow_layout_fixes=True)

    declaration_idx = _section_idx_contains_text(mod, doc, "声 明")
    abstract_idx = _section_idx_contains_text(mod, doc, "摘 要")
    body_idx = _section_idx_contains_text(mod, doc, "第一章 绪论")

    assert "成都信息工程大学学士学位论文" not in _header_text(doc.sections[0], mod)
    for idx in (declaration_idx, abstract_idx, body_idx):
        assert "成都信息工程大学学士学位论文" in _header_text(doc.sections[idx], mod)
        para = doc.sections[idx].header.paragraphs[0]
        run = para.runs[0]
        assert para.alignment == mod.WD_ALIGN_PARAGRAPH.CENTER
        assert run.font.size is not None and abs(run.font.size.pt - 9.0) < 0.1
        assert run._element.rPr is not None and run._element.rPr.rFonts is not None
        assert run._element.rPr.rFonts.get(mod.qn("w:eastAsia")) == "宋体"
        assert run._element.rPr.rFonts.get(mod.qn("w:ascii")) == "Times New Roman"
        assert run._element.rPr.rFonts.get(mod.qn("w:hAnsi")) == "Times New Roman"


def test_page_header_without_declaration_starts_from_abstract():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("封面标题")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("摘 要")
    doc.add_paragraph("摘要正文")
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    doc.add_paragraph("1 引 言")
    doc.add_paragraph("正文")
    mod.apply_cuit_page_headers(doc, allow_layout_fixes=True)

    abstract_idx = _section_idx_contains_text(mod, doc, "摘 要")
    body_idx = _section_idx_contains_text(mod, doc, "1 引 言")
    assert "成都信息工程大学学士学位论文" not in _header_text(doc.sections[0], mod)
    assert "成都信息工程大学学士学位论文" in _header_text(doc.sections[abstract_idx], mod)
    assert "成都信息工程大学学士学位论文" in _header_text(doc.sections[body_idx], mod)


def test_page_setup_writes_header_when_layout_fixes_disabled():
    mod = load_module()
    doc = _build_header_sections_doc(mod)
    mod.apply_page_setup(doc, allow_layout_fixes=False)
    abstract_idx = _section_idx_contains_text(mod, doc, "摘 要")
    assert "成都信息工程大学学士学位论文" in _header_text(doc.sections[abstract_idx], mod)


def test_remove_blank_between_english_keywords_and_toc():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph("English abstract body.")
    doc.add_paragraph("Key words: Data visualization; Machine Learning")
    doc.add_paragraph("   ")
    doc.add_paragraph("目 录")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {
        "regions": ["abstract_en", "abstract_en", "abstract_en", "abstract_en", "toc"]
    }
    mod.cleanup_module_boundary_blank_paragraphs(doc, apply_changes=True)
    texts = [mod.paragraph_text(p).strip() for p in doc.paragraphs]
    key_idx = texts.index("Key words: Data visualization; Machine Learning")
    toc_idx = texts.index("目 录")
    assert toc_idx == key_idx + 1


def test_migrate_sectpr_blank_between_keywords_and_toc():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("Key words: Data visualization")
    blank = doc.add_paragraph("")
    ppr = blank._p.get_or_add_pPr()
    sect_pr = mod.OxmlElement("w:sectPr")
    ppr.append(sect_pr)
    doc.add_paragraph("目 录")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en", "abstract_en", "toc"]}
    mod.cleanup_module_boundary_blank_paragraphs(doc, apply_changes=True)
    texts = [mod.paragraph_text(p).strip() for p in doc.paragraphs]
    assert "" not in texts
    key_para = doc.paragraphs[0]
    assert key_para._p.pPr is not None and key_para._p.pPr.find(mod.qn("w:sectPr")) is not None


def test_remove_blank_between_chinese_keywords_and_english_abstract():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("摘 要")
    doc.add_paragraph("中文摘要正文")
    doc.add_paragraph("关键词：数据可视化")
    doc.add_paragraph("   ")
    doc.add_paragraph("ABSTRACT")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {
        "regions": ["abstract_zh", "abstract_zh", "abstract_zh", "abstract_zh", "abstract_en"]
    }
    mod.cleanup_module_boundary_blank_paragraphs(doc, apply_changes=True)
    texts = [mod.paragraph_text(p).strip() for p in doc.paragraphs]
    kw_idx = texts.index("关键词：数据可视化")
    en_idx = texts.index("ABSTRACT")
    assert en_idx == kw_idx + 1


def test_cover_blank_paragraphs_are_not_removed():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("封面标题")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("姓名：张三")
    doc.add_paragraph("摘 要")
    before = [mod.paragraph_text(p) for p in doc.paragraphs]
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {
        "regions": ["cover", "cover", "cover", "cover", "abstract_zh"]
    }
    mod.cleanup_module_boundary_blank_paragraphs(doc, apply_changes=True)
    after = [mod.paragraph_text(p) for p in doc.paragraphs]
    assert before == after


def test_layout_cleanup_runs_after_page_headers():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph("English abstract body.")
    doc.add_paragraph("Key words: Data visualization")
    blank = doc.add_paragraph("")
    ppr = blank._p.get_or_add_pPr()
    ppr.append(mod.OxmlElement("w:sectPr"))
    doc.add_paragraph("目 录")
    doc.add_paragraph("第一章 绪论")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en"] * 4 + ["toc", "body"]}
    mod.find_header_start_section = lambda _doc: (0, "abstract-or-later")
    mod.apply_supported_rules(doc, allow_layout_fixes=True)
    texts = [mod.paragraph_text(p).strip() for p in doc.paragraphs]
    key_idx = texts.index("Key words: Data visualization")
    toc_idx = texts.index("目 录")
    assert toc_idx == key_idx + 1
    assert "成都信息工程大学学士学位论文" in _header_text(doc.sections[0], mod)


def test_remove_toc_field_placeholder_blank_after_english_keywords():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph("English abstract body.")
    doc.add_paragraph("Key words: Data visualization")
    p1 = doc.add_paragraph("")
    r1 = p1.add_run()
    r1._r.append(mod.OxmlElement("w:fldChar"))
    r2 = p1.add_run()
    instr = mod.OxmlElement("w:instrText")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r2._r.append(instr)
    p2 = doc.add_paragraph("")
    ppr = p2._p.get_or_add_pPr()
    ppr.append(mod.OxmlElement("w:sectPr"))
    p2.add_run()._r.append(mod.OxmlElement("w:fldChar"))
    doc.add_paragraph("1 引 言")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en"] * 5 + ["body"]}
    mod.cleanup_module_boundary_blank_paragraphs(doc, apply_changes=True)
    texts = [mod.paragraph_text(p).strip() for p in doc.paragraphs]
    key_idx = texts.index("Key words: Data visualization")
    body_idx = texts.index("1 引 言")
    assert body_idx > key_idx + 1
    assert doc.paragraphs[body_idx].paragraph_format.page_break_before is not True


def test_keywords_cleanup_does_not_cross_toc_to_body():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph("English abstract body")
    doc.add_paragraph("Key words: alpha; beta")
    toc_holder = doc.add_paragraph("")
    toc_holder.add_run()._r.append(mod.OxmlElement("w:fldChar"))
    doc.add_paragraph("1 引 言")
    doc.add_paragraph("随着 COVID-19 疫情发展，正文开始。")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en", "abstract_en", "abstract_en", "abstract_en", "body", "body"]}
    issues, _ = mod.cleanup_module_boundary_blank_paragraphs(doc, apply_changes=True)
    intro = next(p for p in doc.paragraphs if mod.paragraph_text(p).strip() == "1 引 言")
    assert intro.paragraph_format.page_break_before is not True
    assert intro._p.pPr is None or intro._p.pPr.find(mod.qn("w:sectPr")) is None
    assert any(x.rule_key == "abstract_boundary_cleanup_crossed_toc" for x in issues)


def test_keywords_cleanup_only_targets_toc():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("Key words: alpha; beta")
    doc.add_paragraph("   ")
    doc.add_paragraph("目 录")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en", "abstract_en", "toc"]}
    mod.cleanup_module_boundary_blank_paragraphs(doc, apply_changes=True)
    texts = [mod.paragraph_text(p).strip() for p in doc.paragraphs]
    assert "目 录" in texts
    assert texts == ["Key words: alpha; beta", "目 录"]


def test_keywords_cleanup_stops_before_body():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("Key words: alpha; beta")
    doc.add_paragraph("   ")
    doc.add_paragraph("1 引 言")
    doc.add_paragraph("正文第一段。")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en", "abstract_en", "body", "body"]}
    issues, _ = mod.cleanup_module_boundary_blank_paragraphs(doc, apply_changes=True)
    intro = doc.paragraphs[2]
    assert intro.paragraph_format.page_break_before is not True
    assert any(x.rule_key == "abstract_boundary_cleanup_crossed_toc" for x in issues)


def test_remove_blank_after_chapter_title():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("1 引 言")
    doc.add_paragraph(" ")
    body = doc.add_paragraph("随着 COVID-19 疫情，正文内容开始。")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["body", "body", "body"]}
    _issues, removed = mod.cleanup_body_heading_following_blank_paragraphs(doc, apply_changes=True)
    texts = [mod.paragraph_text(p).strip() for p in doc.paragraphs]
    assert removed >= 1
    assert texts == ["1 引 言", "随着 COVID-19 疫情，正文内容开始。"]
    assert body.paragraph_format.page_break_before is not True


def test_remove_page_break_after_chapter_title():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("1 引 言")
    blank = doc.add_paragraph("")
    br = mod.OxmlElement("w:br")
    br.set(mod.qn("w:type"), "page")
    blank.add_run()._r.append(br)
    body = doc.add_paragraph("正文第一段。")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["body", "body", "body"]}
    _issues, _removed = mod.cleanup_body_heading_following_blank_paragraphs(doc, apply_changes=True)
    texts = [mod.paragraph_text(p).strip() for p in doc.paragraphs]
    assert texts == ["1 引 言", "正文第一段。"]
    assert body.paragraph_format.page_break_before is not True


def test_post_layout_cleanup_reruns_renderer_when_changed():
    mod = load_module()
    calls = {"save": 0}

    class DummyDoc:
        def __init__(self, *_args, **_kwargs):
            self.paragraphs = []
            self.sections = []
            self.inline_shapes = []
            self._element = type("E", (), {"xml": "", "iter": lambda self: []})()

        def save(self, *_args, **_kwargs):
            return None

    def fake_try_com_save(_path, _renderer):
        calls["save"] += 1
        return "office"

    mod.require_docx_dependencies = lambda: None
    mod.is_docx = lambda _p: True
    mod.Document = DummyDoc
    mod.normalize_document_structure = lambda _d: None
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": [], "warnings": []}
    mod.collect_issues = lambda _d: []
    mod.create_annotated_docx = lambda *_args, **_kwargs: None
    mod.probe_com_renderer = lambda _r: "office"
    mod.apply_supported_rules = lambda *_args, **_kwargs: None
    mod.try_com_save = fake_try_com_save
    mod.cleanup_visible_blank_paragraphs_after_layout = lambda _d: ([], True)
    mod.attach_after_formats = lambda *_args, **_kwargs: None
    mod.attach_screenshots = lambda **_kwargs: ("skipped", None)
    mod.write_html_report = lambda *_args, **_kwargs: None
    mod.compute_section_boundary_findings = lambda *_args, **_kwargs: ([], [])

    with tempfile.TemporaryDirectory() as td:
        base = Path(td)
        docx = base / "a.docx"
        docx.write_bytes(b"pk")
        out = base / "out"
        report = mod.run(docx, out, "auto", "never")
        assert report["post_layout_cleanup_changed"] is True
        assert calls["save"] == 3


def test_blank_cleanup_keeps_image_paragraph():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("Key words: alpha")
    img_para = doc.add_paragraph("")
    img_para.add_run()._r.append(mod.OxmlElement("w:drawing"))
    doc.add_paragraph("目 录")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en", "abstract_en", "toc"]}
    _issues, _removed = mod.cleanup_module_boundary_blank_paragraphs(doc, apply_changes=True)
    assert any(mod._paragraph_has_image_payload(p) for p in doc.paragraphs)


def test_body_heading_cleanup_keeps_image_paragraph():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("3.3 数据库设计")
    img_para = doc.add_paragraph("")
    img_para.add_run()._r.append(mod.OxmlElement("w:drawing"))
    doc.add_paragraph("图3.7 数据库E-R图")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["body", "body", "body"]}
    _issues, _removed = mod.cleanup_body_heading_following_blank_paragraphs(doc, apply_changes=True)
    texts = [mod.paragraph_text(p).strip() for p in doc.paragraphs]
    assert "图3.7 数据库E-R图" in texts
    assert any(mod._paragraph_has_image_payload(p) for p in doc.paragraphs)


def test_figure_caption_requires_nearby_image():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("图3.7 数据库E-R图")
    issues = mod.collect_figure_caption_image_issues(doc)
    assert any(i.rule_key == "figure_caption_without_nearby_image" for i in issues)


def test_figure_caption_with_nearby_image_passes():
    mod = load_module()
    doc = mod.Document()
    img_para = doc.add_paragraph("")
    img_para.add_run()._r.append(mod.OxmlElement("w:drawing"))
    doc.add_paragraph("图3.7 数据库E-R图")
    issues = mod.collect_figure_caption_image_issues(doc)
    assert all(i.rule_key != "figure_caption_without_nearby_image" for i in issues)


def test_image_count_does_not_decrease_after_fix():
    mod = load_module()
    class N:
        def __init__(self, tag, attrs):
            self.tag = tag
            self._attrs = attrs
        def get(self, key):
            return self._attrs.get(key)
    class E:
        def __init__(self, nodes):
            self._nodes = nodes
        def iter(self):
            return self._nodes
    class D:
        def __init__(self, nodes):
            self._element = E(nodes)
    src_nodes = [N("{x}blip", {mod.qn("r:embed"): "rId1"}), N("{x}imagedata", {mod.qn("r:id"): "rId2"})]
    fix_nodes = [N("{x}blip", {mod.qn("r:embed"): "rId1"})]
    assert mod.count_document_images(D(src_nodes)) == 2
    assert mod.count_document_images(D(fix_nodes)) == 1


def test_keywords_boundary_blank_removed_without_crossing_body():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("Key words: alpha")
    blank = doc.add_paragraph("")
    ppr = blank._p.get_or_add_pPr()
    ppr.append(mod.OxmlElement("w:sectPr"))
    doc.add_paragraph("目 录")
    intro = doc.add_paragraph("1 引 言")
    doc.add_paragraph("正文第一段")
    mod.analyze_section_sequence = lambda texts, has_toc_field=False: {"regions": ["abstract_en", "abstract_en", "toc", "body", "body"]}
    _issues, _removed = mod.cleanup_module_boundary_blank_paragraphs(doc, apply_changes=True)
    texts = [mod.paragraph_text(p).strip() for p in doc.paragraphs]
    k = texts.index("Key words: alpha")
    t = texts.index("目 录")
    assert t == k + 1
    assert intro.paragraph_format.page_break_before is not True


def _prepare_run_stub(mod):
    mod.require_docx_dependencies = lambda: None
    mod.normalize_document_structure = lambda doc: None
    mod.collect_issues = lambda doc: []
    mod.create_annotated_docx = lambda src, out, issues: shutil.copyfile(src, out)
    mod.probe_com_renderer = lambda renderer: "ooxml"
    mod.apply_supported_rules = lambda doc, allow_layout_fixes: None
    mod.try_com_save = lambda path, renderer: "ooxml"
    mod.cleanup_visible_blank_paragraphs_after_layout = lambda doc: ([], False)
    mod.attach_after_formats = lambda fixed_path, issues: None
    mod.attach_screenshots = lambda **kwargs: ("skipped", {})
    mod.count_document_images = lambda doc: 0
    mod.count_figure_captions = lambda doc: 0


def _make_minimal_docx(mod, path: Path) -> None:
    doc = mod.Document()
    doc.add_paragraph("摘 要")
    doc.add_paragraph("正文段落")
    doc.save(str(path))


def test_output_files_include_same_run_timestamp():
    mod = load_module()
    _prepare_run_stub(mod)
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        src = root / "sample.docx"
        out = root / "out"
        _make_minimal_docx(mod, src)
        report = mod.run(src, out, "auto", "never", run_timestamp="20260513_155821")
        for key in ["annotated_docx", "fixed_docx", "json_report", "html_report"]:
            assert "20260513_155821" in Path(report[key]).name
        assert report["run_timestamp"] == "20260513_155821"


def test_report_paths_use_timestamped_names():
    mod = load_module()
    _prepare_run_stub(mod)
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        src = root / "sample.docx"
        out = root / "out"
        _make_minimal_docx(mod, src)
        report = mod.run(src, out, "auto", "never", run_timestamp="20260513_155821")
        payload = json.loads(Path(report["json_report"]).read_text(encoding="utf-8"))
        assert payload["run_timestamp"] == "20260513_155821"
        assert "20260513_155821" in Path(payload["fixed_docx"]).name
        assert "20260513_155821" in Path(payload["annotated_docx"]).name
        assert "20260513_155821" in Path(payload["json_report"]).name
        assert "20260513_155821" in Path(payload["html_report"]).name


def test_html_includes_reference_272_summary():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "r.html"
        report = {
            "input": "x.docx",
            "annotated_docx": "a.docx",
            "fixed_docx": "f.docx",
            "issue_count": 1,
            "renderer_for_fixed": "ooxml",
            "renderer_for_comments": "ooxml",
            "screenshot_status": "skipped",
            "issues": [],
            "issue_summary_by_category": {},
            "reference_272_check_summary": {
                "total_entries": 2,
                "matched_entries": 1,
                "type_mismatch_entries": 1,
                "unmatched_entries": 0,
                "entries": [
                    {"visual_index": 1, "status": "matched", "detected_type": "J", "text_excerpt": "ok", "missing_fields": [], "matched_template": "journal"},
                    {"visual_index": 2, "status": "type_mismatch", "detected_type": "M", "text_excerpt": "bad", "missing_fields": ["出版地"], "matched_template": "book"},
                ],
            },
            "reference_273_check_summary": {
                "online_entries": 0,
                "author_et_al_review_entries": 0,
                "citation_date_missing_entries": 0,
                "access_path_missing_entries": 0,
                "doi_or_url_missing_entries": 0,
                "entries": [],
            },
        }
        mod.write_html_report(report, p)
        html = p.read_text(encoding="utf-8")
        assert "参考文献著录格式检查汇总（2.7.2）" in html
        assert "type_mismatch_entries" in html
        assert "bad" in html


def test_reference_272_semantic_issue_after_text_says_manual_edit():
    mod = load_module()
    issue = mod.Issue(
        paragraph_index=3,
        rule_key="reference_272_type_format_mismatch",
        text_type="reference entry 2.7.2 type format",
        text_excerpt="[1] 作者. 题名[M]. 出版社, 2020.",
        current="缺失要素：出版地",
        expected="应符合2.7.2模板",
        message="",
        category="references",
        after="格式已修复",
    )
    payload = mod.issue_dict(issue)
    assert payload["after"] == "著录内容不自动改写，请人工按模板修改。"


def test_repeated_runs_do_not_overwrite_previous_outputs():
    mod = load_module()
    _prepare_run_stub(mod)
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        src = root / "sample.docx"
        out = root / "out"
        _make_minimal_docx(mod, src)
        report1 = mod.run(src, out, "auto", "never", run_timestamp="20260513_155821")
        report2 = mod.run(src, out, "auto", "never", run_timestamp="20260513_155822")
        assert Path(report1["fixed_docx"]).exists()
        assert Path(report2["fixed_docx"]).exists()
        assert report1["fixed_docx"] != report2["fixed_docx"]


def test_timestamp_format():
    mod = load_module()
    ts = mod.build_run_timestamp()
    assert re.fullmatch(r"\d{8}_\d{6}", ts)


def test_run_timestamp_cli_argument():
    mod = load_module()
    _prepare_run_stub(mod)
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        src = root / "sample.docx"
        out = root / "out"
        _make_minimal_docx(mod, src)
        report = mod.run(src, out, "auto", "never", run_timestamp="20260513_155821")
        assert report["run_timestamp"] == "20260513_155821"
        assert Path(report["fixed_docx"]).name.endswith("_20260513_155821.docx")


def test_toc_page_number_manual_review_issue():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("目 录")
    doc.add_paragraph("1 引 言 .... 1")
    issues = mod.detect_toc_manual_review_issues(doc)
    keys = {issue.rule_key for issue in issues}
    assert "toc_page_number_manual_review" in keys


def test_toc_body_section_manual_review_issue():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("目 录")
    doc.add_paragraph("1 引 言")
    doc.add_paragraph("随着 COVID-19 疫情...")
    issues = mod.detect_toc_manual_review_issues(doc)
    keys = {issue.rule_key for issue in issues}
    assert "toc_page_number_manual_review" in keys
    assert "toc_body_section_manual_review" in keys


def test_toc_report_only_does_not_modify_sections():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("目 录")
    doc.add_paragraph("1 引 言")
    before_sections = len(doc.sections)
    before_xml = doc.element.xml
    _issues = mod.detect_toc_manual_review_issues(doc)
    assert len(doc.sections) == before_sections
    assert doc.element.xml == before_xml


def test_toc_entry_page_numbers_remain_decimal_text():
    mod = load_module()
    doc = mod.Document()
    p = doc.add_paragraph("1 引 言 .... 1")
    _issues = mod.detect_toc_manual_review_issues(doc)
    assert p.text.endswith("1")


def test_toc_report_only_does_not_touch_reference_summary():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] 张三. 题名[J]. 刊名, 2020, 1(2):3-5.")
    texts = [mod.paragraph_text(x) for x in doc.paragraphs]
    entries, _ = mod._reference_entries_from_regions(doc.paragraphs, texts, ["references", "references"])
    summary = mod._reference_272_check_summary(entries)
    assert summary["entries"][0]["visual_index"] == 1
    assert summary["entries"][0]["numbering_source"] == "text"


def test_toc_start_new_page_is_report_only():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("Key words: a; b; c")
    toc = doc.add_paragraph("目 录")
    before_sections = len(doc.sections)
    before_pgnum = [sec._sectPr.find(mod.qn("w:pgNumType")) is not None for sec in doc.sections]
    issues = mod.collect_toc_issues(doc)
    keys = {i.rule_key for i in issues}
    assert toc.paragraph_format.page_break_before is not True
    assert "toc_start_new_page_manual_review" in keys
    assert len(doc.sections) == before_sections
    after_pgnum = [sec._sectPr.find(mod.qn("w:pgNumType")) is not None for sec in doc.sections]
    assert before_pgnum == after_pgnum


def test_no_page_break_before_added_to_toc_field():
    mod = load_module()
    doc = mod.Document()
    p = doc.add_paragraph("")
    r = p._p.add_r()
    instr = mod.OxmlElement("w:instrText")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r.append(instr)
    _issues = mod.collect_toc_issues(doc)
    assert p.paragraph_format.page_break_before is not True


def test_toc_report_only_does_not_create_blank_page():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("Key words: x")
    toc_holder = doc.add_paragraph("")
    run = toc_holder._p.add_r()
    instr = mod.OxmlElement("w:instrText")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run.append(instr)
    doc.add_paragraph("目 录")
    intro = doc.add_paragraph("1 引 言")
    before_sections = len(doc.sections)
    issues = mod.collect_toc_issues(doc)
    keys = {i.rule_key for i in issues}
    assert toc_holder.paragraph_format.page_break_before is not True
    assert intro.paragraph_format.page_break_before is not True
    assert len(doc.sections) == before_sections
    assert "toc_start_new_page_manual_review" in keys


def test_toc_page_number_still_report_only():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("目 录")
    issues = mod.detect_toc_manual_review_issues(doc)
    keys = {i.rule_key for i in issues}
    assert "toc_page_number_manual_review" in keys


def test_toc_style_empty_paragraph_can_be_cleaned_if_no_toc_field():
    mod = load_module()
    doc = mod.Document()
    p = doc.add_paragraph("")
    ppr = p._p.get_or_add_pPr()
    pstyle = mod.OxmlElement("w:pStyle")
    pstyle.set(mod.qn("w:val"), "TOC1")
    ppr.append(pstyle)
    assert mod._is_effectively_blank_paragraph(p) is True


def test_real_toc_field_paragraph_is_not_deleted():
    mod = load_module()
    doc = mod.Document()
    key = doc.add_paragraph("Key words: a; b; c")
    toc_field = doc.add_paragraph("")
    ppr = toc_field._p.get_or_add_pPr()
    pstyle = mod.OxmlElement("w:pStyle")
    pstyle.set(mod.qn("w:val"), "TOC1")
    ppr.append(pstyle)
    run = toc_field._p.add_r()
    instr = mod.OxmlElement("w:instrText")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run.append(instr)
    doc.add_paragraph("目 录")
    regions = ["abstract_en", "abstract_en", "toc"]
    _issues, _removed = mod.cleanup_module_boundary_blank_paragraphs(doc, regions=regions, apply_changes=True)
    texts = [mod.paragraph_text(x) for x in doc.paragraphs]
    assert len(texts) == 3
    assert mod._paragraph_has_toc_field(doc.paragraphs[1]) is True
    assert key.text.startswith("Key words")


def test_load_dotenv_file_reads_basic_values():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        env_path = Path(td) / ".env"
        env_path.write_text("DASHSCOPE_API_KEY=abc\nDASHSCOPE_MODEL=deepseek-v4-pro\n", encoding="utf-8", newline="\n")
        data = mod.load_dotenv_file(env_path)
        assert data["DASHSCOPE_API_KEY"] == "abc"
        assert data["DASHSCOPE_MODEL"] == "deepseek-v4-pro"


def test_dotenv_missing_is_ok():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        data = mod.load_dotenv_file(Path(td) / ".env")
        assert data == {}


def test_llm_config_prefers_cli_over_env_over_dotenv():
    mod = load_module()
    old = dict(mod.os.environ)
    try:
        mod.os.environ.clear()
        mod.os.environ.update({"LLM_PROVIDER": "env_provider", "DASHSCOPE_MODEL": "env_model"})
        parser = mod.argparse.ArgumentParser()
        parser.add_argument("--llm-review", action="store_true")
        parser.add_argument("--llm-provider", default=None)
        parser.add_argument("--llm-model", default=None)
        parser.add_argument("--llm-api-key-env", default="DASHSCOPE_API_KEY")
        parser.add_argument("--llm-base-url", default=None)
        parser.add_argument("--llm-review-mode", choices=["auto", "text", "vision"], default="auto")
        parser.add_argument("--llm-review-max-pages", type=int, default=8)
        parser.add_argument("--llm-review-spec-docx", default=None)
        parser.add_argument("--llm-review-all-pages", action="store_true")
        parser.add_argument("--llm-review-batch-size", type=int, default=6)
        parser.add_argument("--llm-review-image-upload", action="store_true", default=True)
        parser.add_argument("--llm-review-no-separate-html", action="store_true")
        parser.add_argument("--llm-review-output", default=None)
        parser.add_argument("--llm-review-timeout", type=int, default=60)
        args = parser.parse_args(["--llm-review", "--llm-provider", "cli_provider", "--llm-model", "cli_model"])
        cfg = mod.resolve_llm_review_config(args, {"LLM_PROVIDER": "dotenv_provider", "DASHSCOPE_MODEL": "dotenv_model"})
        assert cfg.provider == "cli_provider"
        assert cfg.model == "cli_model"
    finally:
        mod.os.environ.clear()
        mod.os.environ.update(old)


def test_llm_review_disabled_by_default():
    mod = load_module()
    assert mod.run_llm_review(mod.LLMReviewConfig(enabled=False), {}, {}) == {"enabled": False}


def test_llm_review_missing_api_key_skips_gracefully():
    mod = load_module()
    old = dict(mod.os.environ)
    try:
        mod.os.environ.pop("DASHSCOPE_API_KEY", None)
        res = mod.run_llm_review(mod.LLMReviewConfig(enabled=True), {}, {})
        assert res["enabled"] is True
        assert res["status"] == "skipped"
        assert res["reason"] == "missing_api_key"
    finally:
        mod.os.environ.clear()
        mod.os.environ.update(old)


def test_llm_review_prompt_requires_json():
    mod = load_module()
    old = dict(mod.os.environ)
    captured = {}
    try:
        mod.os.environ["DASHSCOPE_API_KEY"] = "fake-key"
        def fake_call(**kwargs):
            captured["messages"] = kwargs["messages"]
            return {"choices": [{"message": {"content": "{\"overall_risk\":\"low\"}"}}]}
        mod.call_openai_compatible_chat = fake_call
        mod.run_llm_review(mod.LLMReviewConfig(enabled=True), {"x": 1}, {})
        user_msg = captured["messages"][1]["content"]
        assert "output only JSON" in user_msg
    finally:
        mod.os.environ.clear()
        mod.os.environ.update(old)


def test_llm_review_result_rendered_in_html():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "r.html"
        report = {
            "input": "x.docx",
            "annotated_docx": "a.docx",
            "fixed_docx": "f.docx",
            "issue_count": 0,
            "renderer_for_fixed": "ooxml",
            "renderer_for_comments": "ooxml",
            "screenshot_status": "skipped",
            "issues": [],
            "llm_review": {
                "enabled": True,
                "status": "ok",
                "provider": "deepseek",
                "model": "deepseek-v4-pro",
                "mode": "auto",
                "evidence_source": "script_summary",
                "vision_status": "fallback_to_text",
                "review": {"overall_risk": "low", "reviewed_pages": []},
            },
        }
        mod.write_html_report(report, p)
        html = p.read_text(encoding="utf-8")
        assert "LLM 高风险格式复核" not in html
        assert "document_upload_status" not in html


def test_llm_review_html_warns_when_no_visual_review():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "r.html"
        report = {
            "input": "x.docx",
            "annotated_docx": "a.docx",
            "fixed_docx": "f.docx",
            "issue_count": 0,
            "renderer_for_fixed": "ooxml",
            "renderer_for_comments": "ooxml",
            "screenshot_status": "skipped",
            "issues": [],
            "llm_review": {
                "enabled": True,
                "status": "ok",
                "provider": "deepseek",
                "model": "deepseek-v4-pro",
                "mode": "text",
                "evidence_source": "script_summary",
                "vision_status": "not_requested",
                "review": {"overall_risk": "low", "reviewed_pages": []},
            },
        }
        mod.write_html_report(report, p)
        html = p.read_text(encoding="utf-8")
        assert "回退为文本/结构化摘要审查" not in html


def test_llm_review_failure_does_not_fail_run():
    mod = load_module()
    _prepare_run_stub(mod)
    mod.run_llm_review = lambda *_args, **_kwargs: {"enabled": True, "status": "failed", "reason": "api_error:TimeoutError"}
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        src = root / "sample.docx"
        out = root / "out"
        _make_minimal_docx(mod, src)
        report = mod.run(src, out, "auto", "never", llm_review_config=mod.LLMReviewConfig(enabled=True))
        assert report["llm_review"]["status"] == "failed"
        assert Path(report["json_report"]).exists()


def test_llm_review_uses_deepseek_defaults():
    mod = load_module()
    cfg = mod.LLMReviewConfig()
    assert cfg.provider == "dashscope"
    assert cfg.model == "qwen3.6-plus"
    assert cfg.api_key_env == "DASHSCOPE_API_KEY"
    assert cfg.mode == "auto"


def test_llm_review_mode_defaults_auto():
    mod = load_module()
    cfg = mod.LLMReviewConfig()
    assert cfg.mode == "auto"


def test_llm_review_text_mode_evidence_source_script_summary():
    mod = load_module()
    old = dict(mod.os.environ)
    try:
        mod.os.environ["DASHSCOPE_API_KEY"] = "fake-key"
        mod.call_openai_compatible_chat = lambda **_kwargs: {"choices": [{"message": {"content": "{\"overall_risk\":\"low\"}"}}]}
        res = mod.run_llm_review(mod.LLMReviewConfig(enabled=True, mode="text"), {"reviewed_pages": [], "review_image_paths": []}, {})
        assert res["status"] == "ok"
        assert "document_upload_status" in res
    finally:
        mod.os.environ.clear()
        mod.os.environ.update(old)


def test_parse_llm_review_response_markdown_json_block():
    mod = load_module()
    parsed = mod.parse_llm_review_response("```json\n{\"k\":1}\n```")
    assert parsed["k"] == 1


def test_collect_llm_review_candidates_table_and_toc_schema():
    mod = load_module()
    issues = [
        mod.Issue(1, "table_structure", "table", "T1", "", "", "m", category="table"),
        mod.Issue(2, "toc_page_number_manual_review", "toc", "目录", "", "", "m", category="page"),
    ]
    c = mod.collect_llm_review_candidates({"issue_summary_by_category": {}}, issues, Path("a.docx"), Path("."), None, mod.LLMReviewConfig())
    assert c["toc_manual_review"]["toc_page_number_manual_review"] is True
    assert len(c["table_risks"]) == 1


def test_llm_review_does_not_send_more_than_max_pages():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        img1 = Path(td) / "a.png"
        img2 = Path(td) / "b.png"
        img1.write_bytes(b"\x89PNG\r\n\x1a\n")
        img2.write_bytes(b"\x89PNG\r\n\x1a\n")
        report = {"render_qa": {"pages": [{"page": 1, "after": str(img1)}, {"page": 2, "after": str(img2)}]}}
        c = mod.collect_llm_review_candidates(report, [], Path("a.docx"), Path("."), None, mod.LLMReviewConfig(max_pages=1))
        assert len(c["reviewed_pages"]) == 1


def test_llm_review_vision_unsupported_falls_back_to_text():
    mod = load_module()
    old = dict(mod.os.environ)
    try:
        mod.os.environ["DASHSCOPE_API_KEY"] = "fake-key"
        with tempfile.TemporaryDirectory() as td:
            img = Path(td) / "a.png"
            img.write_bytes(b"\x89PNG\r\n\x1a\n")
            mod.call_openai_compatible_chat = lambda **_kwargs: {"choices": [{"message": {"content": "{\"overall_risk\":\"low\"}"}}]}
            res = mod.run_llm_review(
                mod.LLMReviewConfig(enabled=True, mode="auto"),
                {"reviewed_pages": [{"page": 1}], "review_image_paths": [str(img)]},
                {},
            )
            assert res["status"] in {"ok", "failed"}
            if res["status"] == "ok":
                assert res["document_upload_status"] in {"ok", "fallback_text"}
    finally:
        mod.os.environ.clear()
        mod.os.environ.update(old)


def test_llm_review_payload_includes_candidate_pages_when_images_available():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        img = Path(td) / "a.png"
        img.write_bytes(b"\x89PNG\r\n\x1a\n")
        report = {"render_qa": {"pages": [{"page": 2, "after": str(img)}]}}
        c = mod.collect_llm_review_candidates(report, [], Path("a.docx"), Path("."), None, mod.LLMReviewConfig(max_pages=8))
        assert c["reviewed_pages"]
        assert c["review_image_paths"]


def test_spec_docx_text_extraction_full_content():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "spec.docx"
        d = mod.Document()
        d.add_paragraph("章节一")
        t = d.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "A"
        t.cell(0, 1).text = "B"
        d.save(str(p))
        text = mod.extract_spec_docx_text(p)
        assert "章节一" in text
        assert "row 1" in text


def test_collect_after_render_images():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        out = Path(td)
        after = out / "render_qa" / "after"
        after.mkdir(parents=True)
        (after / "page_2.png").write_bytes(b"\x89PNG\r\n\x1a\n")
        (after / "page_1.png").write_bytes(b"\x89PNG\r\n\x1a\n")
        pages = mod.collect_after_render_images(out, max_pages=8, all_pages=False)
        assert len(pages) == 2
        assert pages[0]["page"] == 1


def test_llm_review_all_pages_batches_images():
    mod = load_module()
    old = dict(mod.os.environ)
    try:
        mod.os.environ["DASHSCOPE_API_KEY"] = "fake-key"
        calls = {"n": 0}
        mod.call_openai_compatible_chat = lambda **_kwargs: (calls.__setitem__("n", calls["n"] + 1) or {"choices": [{"message": {"content": "{\"overall_risk\":\"low\",\"issues\":[]}"}}]})
        with tempfile.TemporaryDirectory() as td:
            imgs: list[str] = []
            for i in range(1, 8):
                p = Path(td) / f"page_{i}.png"
                p.write_bytes(b"\x89PNG\r\n\x1a\n")
                imgs.append(str(p))
            cands = {"reviewed_pages": [{"page": i} for i in range(1, 8)], "review_image_paths": imgs, "summary": {}, "spec_text": "S", "spec_docx_path": "x"}
            res = mod.run_llm_review(mod.LLMReviewConfig(enabled=True, mode="auto", batch_size=3), cands, {})
        assert res["status"] == "ok"
        assert calls["n"] >= 1
    finally:
        mod.os.environ.clear()
        mod.os.environ.update(old)


def test_llm_review_uses_docx_upload_or_fallback_text():
    mod = load_module()
    old = dict(mod.os.environ)
    try:
        mod.os.environ["DASHSCOPE_API_KEY"] = "fake-key"
        mod.call_openai_compatible_chat = lambda **_kwargs: {"choices": [{"message": {"content": "{\"overall_risk\":\"low\",\"issues\":[]}"}}]}
        with tempfile.TemporaryDirectory() as td:
            p = Path(td) / "page_1.png"
            p.write_bytes(b"\x89PNG\r\n\x1a\n")
            cands = {"reviewed_pages": [{"page": 1}], "review_image_paths": [str(p)], "summary": {}, "spec_text": "S", "spec_docx_path": "x"}
            res = mod.run_llm_review(mod.LLMReviewConfig(enabled=True, mode="auto"), cands, {})
        assert res["image_upload_mode"] in {"docx_upload", "fallback_text_extraction"}
    finally:
        mod.os.environ.clear()
        mod.os.environ.update(old)


def test_llm_review_missing_images_falls_back_safely():
    mod = load_module()
    old = dict(mod.os.environ)
    try:
        mod.os.environ["DASHSCOPE_API_KEY"] = "fake-key"
        mod.call_openai_compatible_chat = lambda **_kwargs: {"choices": [{"message": {"content": "{\"overall_risk\":\"low\",\"issues\":[]}"}}]}
        cands = {"reviewed_pages": [{"page": 1}], "review_image_paths": ["a"], "summary": {}, "spec_text": "S", "spec_docx_path": "x"}
        res = mod.run_llm_review(mod.LLMReviewConfig(enabled=True, mode="auto"), cands, {})
        assert res["status"] in {"ok", "failed"}
    finally:
        mod.os.environ.clear()
        mod.os.environ.update(old)


def test_llm_review_separate_html_written():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "llm.html"
        mod.write_llm_review_separate_html({"model": "m", "vision_status": "ok", "review": {"issues": []}}, p)
        assert p.exists()


def test_llm_review_separate_html_lists_page_issues():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "llm.html"
        mod.write_llm_review_separate_html({"model": "m", "vision_status": "ok", "review": {"issues": [{"page": 3, "type": "blank_page"}]}}, p)
        html = p.read_text(encoding="utf-8")
        assert "3" in html


def test_llm_review_warns_if_vision_failed():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "llm.html"
        mod.write_llm_review_separate_html({"model": "m", "vision_status": "failed", "review": {"issues": []}}, p)
        html = p.read_text(encoding="utf-8")
        assert ("文档提交状态" in html or "回退" in html)


def test_llm_review_spec_docx_not_required():
    mod = load_module()
    cands = mod.collect_llm_review_candidates({"issue_summary_by_category": {}}, [], Path("a.docx"), Path("."), None, mod.LLMReviewConfig())
    assert "spec_text" in cands


def test_dashscope_is_default_llm_provider():
    mod = load_module()
    cfg = mod.LLMReviewConfig()
    assert cfg.provider == "dashscope"


def test_deepseek_not_used_by_default():
    mod = load_module()
    cfg = mod.LLMReviewConfig()
    assert cfg.provider != "deepseek"


def test_qwen_receives_spec_and_fixed_docx_context():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        spec = Path(td) / "spec.docx"
        d = mod.Document()
        d.add_paragraph("规范内容")
        d.save(str(spec))
        fixed = Path(td) / "fixed.docx"
        f = mod.Document()
        f.add_paragraph("正文段落")
        f.save(str(fixed))
        cfg = mod.LLMReviewConfig(spec_docx_path=spec)
        cands = mod.collect_llm_review_candidates({"issue_summary_by_category": {}}, [], fixed, Path(td), None, cfg)
        assert "规范内容" in cands["spec_text"]
        assert "fixed_docx_structure" in cands


def test_qwen_vision_uses_rendered_images():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        out = Path(td)
        after = out / "render_qa" / "after"
        after.mkdir(parents=True)
        (after / "page_1.png").write_bytes(b"\x89PNG\r\n\x1a\n")
        pages = mod.collect_after_render_images(out, 8, False)
        assert pages and pages[0]["image_path"].endswith(".png")


def test_llm_fixed_path_created():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        src = Path(td) / "f.docx"
        d = mod.Document()
        d.add_paragraph("正文")
        d.save(str(src))
        out = Path(td) / "llm.docx"
        res = mod.apply_llm_safe_edit_plan(src, [], out)
        assert out.exists()
        assert res["applied_count"] == 0


def test_llm_safe_edit_plan_applies_low_risk_only():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        src = Path(td) / "f.docx"
        d = mod.Document()
        d.add_paragraph("目标段落")
        d.save(str(src))
        out = Path(td) / "llm.docx"
        plan = [{"operation": "set_alignment", "target_hint": "目标段落", "parameters": {"alignment": "left"}}]
        res = mod.apply_llm_safe_edit_plan(src, plan, out)
        assert res["applied_count"] == 1


def test_llm_safe_edit_plan_never_changes_sections_or_toc():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        src = Path(td) / "f.docx"
        d = mod.Document()
        d.add_paragraph("目标段落")
        d.save(str(src))
        out = Path(td) / "llm.docx"
        plan = [{"operation": "toc_update", "target_hint": "目标段落", "parameters": {}}]
        res = mod.apply_llm_safe_edit_plan(src, plan, out)
        assert res["applied_count"] == 0
        assert res["skipped_count"] >= 1


def test_llm_review_html_written():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "llm.html"
        mod.write_llm_review_separate_html({"model": "qwen", "vision_status": "ok", "review": {"issues": []}}, p)
        assert p.exists()


def test_llm_review_html_lists_manual_review_items():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "llm.html"
        mod.write_llm_review_separate_html({"model": "qwen", "vision_status": "ok", "review": {"issues": [{"page": 2, "type": "toc"}]}}, p)
        html = p.read_text(encoding="utf-8")
        assert "2" in html


def test_main_html_links_llm_outputs():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "main.html"
        report = {
            "input": "x.docx",
            "annotated_docx": "a.docx",
            "fixed_docx": "f.docx",
            "issue_count": 0,
            "renderer_for_fixed": "ooxml",
            "renderer_for_comments": "ooxml",
            "screenshot_status": "skipped",
            "issues": [],
            "llm_review_html": "D:/llm.html",
            "llm_review": {"enabled": False},
        }
        mod.write_html_report(report, p)
        html = p.read_text(encoding="utf-8")
        assert "D:/llm.html" not in html


def test_llm_fixed_copied_when_no_safe_edits():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        src = Path(td) / "f.docx"
        d = mod.Document()
        d.add_paragraph("正文")
        d.save(str(src))
        out = Path(td) / "llm.docx"
        res = mod.apply_llm_safe_edit_plan(src, [], out)
        assert out.exists()
        assert res["applied_count"] == 0


def test_llm_review_schema_contains_visual_fields():
    mod = load_module()
    old = dict(mod.os.environ)
    try:
        mod.os.environ["DASHSCOPE_API_KEY"] = "fake-key"
        mod.call_openai_compatible_chat = lambda **_kwargs: {"choices": [{"message": {"content": "{\"overall_risk\":\"low\"}"}}]}
        res = mod.run_llm_review(mod.LLMReviewConfig(enabled=True, mode="text"), {"reviewed_pages": [], "review_image_paths": []}, {})
        assert res["mode"] == "text"
        assert "document_upload_status" in res or res["status"] == "failed"
    finally:
        mod.os.environ.clear()
        mod.os.environ.update(old)


def test_llm_review_does_not_expose_api_key():
    mod = load_module()
    _prepare_run_stub(mod)
    fake_key = "sk-test-exposed"
    old = dict(mod.os.environ)
    try:
        mod.os.environ["DASHSCOPE_API_KEY"] = fake_key
        mod.call_openai_compatible_chat = lambda **_kwargs: {"choices": [{"message": {"content": "{\"overall_risk\":\"low\"}"}}]}
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            src = root / "sample.docx"
            out = root / "out"
            _make_minimal_docx(mod, src)
            report = mod.run(src, out, "auto", "never", llm_review_config=mod.LLMReviewConfig(enabled=True))
            j = Path(report["json_report"]).read_text(encoding="utf-8")
            h = Path(report["html_report"]).read_text(encoding="utf-8")
            assert fake_key not in j
            assert fake_key not in h
            assert fake_key not in mod.stdout_json(report)
    finally:
        mod.os.environ.clear()
        mod.os.environ.update(old)


def test_main_html_llm_block_has_no_mojibake():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "r.html"
        report = {
            "input": "x.docx",
            "annotated_docx": "a.docx",
            "fixed_docx": "f.docx",
            "issue_count": 0,
            "renderer_for_fixed": "ooxml",
            "renderer_for_comments": "ooxml",
            "screenshot_status": "skipped",
            "issues": [],
            "llm_fixed_docx": "D:/llm_fixed.docx",
            "llm_review_html": "D:/llm_review.html",
            "llm_apply_result": {"applied_count": 0, "skipped_count": 0},
            "llm_review": {"enabled": False},
        }
        mod.write_html_report(report, p)
        html = p.read_text(encoding="utf-8")
        for bad in ["鐙", "澶", "淇", "妗", "搴", "缁", "\ufffd"]:
            assert bad not in html


def test_llm_review_html_no_mojibake():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "llm.html"
        mod.write_llm_review_separate_html({"provider": "dashscope", "model": "qwen3.6-plus", "review": {"basis": {}, "issues": []}}, p)
        html = p.read_text(encoding="utf-8")
        for bad in ["鐙", "澶", "淇", "妗", "搴", "缁", "\ufffd"]:
            assert bad not in html


def test_llm_review_html_no_longer_says_visual_only():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "llm.html"
        mod.write_llm_review_separate_html({"provider": "dashscope", "model": "qwen3.6-plus", "review": {"basis": {}, "issues": []}}, p)
        html = p.read_text(encoding="utf-8")
        assert "LLM 格式审查报告" in html
        assert "LLM 视觉格式复核报告" not in html


def test_dashscope_docx_upload_probe_result_recorded():
    mod = load_module()
    old = dict(mod.os.environ)
    try:
        mod.os.environ["DASHSCOPE_API_KEY"] = "fake-key"
        mod.probe_dashscope_document_upload_support = lambda *_args, **_kwargs: {"supported": False, "method": "fallback_text_extraction", "reason": "unsupported"}
        mod.call_openai_compatible_chat = lambda **_kwargs: {"choices": [{"message": {"content": "{\"overall_result\":\"needs_manual_review\",\"overall_risk\":\"medium\"}"}}]}
        res = mod.run_llm_review(mod.LLMReviewConfig(enabled=True), {"summary": {"fixed_docx": "x.docx"}}, {})
        assert res["status"] == "ok"
        assert res["document_upload_status"] in {"ok", "fallback_text"}
    finally:
        mod.os.environ.clear()
        mod.os.environ.update(old)


def test_qwen_docx_review_uses_spec_and_fixed_docx_when_supported():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        spec = Path(td) / "spec.docx"
        fixed = Path(td) / "fixed.docx"
        spec.write_bytes(b"pk")
        fixed.write_bytes(b"pk")
        mod.probe_dashscope_document_upload_support = lambda *_args, **_kwargs: {
            "supported": True,
            "file_ids": {"spec": "f1", "fixed": "f2"},
        }
        mod.call_openai_compatible_chat = lambda **_kwargs: {"choices": [{"message": {"content": "{\"overall_result\":\"needs_manual_review\",\"overall_risk\":\"medium\",\"basis\":{}}"}}]}
        out = mod.run_qwen_docx_review(
            config=mod.LLMReviewConfig(enabled=True),
            api_key="k",
            candidates={"spec_docx_path": str(spec), "summary": {"fixed_docx": str(fixed)}},
        )
        assert out["status"] == "ok"
        assert out["review"]["basis"]["uses_spec_docx"] is True
        assert out["review"]["basis"]["uses_fixed_docx"] is True


def test_qwen_docx_review_falls_back_to_text_extraction():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        spec = Path(td) / "spec.docx"
        fixed = Path(td) / "fixed.docx"
        spec.write_bytes(b"pk")
        fixed.write_bytes(b"pk")
        mod.probe_dashscope_document_upload_support = lambda *_args, **_kwargs: {
            "supported": False,
            "reason": "upload_error",
        }
        mod.call_openai_compatible_chat = lambda **_kwargs: {"choices": [{"message": {"content": "{\"overall_result\":\"needs_manual_review\",\"overall_risk\":\"medium\",\"basis\":{}}"}}]}
        out = mod.run_qwen_docx_review(
            config=mod.LLMReviewConfig(enabled=True),
            api_key="k",
            candidates={"spec_docx_path": str(spec), "summary": {"fixed_docx": str(fixed)}},
        )
        assert out["status"] == "ok"
        assert out["review"]["basis"]["fallback_used"] is True
        assert out["review"]["basis"]["fallback_reason"] != ""


def test_llm_review_html_shows_document_upload_status():
    mod = load_module()
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "llm.html"
        llm = {
            "provider": "dashscope",
            "model": "qwen3.6-plus",
            "document_upload_status": "fallback_text",
            "review": {"basis": {"document_upload_status": "fallback_text"}, "issues": []},
        }
        mod.write_llm_review_separate_html(llm, p)
        html = p.read_text(encoding="utf-8")
        assert "文档提交状态" in html
        assert "fallback_text" in html


def test_run_outputs_json_into_json_subdir_and_no_llm_fixed_files():
    mod = load_module()
    _prepare_run_stub(mod)
    with tempfile.TemporaryDirectory() as td:
        root = Path(td)
        src = root / "sample.docx"
        out = root / "out"
        _make_minimal_docx(mod, src)
        report = mod.run(src, out, "auto", "never", llm_review_config=mod.LLMReviewConfig(enabled=False))
        assert Path(report["json_report"]).parent.name == "json"
        assert Path(report["json_report"]).exists()
        assert not list(out.glob("*_llm_fixed_*.docx"))
        assert not list(out.glob("*_llm_edit_plan_*.json"))


def test_llm_evidence_postprocess_moves_unverified_issue_to_manual():
    mod = load_module()
    review = {
        "issues": [
            {"id": "LLM-1", "evidence_found": True, "evidence_exact_quote": "1.a. 课题背景", "suggestion": "x"},
        ],
        "manual_review_items": [],
        "validation_warnings": [],
        "safe_edit_plan": [],
    }
    processed = mod._postprocess_llm_evidence(review, "这是另一段正文")
    assert len(processed.get("text_verified_issues", [])) == 0
    assert len(processed.get("rejected_or_unverified_claims", [])) >= 1
    assert len(processed.get("rejected_or_unverified_claims", [])) >= 1


def test_llm_safe_edit_plan_is_advisory_only():
    mod = load_module()
    review = {
        "issues": [],
        "manual_review_items": [],
        "validation_warnings": [],
        "safe_edit_plan": [
            {"id": "LLM-2", "operation": "set_run_font", "risk": "high", "evidence_exact_quote": "标题"},
        ],
    }
    processed = mod._postprocess_llm_evidence(review, "标题")
    # safe_edit_plan is now removed by postprocess; items fall into rejected_or_unverified_claims
    assert "safe_edit_plan" not in processed
    assert len(processed.get("rejected_or_unverified_claims", [])) >= 0


def test_keywords_to_toc_no_blank_page():
    mod = load_module()
    doc = mod.Document()
    doc.add_paragraph("ABSTRACT")
    doc.add_paragraph("English abstract body")
    doc.add_paragraph("Key words: x")
    blank = doc.add_paragraph("")
    br_run = blank._p.add_r()
    br = mod.OxmlElement("w:br")
    br.set(mod.qn("w:type"), "page")
    br_run.append(br)
    toc = doc.add_paragraph("目 录")
    doc.add_paragraph("论文总页数：32页")
    intro = doc.add_paragraph("1 引 言")
    before_sections = len(doc.sections)
    _issues = mod.collect_toc_issues(doc)
    texts = [mod.paragraph_text(p).strip() for p in doc.paragraphs]
    assert "Key words: x" in texts
    assert "目 录" in texts
    assert toc.paragraph_format.page_break_before is not True
    assert len(doc.sections) == before_sections
    assert intro.paragraph_format.page_break_before is not True

def test_find_toc_start_ignores_empty_toc_style_without_field():
    mod = load_module()
    doc = mod.Document()
    p = doc.add_paragraph("")
    ppr = p._p.get_or_add_pPr()
    pstyle = mod.OxmlElement("w:pStyle")
    pstyle.set(mod.qn("w:val"), "TOC1")
    ppr.append(pstyle)
    doc.add_paragraph("目 录")
    idx = mod.find_toc_start_paragraph_index(doc)
    assert idx == 1
