#!/usr/bin/env python3
"""Build a split-table DOCX review map for CUIT thesis format rule locations."""

from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SCRIPT_DIR = Path(__file__).resolve().parent
ROOT = SCRIPT_DIR.parents[1]
BASE_BUILDER = SCRIPT_DIR / "build_rule_position_review_doc.py"
OUTPUT_DOCX = ROOT / "cuit-thesis-rule-position-review.docx"


def load_base_builder():
    spec = importlib.util.spec_from_file_location("base_rule_position_builder", BASE_BUILDER)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    spec.loader.exec_module(module)
    return module


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.line_spacing = 1.05
    normal.paragraph_format.space_after = Pt(2)

    for style_name, size, color in [
        ("Title", 19, "1F4E79"),
        ("Heading 1", 14, "1F4E79"),
        ("Heading 2", 11, "365F91"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    for idx, (text, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[idx]
        cell.width = Cm(width)
        cell.text = text
        set_cell_shading(cell, "D9EAF7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for idx, (text, width) in enumerate(zip(row, widths)):
            cell = cells[idx]
            cell.width = Cm(width)
            cell.text = text
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
    doc.add_paragraph()


def table_rows(table) -> list[list[str]]:
    return [[cell.text for cell in row.cells] for row in table.rows]


def main() -> None:
    # Rebuild the base document first so the source tables come from the UTF-8 builder.
    load_base_builder().main()
    base_doc = Document(str(OUTPUT_DOCX))
    structure_rows = table_rows(base_doc.tables[0])[1:]
    rule_rows_full = table_rows(base_doc.tables[1])
    rule_header = rule_rows_full[0]
    rule_rows = rule_rows_full[1:]
    bug_rows = table_rows(base_doc.tables[2])[1:]
    checklist = [p.text for p in base_doc.paragraphs if p.text.startswith("□ ")]

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("成都信息工程大学学士学位论文格式规则位置审核稿").bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("来源文档：成都信息工程大学学士学位论文规范.docx").italic = True

    note = doc.add_paragraph()
    note.add_run("审核重点：").bold = True
    note.add_run(
        "本稿只说明各格式规则在规范中的适用位置和来源依据，供人工审核后再更新 checker 规则。"
        "当前已确认“论文提交日期”仅属于封面项目，不应在正文中按封面日期规则处理。"
    )

    doc.add_heading("一、结构位置边界", level=1)
    add_table(doc, ["位置", "规范来源", "包含内容", "分类边界"], structure_rows, [2.2, 4.2, 8.0, 8.4])

    doc.add_heading("二、格式规则来源清单", level=1)
    groups = [
        ("2.1 基础、封面、摘要、目录规则", rule_rows[:10]),
        ("2.2 正文、图表、公式、页眉页码规则", rule_rows[10:19]),
        ("2.3 后置部分规则", rule_rows[19:]),
    ]
    for index, (heading, rows) in enumerate(groups):
        doc.add_heading(heading, level=2)
        add_table(doc, rule_header, rows, [2.5, 2.0, 5.0, 8.8, 6.8])
        if index == 0:
            doc.add_page_break()

    doc.add_heading("三、当前误判说明与拟调整点", level=1)
    add_table(doc, ["问题", "当前证据", "应使用的规范位置", "建议调整"], bug_rows, [5.0, 6.2, 7.2, 7.2])

    doc.add_heading("四、待人工审核的问题", level=1)
    for item in checklist:
        paragraph = doc.add_paragraph()
        paragraph.add_run("□ ").bold = True
        paragraph.add_run(item[2:] if item.startswith("□ ") else item)

    final_note = doc.add_paragraph()
    final_note.add_run("下一步建议：").bold = True
    final_note.add_run(
        "请先在本审核稿中确认每条规则的位置边界。确认后，再更新 rules.json 与分类逻辑，"
        "重点修复 cover_date 只能出现在封面字段的问题。"
    )

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
