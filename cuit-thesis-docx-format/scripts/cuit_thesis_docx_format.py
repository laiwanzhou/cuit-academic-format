#!/usr/bin/env python3
"""Check and fix bachelor thesis DOCX formatting."""

from __future__ import annotations

import argparse
import base64
from copy import deepcopy
import html
import json
import mimetypes
import os
import re
import shutil
import subprocess
import sys
import urllib.error
import urllib.request
import zipfile
from dataclasses import dataclass, replace
from datetime import datetime, timezone
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Iterable
from xml.etree import ElementTree as ET

SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from dashscope_doc_review import (  # type: ignore
    delete_uploaded_file as dashscope_delete_uploaded_file,
    get_dashscope_client,
    resolve_spec_file_id,
    run_qwen_long_docx_review,
    upload_target_docx,
)

try:
    from docx import Document
    from docx.enum.section import WD_SECTION_START
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ModuleNotFoundError as exc:
    Document = None
    WD_SECTION_START = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    qn = None
    Cm = None
    Pt = None
    DOCX_IMPORT_ERROR = exc
else:
    DOCX_IMPORT_ERROR = None


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
PKG_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CONTENT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"

ET.register_namespace("w", W_NS)
ET.register_namespace("r", R_NS)
ET.register_namespace("", CONTENT_NS)


if WD_ALIGN_PARAGRAPH is None:
    ALIGN_LABELS = {}
    ALIGN_LABELS_ZH = {}
    ALIGNMENT_VALUES = {"left": None, "center": None, "right": None, "justified": None, None: None}
else:
    ALIGN_LABELS = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justified",
    }

    ALIGN_LABELS_ZH = {
        WD_ALIGN_PARAGRAPH.LEFT: "左对齐",
        WD_ALIGN_PARAGRAPH.CENTER: "居中",
        WD_ALIGN_PARAGRAPH.RIGHT: "右对齐",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "两端对齐",
    }

    ALIGNMENT_VALUES = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
        None: None,
    }

ALIGN_LEFT = ALIGNMENT_VALUES["left"]
ALIGN_CENTER = ALIGNMENT_VALUES["center"]
ALIGN_RIGHT = ALIGNMENT_VALUES["right"]
ALIGN_JUSTIFIED = ALIGNMENT_VALUES["justified"]

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
RULE_CONFIG_PATH = SKILL_DIR / "references" / "rules.json"
DEFAULT_DOTENV_PATH = SKILL_DIR.parent / ".env"


def require_docx_dependencies() -> None:
    if DOCX_IMPORT_ERROR is None:
        return
    raise RuntimeError(
        "Missing required DOCX dependency. Run `python -m pip install -r requirements.txt` "
        "from the cuit-thesis-docx-format skill directory, then retry."
    ) from DOCX_IMPORT_ERROR


@dataclass(frozen=True)
class Rule:
    key: str
    label: str
    font_east_asia: str
    font_ascii: str
    size_pt: float
    bold: bool | None
    alignment: int | None
    line_spacing_pt: float | None = 20
    first_line_indent_cm: float | None = None
    first_line_indent_chars: float | None = None
    hanging_indent_chars: float | None = None
    space_before_pt: float | None = None
    space_after_pt: float | None = None
    space_before_lines: float | None = None
    space_after_lines: float | None = None
    expected: str = ""
    category: str = "format"


@dataclass
class Issue:
    paragraph_index: int
    rule_key: str
    text_type: str
    text_excerpt: str
    current: str
    expected: str
    message: str
    category: str = "format"
    location: str | None = None
    after: str | None = None
    page: int | None = None
    before_screenshot: str | None = None
    after_screenshot: str | None = None
    screenshot_note: str | None = None


@dataclass(frozen=True)
class LLMReviewConfig:
    enabled: bool = False
    provider: str = "dashscope"
    review_model: str = "qwen3.6-plus"
    doc_model: str = "qwen-long"
    model: str = "qwen3.6-plus"
    api_key_env: str = "DASHSCOPE_API_KEY"
    base_url: str = "https://dashscope.aliyuncs.com/compatible-mode/v1"
    max_pages: int = 8
    timeout: int = 600
    mode: str = "auto"
    spec_docx_path: Path | None = None
    all_pages: bool = False
    batch_size: int = 6
    separate_html: bool = True
    image_upload: bool = True
    image_transport: str = "data_uri"
    spec_file_id: str | None = None
    doc_target: str = "fixed"
    upload_target_docx: bool = True
    delete_uploaded_target_after_review: bool = False
    output_path: Path | None = None


LLM_REVIEW_SYSTEM_PROMPT = (
    "你是论文格式高风险复核助手。你只做格式风险复核，不直接修改文档。请根据给定的脚本诊断摘要，判断是否存在需要人工复核的高风险问题，"
    "包括空白页、空白段、目录页码、目录是否另起页、三线表、图题/表题位置、图片是否丢失、续表表头、参考文献著录格式风险。"
    "你必须只输出 JSON，不要输出 Markdown，不要输出解释性前后缀。看不清或证据不足时，使用 not_sure 或 manual_review_needed，不要编造。"
    "特别注意：1. 目录条目右侧页码应为阿拉伯数字，例如 1、3、7、12。2. 目录页底部页码属于前置页，应为小写罗马数字，例如 iii、iv。"
    "3. 三线表一般只保留顶线、表头下线、底线，不应有竖线和全网格线。"
    "4. LLM 复核结果只作为人工检查参考，不能替代脚本确定性检查。"
    "5. 你不能要求自动修改 section、TOC field、PAGE field、表格边框或图片位置。"
)


def load_dotenv_file(path: Path) -> dict[str, str]:
    if not path.exists():
        return {}
    values: dict[str, str] = {}
    for raw_line in path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip()
        if (value.startswith('"') and value.endswith('"')) or (value.startswith("'") and value.endswith("'")):
            value = value[1:-1]
        if key:
            values[key] = value
    return values


def get_config_value(name: str, cli_value: str | None, env: dict[str, str], dotenv: dict[str, str], default: str) -> str:
    if cli_value is not None:
        return cli_value
    if env.get(name):
        return env[name]
    if dotenv.get(name):
        return dotenv[name]
    return default


def _safe_positive_int(value: str | None, default: int) -> int:
    try:
        parsed = int(value or "")
    except ValueError:
        return default
    return parsed if parsed > 0 else default


def resolve_llm_review_config(args: argparse.Namespace, dotenv: dict[str, str]) -> LLMReviewConfig:
    env = os.environ
    default_spec = SKILL_DIR.parent / "成都信息工程大学学士学位论文规范.docx"
    spec_path = Path(args.llm_review_spec_docx) if args.llm_review_spec_docx else (default_spec if default_spec.exists() else None)
    review_model = (
        getattr(args, "llm_review_model", None)
        or getattr(args, "llm_model", None)
        or env.get("DASHSCOPE_REVIEW_MODEL")
        or dotenv.get("DASHSCOPE_REVIEW_MODEL")
        or env.get("DASHSCOPE_MODEL")
        or dotenv.get("DASHSCOPE_MODEL")
        or "qwen3.6-plus"
    )
    doc_model = (
        getattr(args, "llm_doc_model", None)
        or env.get("DASHSCOPE_DOC_MODEL")
        or dotenv.get("DASHSCOPE_DOC_MODEL")
        or "qwen-long"
    )
    return LLMReviewConfig(
        enabled=bool(args.llm_review),
        provider=get_config_value("LLM_PROVIDER", args.llm_provider, env, dotenv, "dashscope"),
        review_model=review_model,
        doc_model=doc_model,
        model=review_model,
        api_key_env=args.llm_api_key_env or "DASHSCOPE_API_KEY",
        base_url=get_config_value("DASHSCOPE_BASE_URL", args.llm_base_url, env, dotenv, "https://dashscope.aliyuncs.com/compatible-mode/v1"),
        max_pages=_safe_positive_int(get_config_value("LLM_REVIEW_MAX_PAGES", str(args.llm_review_max_pages), env, dotenv, "8"), 8),
        timeout=_safe_positive_int(get_config_value("LLM_REVIEW_TIMEOUT", str(args.llm_review_timeout), env, dotenv, "600"), 600),
        mode=args.llm_review_mode,
        spec_docx_path=spec_path,
        all_pages=bool(args.llm_review_all_pages),
        batch_size=max(1, int(args.llm_review_batch_size)),
        separate_html=not bool(args.llm_review_no_separate_html),
        image_upload=bool(args.llm_review_image_upload),
        image_transport=getattr(args, "llm_review_image_transport", "data_uri"),
        spec_file_id=getattr(args, "llm_review_spec_file_id", None),
        doc_target=getattr(args, "llm_review_doc_target", "fixed"),
        upload_target_docx=bool(getattr(args, "llm_review_upload_target_docx", True)),
        delete_uploaded_target_after_review=bool(getattr(args, "llm_review_delete_uploaded_target_after_review", False)),
        output_path=Path(args.llm_review_output) if args.llm_review_output else None,
    )


def call_openai_compatible_chat(
    *,
    base_url: str,
    api_key: str,
    model: str,
    messages: list[dict[str, str]],
    timeout: int = 60,
) -> dict:
    endpoint = f"{base_url.rstrip('/')}/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    body = {"model": model, "messages": messages, "temperature": 0, "response_format": {"type": "json_object"}}
    req = urllib.request.Request(
        endpoint,
        data=json.dumps(body, ensure_ascii=False).encode("utf-8"),
        headers=headers,
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError:
        fallback = {"model": model, "messages": messages, "temperature": 0}
        fallback_req = urllib.request.Request(
            endpoint,
            data=json.dumps(fallback, ensure_ascii=False).encode("utf-8"),
            headers=headers,
            method="POST",
        )
        with urllib.request.urlopen(fallback_req, timeout=timeout) as resp:
            return json.loads(resp.read().decode("utf-8"))


def parse_llm_review_response(text: str) -> dict:
    content = text.strip()
    if content.startswith("```"):
        content = re.sub(r"^```(?:json)?\s*", "", content, flags=re.IGNORECASE)
        content = re.sub(r"\s*```$", "", content)
    try:
        payload = json.loads(content)
    except json.JSONDecodeError:
        return {"enabled": True, "status": "failed", "reason": "invalid_json"}
    if not isinstance(payload, dict):
        return {"enabled": True, "status": "failed", "reason": "invalid_json"}
    return payload

RULES: dict[str, Rule] = {
    "chapter": Rule(
        "chapter",
        "chapter title",
        "宋体",
        "Times New Roman",
        16,
        True,
        ALIGN_CENTER,
        space_before_lines=0.5,
        space_after_lines=0.5,
        expected="宋体三号16pt，加粗，居中，固定20磅行距，段前0.5行，段后0.5行，章序号与章名间空一格。",
    ),
    "heading2": Rule(
        "heading2",
        "second-level heading",
        "宋体",
        "Times New Roman",
        14,
        True,
        ALIGN_LEFT,
        space_before_lines=0.5,
        space_after_lines=0.5,
        expected="宋体四号14pt，加粗，左对齐，固定20磅行距，段前0.5行，段后0.5行，序号与题名间空一格。",
    ),
    "heading3": Rule(
        "heading3",
        "third-level heading",
        "宋体",
        "Times New Roman",
        12,
        True,
        ALIGN_LEFT,
        space_before_lines=0.5,
        space_after_lines=0.5,
        expected="宋体小四12pt，加粗，左对齐，固定20磅行距，段前0.5行，段后0.5行，序号与题名间空一格。",
    ),
    "body": Rule(
        "body",
        "body paragraph",
        "宋体",
        "Times New Roman",
        12,
        None,
        ALIGN_JUSTIFIED,
        first_line_indent_chars=2,
        space_before_pt=0,
        space_after_pt=0,
        expected="中文宋体小四12pt，英文和数字Times New Roman 12pt，两端对齐，首行缩进2个汉字符，段前段后0磅，固定20磅行距。",
    ),
    "figure_caption": Rule(
        "figure_caption",
        "figure caption",
        "宋体",
        "Times New Roman",
        10.5,
        None,
        ALIGN_CENTER,
        first_line_indent_chars=0,
        space_before_lines=0,
        space_after_lines=1,
        expected="图序、图名、图注：宋体五号10.5pt，居中，固定20磅行距，段前0行，段后1行，置于图下方，图序与图名之间空一格。",
    ),
    "table_caption": Rule(
        "table_caption",
        "table caption",
        "宋体",
        "Times New Roman",
        10.5,
        None,
        ALIGN_CENTER,
        first_line_indent_chars=0,
        space_before_lines=1,
        space_after_lines=0,
        expected="表序、表名、表注：按章编号，置于表上方，宋体五号10.5pt，居中，固定20磅行距，段前1行，段后0行，表序与表名之间空一个空格；续表头右顶格，续表应重复表编号和表头。",
    ),
    "abstract_title_zh": Rule(
        "abstract_title_zh",
        "Chinese abstract title",
        "宋体",
        "Times New Roman",
        16,
        True,
        ALIGN_CENTER,
        space_before_lines=0.5,
        space_after_lines=0.5,
        expected="摘要标题写作“摘 要”，宋体三号16pt，加粗，居中，段前0.5行，段后0.5行，固定20磅行距。",
    ),
    "abstract_title_en": Rule(
        "abstract_title_en",
        "English abstract title",
        "Times New Roman",
        "Times New Roman",
        14,
        True,
        ALIGN_CENTER,
        space_before_lines=0.5,
        space_after_lines=0.5,
        expected="英文摘要标题必须写作 ABSTRACT，所有字母大写，且应为独立标题段落；Times New Roman四号14pt，加粗，居中，段前0.5行，段后0.5行。",
    ),
    "keywords": Rule(
        "keywords",
        "keywords paragraph",
        "宋体",
        "Times New Roman",
        12,
        None,
        ALIGN_LEFT,
        expected="关键词段落宋体小四12pt，英文Times New Roman 12pt，关键词标签加粗，关键词之间用分号隔开。",
    ),
    "toc_title": Rule(
        "toc_title",
        "table of contents title",
        "宋体",
        "Times New Roman",
        16,
        True,
        ALIGN_CENTER,
        space_before_lines=0.5,
        space_after_lines=0.5,
        expected="目录标题写作“目 录”，宋体三号16pt，加粗，居中，段前0.5行，段后0.5行。",
    ),
}


def _rule_from_config(key: str, data: dict[str, object]) -> Rule:
    return Rule(
        key=key,
        label=str(data.get("label", key)),
        font_east_asia=str(data.get("font_east_asia", "宋体")),
        font_ascii=str(data.get("font_ascii", "Times New Roman")),
        size_pt=float(data.get("size_pt", 12)),
        bold=data.get("bold") if data.get("bold") is None else bool(data.get("bold")),
        alignment=ALIGNMENT_VALUES.get(data.get("alignment"), None),
        line_spacing_pt=data.get("line_spacing_pt") if data.get("line_spacing_pt") is None else float(data.get("line_spacing_pt")),
        first_line_indent_cm=data.get("first_line_indent_cm") if data.get("first_line_indent_cm") is None else float(data.get("first_line_indent_cm")),
        first_line_indent_chars=data.get("first_line_indent_chars") if data.get("first_line_indent_chars") is None else float(data.get("first_line_indent_chars")),
        hanging_indent_chars=data.get("hanging_indent_chars") if data.get("hanging_indent_chars") is None else float(data.get("hanging_indent_chars")),
        space_before_pt=data.get("space_before_pt") if data.get("space_before_pt") is None else float(data.get("space_before_pt")),
        space_after_pt=data.get("space_after_pt") if data.get("space_after_pt") is None else float(data.get("space_after_pt")),
        space_before_lines=data.get("space_before_lines") if data.get("space_before_lines") is None else float(data.get("space_before_lines")),
        space_after_lines=data.get("space_after_lines") if data.get("space_after_lines") is None else float(data.get("space_after_lines")),
        expected=str(data.get("expected", "")),
        category=str(data.get("category", "format")),
    )


def load_rules() -> dict[str, Rule]:
    if not RULE_CONFIG_PATH.exists():
        return RULES
    data = json.loads(RULE_CONFIG_PATH.read_text(encoding="utf-8"))
    return {key: _rule_from_config(key, value) for key, value in data["rules"].items()}


RULES = load_rules()
if "keywords" in RULES:
    RULES.setdefault("keywords_zh", replace(RULES["keywords"], key="keywords_zh", label="Chinese keywords paragraph"))
    RULES.setdefault("keywords_en", replace(RULES["keywords"], key="keywords_en", label="English keywords paragraph"))
if "cover_title_zh" in RULES:
    RULES.setdefault("thesis_title_zh", replace(RULES["cover_title_zh"], key="thesis_title_zh", label="Chinese thesis title"))
if "cover_title_en" in RULES:
    RULES.setdefault("thesis_title_en", replace(RULES["cover_title_en"], key="thesis_title_en", label="English thesis title"))


def is_docx(path: Path) -> bool:
    return path.suffix.lower() == ".docx"


def paragraph_text(paragraph) -> str:
    return re.sub(r"\s+", " ", paragraph.text or "").strip()


SECTION_ORDER = {
    "cover": 0,
    "declaration": 1,
    "abstract_zh": 2,
    "abstract_en": 2,
    "toc": 3,
    "symbols": 4,
    "body": 5,
    "references": 6,
    "appendix": 7,
    "acknowledgement": 8,
}

SECTION_SEQUENCE = ["cover", "declaration", "abstract", "toc", "symbols", "body", "references", "appendix", "acknowledgement"]

SECTION_LABELS = {
    "cover": "封面",
    "declaration": "封二",
    "abstract": "摘要和关键词",
    "toc": "目录",
    "symbols": "符号说明",
    "body": "正文",
    "references": "参考文献",
    "appendix": "附录",
    "acknowledgement": "致谢",
}


def section_module_name(marker: str) -> str:
    return "abstract" if marker in {"abstract_zh", "abstract_en"} else marker


def compact_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def detect_section_marker(text: str) -> str | None:
    compact = compact_text(text)
    if not compact:
        return None
    if "原创性声明" in compact or "版权使用授权书" in compact:
        return "declaration"
    if compact in {"摘要", "摘要:"} or text == "摘 要" or re.match(r"^摘要\s*[:：]", text):
        return "abstract_zh"
    if compact.upper() == "ABSTRACT" or re.match(r"^Abstract\s*[:：]", text, re.I):
        return "abstract_en"
    if compact == "目录" or text == "目 录":
        return "toc"
    if compact == "符号说明":
        return "symbols"
    if (
        re.match(r"^第[一二三四五六七八九十百\d]+章\S*", compact)
        or compact in {"引言", "绪论"}
        or re.match(r"^\d+(引言|绪论)$", compact)
    ):
        return "body"
    if compact == "参考文献":
        return "references"
    if compact in {"附录", "附錄"}:
        return "appendix"
    if compact in {"致谢", "致謝"}:
        return "acknowledgement"
    return None


def analyze_section_sequence(texts: list[str], has_toc_field: bool = False) -> dict[str, object]:
    regions: list[str] = []
    warnings: list[str] = []
    present: set[str] = set()
    current = "cover"
    current_order = SECTION_ORDER[current]
    saw_nonempty = False
    entered_modules: set[str] = {"cover"}

    def module_name(region: str) -> str:
        return "abstract" if region in {"abstract_zh", "abstract_en"} else region

    for index, text in enumerate(texts):
        text = text.strip()
        if not text:
            regions.append(current)
            continue
        saw_nonempty = True
        marker = detect_section_marker(text)
        if marker is not None:
            if current == "toc" and marker == "body":
                # TOC lines can contain chapter patterns like "第1章 ... 1"; do not
                # advance to body until a real body heading appears outside TOC-entry shape.
                compact = compact_text(text)
                looks_like_toc_entry = (
                    ("." in text and re.search(r"\d+\s*$", text) is not None)
                    or ("……" in text or "..." in text)
                    or bool(re.match(r"^第[一二三四五六七八九十百\d]+章", compact) and re.search(r"\d+\s*$", compact))
                )
                if looks_like_toc_entry:
                    marker = None
            if marker is None:
                regions.append(current)
                present.add(current)
                continue
            marker_order = SECTION_ORDER[marker]
            if marker_order < current_order:
                warnings.append(
                    f"组成部分顺序疑似错位：段落 {index} 出现“{text[:40]}”，应属于{SECTION_LABELS.get(marker, marker)}，"
                    f"但当前已经进入{SECTION_LABELS.get(current, current)}。"
                )
            else:
                if marker != current:
                    marker_module = module_name(marker)
                    current_module = module_name(current)
                    # abstract_zh -> abstract_en is an in-module transition, not duplication
                    if marker_module == current_module:
                        pass
                    elif marker_module in entered_modules:
                        warnings.append(
                            f"组成部分疑似重复出现：段落 {index} 再次进入“{SECTION_LABELS.get(marker_module, marker_module)}”模块。"
                        )
                    else:
                        entered_modules.add(marker_module)
                current = marker
                current_order = marker_order
        regions.append(current)
        present.add(current)

    if saw_nonempty:
        present.add("cover")
    if "abstract_zh" in present or "abstract_en" in present:
        present.add("abstract")
    if has_toc_field:
        present.add("toc")

    required = ["cover", "declaration", "abstract", "toc", "body", "references", "acknowledgement"]
    optional = ["symbols", "appendix"]
    for key in required:
        if key not in present:
            warnings.append(f"缺失必需组成部分：{SECTION_LABELS[key]}。")
    for key in optional:
        if key not in present:
            warnings.append(f"未检测到非必需组成部分：{SECTION_LABELS[key]}；如论文需要该部分，请补充并按规范排版。")

    return {
        "regions": regions,
        "present": sorted(present),
        "warnings": warnings,
    }


def compute_section_boundary_findings(texts: list[str], regions: list[str]) -> tuple[list[str], list[str]]:
    errors: list[str] = []
    warnings: list[str] = []
    entered_modules: set[str] = {"cover"}
    ended_modules: set[str] = set()
    current_module = "cover"
    for idx, text in enumerate(texts):
        stripped = text.strip()
        if not stripped:
            continue
        marker = detect_section_marker(stripped)
        if marker is not None:
            marker_module = section_module_name(marker)
            region_module = section_module_name(regions[idx]) if idx < len(regions) else current_module
            marker_order = SECTION_SEQUENCE.index(marker_module) if marker_module in SECTION_SEQUENCE else -1
            current_order = SECTION_SEQUENCE.index(current_module) if current_module in SECTION_SEQUENCE else -1
            if marker_order != -1 and current_order != -1 and marker_order < current_order:
                errors.append(
                    f"结构错误：段落 {idx} 出现“{stripped[:40]}”，疑似{SECTION_LABELS.get(marker_module, marker_module)}错位；当前已进入{SECTION_LABELS.get(current_module, current_module)}。"
                )
            if marker_module != current_module:
                ended_modules.add(current_module)
                if marker_module in ended_modules:
                    errors.append(f"结构错误：段落 {idx} 再次进入“{SECTION_LABELS.get(marker_module, marker_module)}”模块。")
                if marker_module in entered_modules and marker_module != "abstract":
                    errors.append(f"结构错误：段落 {idx} 检测到模块重复出现：{SECTION_LABELS.get(marker_module, marker_module)}。")
                entered_modules.add(marker_module)
                current_module = region_module if region_module in SECTION_SEQUENCE else marker_module
        compact = compact_text(stripped)
        if (
            "\u4f5c\u8005\u7b80\u5386" in compact
            or "\u653b\u8bfb\u5b66\u4f4d\u671f\u95f4\u53d1\u8868" in compact
            or "\u5b66\u672f\u8bba\u6587\u4e0e\u7814\u7a76\u6210\u679c" in compact
        ) and current_module != "appendix":
            warnings.append("\u68c0\u6d4b\u5230\u4e5d\u5927\u6a21\u5757\u4e4b\u5916\u7684\u7591\u4f3c\u989d\u5916\u7ec4\u6210\u90e8\u5206\uff0c\u9700\u8981\u4eba\u5de5\u786e\u8ba4\u3002")
    return errors, warnings


def detect_abstract_thesis_titles(texts: list[str]) -> dict[int, str]:
    title_rules: dict[int, str] = {}
    previous_nonempty: int | None = None
    for index, text in enumerate(texts):
        stripped = text.strip()
        if not stripped:
            continue
        marker = detect_section_marker(stripped)
        if marker == "abstract_zh" and previous_nonempty is not None:
            title_rules.setdefault(previous_nonempty, "thesis_title_zh")
        elif marker == "abstract_en" and previous_nonempty is not None:
            title_rules.setdefault(previous_nonempty, "thesis_title_en")
        previous_nonempty = index
    return title_rules


def style_heading_key(paragraph, region: str) -> str | None:
    if region != "body":
        return None
    style_name = (paragraph.style.name if paragraph.style else "").strip().lower()
    style_id = (paragraph.style.style_id if paragraph.style else "").strip().lower()
    style_token = f"{style_name} {style_id}"
    if "heading 1" in style_token or "heading1" in style_token:
        return "chapter"
    if "heading 2" in style_token or "heading2" in style_token:
        return "heading2"
    if "heading 3" in style_token or "heading3" in style_token:
        return "heading3"
    return None


def classify_checked_paragraph(paragraph, idx: int, text: str, region: str, title_overrides: dict[int, str]) -> str | None:
    style_name = (paragraph.style.name if paragraph.style else "").lower()
    if style_name.startswith("toc") and section_module_name(detect_section_marker(text) or "") != "toc":
        return "toc_level3" if style_name.endswith("3") else "toc_level2" if style_name.endswith("2") else "toc_level1"
    key = title_overrides.get(idx) or style_heading_key(paragraph, region) or classify_paragraph(text, in_body=region == "body", region=region)
    if key is None:
        return None
    region_scopes: dict[str, set[str]] = {
        "cover": {"cover_title_zh", "cover_title_en", "cover_field_zh", "cover_field_en", "cover_date", "thesis_title_zh", "thesis_title_en"},
        "declaration": set(),
        "abstract_zh": {"abstract_title_zh", "abstract_body_zh", "keywords", "keywords_zh", "thesis_title_zh"},
        "abstract_en": {"abstract_title_en", "abstract_body_en", "keywords", "keywords_en", "thesis_title_en"},
        "toc": {"toc_title", "toc_total_pages", "toc_level1", "toc_level2", "toc_level3"},
        "symbols": {"symbols_title", "symbols_body"},
        "body": {"chapter", "heading2", "heading3", "body", "figure_caption", "table_caption", "formula"},
        "references": {"reference_title", "reference_entry"},
        "appendix": {"appendix_title", "appendix_body"},
        "acknowledgement": {"acknowledgement_title", "acknowledgement_body"},
    }
    allowed = region_scopes.get(region, set())
    return key if key in allowed else None
def classify_paragraph(text: str, in_body: bool, region: str = "front") -> str | None:
    compact = compact_text(text)
    if re.match(r"^论文总页数[:：]\s*\d+\s*页", text):
        return "toc_total_pages"
    if compact in {"摘要", "摘要:"} or text == "摘 要":
        return "abstract_title_zh"
    if compact.upper() == "ABSTRACT":
        return "abstract_title_en"
    if region == "abstract_zh" and re.match(r"^(关键词|关键字)\s*[:：]", text):
        return "keywords_zh"
    if region == "abstract_en" and re.match(r"^(Key\s*words?|Keywords?)\s*[:：]", text, re.I):
        return "keywords_en"
    if compact == "目录" or text == "目 录":
        return "toc_title"
    if re.match(r"^第[一二三四五六七八九十百\d]+章\s*\S+", text):
        return "chapter"
    if region == "body" and re.match(r"^\d+(引言|绪论)$", compact):
        return "chapter"
    if compact == "参考文献":
        return "reference_title"
    if compact == "符号说明":
        return "symbols_title"
    if compact in {"附录", "附錄"}:
        return "appendix_title"
    if compact in {"致谢", "致謝"}:
        return "acknowledgement_title"
    if region == "toc":
        if re.match(r"^第[一二三四五六七八九十百\d]+章\s+\S+", text):
            return "toc_level1"
        if re.match(r"^\d+\.\d+\.\d+\s+\S+", text):
            return "toc_level3"
        if re.match(r"^\d+\.\d+\s+\S+", text):
            return "toc_level2"
    if region == "body" and re.match(r"^\d+\.\d+\.\d+\s*\S+", text):
        return "heading3"
    if region == "body" and re.match(r"^\d+\.\d+\s*\S+", text):
        return "heading2"
    if re.match(r"^图\s*\d+[\.\-－]\d+\s*\S+", text):
        return "figure_caption"
    if re.match(r"^表\s*\d+[\.\-－]\d+\s*\S+", text):
        return "table_caption"
    if re.match(r"^（?\(?\d+[-－]\d+）?\)?$", text):
        return "formula"
    if region == "references" and _reference_entry_start_match(text):
        return "reference_entry"
    if region == "references" and compact and compact != "参考文献":
        return "reference_entry"
    if region == "symbols" and len(text) >= 3:
        return "symbols_body"
    if region == "appendix" and len(text) >= 10:
        return "appendix_body"
    if region == "acknowledgement" and len(text) >= 10:
        return "acknowledgement_body"
    if region == "abstract_zh" and len(text) >= 20:
        return "abstract_body_zh"
    if region == "abstract_en" and len(text) >= 20:
        return "abstract_body_en"
    if region == "cover":
        if "论文提交日期" in compact:
            return "cover_date"
        if re.search(r"论文题目|题目", text):
            return "cover_title_zh"
        if re.search(r"姓名|学号|学院|专业|研究方向|指导教师|辅助指导教师", text):
            return "cover_field_zh"
    if region == "body" and len(text) >= 1:
        return "body"
    return None


def next_region(text: str, current: str) -> str:
    compact = text.replace(" ", "")
    if compact == "摘要":
        return "abstract_zh"
    if compact.upper() == "ABSTRACT":
        return "abstract_en"
    if compact == "目录":
        return "toc"
    if re.match(r"^第[一二三四五六七八九十百\d]+章\s*\S+", text):
        return "body"
    if compact == "参考文献":
        return "references"
    if compact == "符号说明":
        return "symbols"
    if compact in {"附录", "附錄"}:
        return "appendix"
    if compact in {"致谢", "致謝"}:
        return "acknowledgement"
    return current


def font_name(run, attr: str) -> str | None:
    rpr = run._element.rPr
    if rpr is None or rpr.rFonts is None:
        return run.font.name
    return rpr.rFonts.get(qn(f"w:{attr}")) or run.font.name


def _style_font_name(style, attr: str) -> str | None:
    while style is not None:
        rpr = style.element.rPr
        if rpr is not None and rpr.rFonts is not None:
            value = rpr.rFonts.get(qn(f"w:{attr}"))
            if value:
                return value
        if style.font and style.font.name:
            return style.font.name
        style = style.base_style
    return None


def _style_font_size(style) -> float | None:
    while style is not None:
        if style.font and style.font.size:
            return style.font.size.pt
        style = style.base_style
    return None


def _style_bold(style) -> bool | None:
    while style is not None:
        if style.font and style.font.bold is not None:
            return style.font.bold
        style = style.base_style
    return None


def common_run_format(paragraph) -> dict[str, object]:
    runs = [run for run in paragraph.runs if run.text.strip()]
    run = runs[0] if runs else None
    style = paragraph.style
    return {
        "eastAsia": font_name(run, "eastAsia") if run else _style_font_name(style, "eastAsia"),
        "ascii": font_name(run, "ascii") if run else _style_font_name(style, "ascii"),
        "size": run.font.size.pt if run and run.font.size else _style_font_size(style),
        "bold": run.bold if run and run.bold is not None else _style_bold(style),
    }


def run_effective_format(run, paragraph) -> dict[str, object]:
    style = paragraph.style
    return {
        "eastAsia": font_name(run, "eastAsia") or _style_font_name(style, "eastAsia"),
        "ascii": font_name(run, "ascii") or _style_font_name(style, "ascii"),
        "size": run.font.size.pt if run.font.size else _style_font_size(style),
        "bold": run.bold if run.bold is not None else _style_bold(style),
    }


def contains_cjk(text: str) -> bool:
    return bool(re.search(r"[\u3400-\u9fff]", text))


def contains_latin_or_digit(text: str) -> bool:
    return bool(re.search(r"[A-Za-z0-9]", text))


CJK_SPACING_BOUNDARY = r"\u3400-\u9fff\u3000-\u303f\uff00-\uffef"
BODY_CJK_SPACING_EXPECTED = "正文中只允许非中文之间保留空格；汉字两端不允许有空格；中文与英文、数字、中文标点之间的手动空格应删除。"


def normalize_body_cjk_spacing(text: str) -> str:
    normalized = text
    normalized = re.sub(rf"([{CJK_SPACING_BOUNDARY}])\s+", r"\1", normalized)
    normalized = re.sub(rf"\s+([{CJK_SPACING_BOUNDARY}])", r"\1", normalized)
    return normalized


def find_body_cjk_spacing_issues(text: str) -> list[str]:
    normalized = normalize_body_cjk_spacing(text)
    if normalized == text:
        return []
    return [f"{text} -> {normalized}"]


def _apply_text_to_runs(paragraph, new_text: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        if paragraph.text != new_text:
            paragraph.text = new_text
        return
    old_text = "".join(run.text for run in runs)
    if old_text == new_text:
        return
    cursor = 0
    for idx, run in enumerate(runs):
        if idx == len(runs) - 1:
            run.text = new_text[cursor:]
            break
        take = min(len(run.text), max(len(new_text) - cursor, 0))
        run.text = new_text[cursor : cursor + take]
        cursor += take


def _replace_paragraph_runs(paragraph, parts: list[tuple[str, bool | None]]) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    for text, bold in parts:
        if text == "":
            continue
        run = paragraph.add_run(text)
        if bold is not None:
            run.bold = bold


def run_format_summary(run, paragraph) -> str:
    fmt = run_effective_format(run, paragraph)
    return (
        f"文字片段={run.text[:40]!r}；中文字体={fmt['eastAsia'] or '继承/未直接设置'}；"
        f"西文字体={fmt['ascii'] or '继承/未直接设置'}；字号={fmt['size'] or '继承/未直接设置'}pt；"
        f"加粗={bold_label(fmt['bold'])}"
    )


def normalize_font_name(name: str | None) -> str | None:
    if name is None:
        return None
    text = str(name).strip()
    if not text:
        return None
    lowered = re.sub(r"\s+", " ", text).strip().lower()
    lowered = re.sub(r"\s+regular$", "", lowered)
    if lowered in {"times new roman", "timesnewroman"}:
        return "Times New Roman"
    if lowered in {"simsun", "songti", "宋体"}:
        return "宋体"
    return text


def run_matches_rule(run, paragraph, rule: Rule) -> bool:
    text = run.text.strip()
    if not text:
        return True
    fmt = run_effective_format(run, paragraph)
    if not approx(fmt["size"], rule.size_pt):
        return False
    if rule.bold is not None and fmt["bold"] is not rule.bold:
        return False
    actual_east = normalize_font_name(fmt["eastAsia"])
    expected_east = normalize_font_name(rule.font_east_asia)
    if contains_cjk(text) and actual_east and actual_east != expected_east:
        return False
    actual_ascii = normalize_font_name(fmt["ascii"])
    expected_ascii = normalize_font_name(rule.font_ascii)
    if contains_latin_or_digit(text) and actual_ascii and actual_ascii != expected_ascii:
        return False
    return True


def collect_run_format_issues(paragraph, paragraph_index: int, rule: Rule, text: str) -> list[Issue]:
    issues: list[Issue] = []
    bad_runs = [
        run_format_summary(run, paragraph)
        for run in paragraph.runs
        if run.text.strip() and not run_matches_rule(run, paragraph, rule)
    ]
    if not bad_runs:
        return issues
    current = " | ".join(bad_runs[:6])
    if len(bad_runs) > 6:
        current += f" | ... +{len(bad_runs) - 6} more"
    issues.append(
        Issue(
            paragraph_index=paragraph_index,
            rule_key=f"{rule.key}_runs",
            text_type=f"{rule.label} run formatting",
            text_excerpt=text[:160],
            current=current,
            expected=rule.expected,
            message=f"文本类型：{rule.label} run formatting\n当前格式：{current}\n应改为：{rule.expected}",
            category="run-format",
            location=f"paragraph {paragraph_index}",
        )
    )
    return issues


def pt_value(value) -> float | None:
    return round(value.pt, 2) if value is not None else None


def value_label(value: object | None, unit: str = "") -> str:
    if value is None:
        return "继承/未直接设置"
    return f"{value}{unit}"


def bold_label(value: object | None) -> str:
    if value is True:
        return "是"
    if value is False:
        return "否"
    return "继承/未直接设置"


def line_spacing_pt(paragraph) -> float | None:
    pf = paragraph.paragraph_format
    style_pf = paragraph.style.paragraph_format if paragraph.style is not None else None
    line_spacing_value = pf.line_spacing if pf.line_spacing is not None else (style_pf.line_spacing if style_pf else None)
    return line_spacing_value.pt if hasattr(line_spacing_value, "pt") else None


def first_line_indent_cm(paragraph) -> float | None:
    pf = paragraph.paragraph_format
    style_pf = paragraph.style.paragraph_format if paragraph.style is not None else None
    first_indent_value = pf.first_line_indent if pf.first_line_indent is not None else (style_pf.first_line_indent if style_pf else None)
    return first_indent_value.cm if first_indent_value is not None else None


def first_line_indent_chars(paragraph) -> float | None:
    ppr = paragraph._p.pPr
    if ppr is not None:
        ind = ppr.find(qn("w:ind"))
        if ind is not None:
            value = ind.get(qn("w:firstLineChars"))
            if value is not None:
                return float(value) / 100
    if paragraph.style is not None:
        style_ppr = paragraph.style._element.pPr
        if style_ppr is not None:
            ind = style_ppr.find(qn("w:ind"))
            if ind is not None:
                value = ind.get(qn("w:firstLineChars"))
                if value is not None:
                    return float(value) / 100
    return None


def hanging_indent_chars(paragraph) -> float | None:
    ppr = paragraph._p.pPr
    if ppr is not None:
        ind = ppr.find(qn("w:ind"))
        if ind is not None:
            value = ind.get(qn("w:hangingChars"))
            if value is not None:
                return float(value) / 100
    if paragraph.style is not None:
        style_ppr = paragraph.style._element.pPr
        if style_ppr is not None:
            ind = style_ppr.find(qn("w:ind"))
            if ind is not None:
                value = ind.get(qn("w:hangingChars"))
                if value is not None:
                    return float(value) / 100
    return None


def has_twips_hanging_indent(paragraph) -> bool:
    ppr = paragraph._p.pPr
    if ppr is not None:
        ind = ppr.find(qn("w:ind"))
        if ind is not None and ind.get(qn("w:hanging")) is not None:
            return True
    if paragraph.style is not None:
        style_ppr = paragraph.style._element.pPr
        if style_ppr is not None:
            ind = style_ppr.find(qn("w:ind"))
            if ind is not None and ind.get(qn("w:hanging")) is not None:
                return True
    return False


def _set_first_line_indent_chars(paragraph, chars: float) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    for attr in ("w:firstLine", "w:firstLineChars"):
        qattr = qn(attr)
        if qattr in ind.attrib:
            del ind.attrib[qattr]
    ind.set(qn("w:firstLineChars"), str(int(round(chars * 100))))


def _set_hanging_indent_chars(paragraph, chars: float) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    for attr in ("w:hanging", "w:hangingChars"):
        qattr = qn(attr)
        if qattr in ind.attrib:
            del ind.attrib[qattr]
    ind.set(qn("w:hangingChars"), str(int(round(chars * 100))))


def spacing_before_after_pt(paragraph) -> tuple[float | None, float | None]:
    pf = paragraph.paragraph_format
    style_pf = paragraph.style.paragraph_format if paragraph.style is not None else None
    before_value = pf.space_before if pf.space_before is not None else (style_pf.space_before if style_pf else None)
    after_value = pf.space_after if pf.space_after is not None else (style_pf.space_after if style_pf else None)
    return pt_value(before_value), pt_value(after_value)


def _spacing_element_from_paragraph(paragraph):
    ppr = paragraph._p.pPr
    if ppr is not None:
        spacing = ppr.find(qn("w:spacing"))
        if spacing is not None:
            return spacing
    if paragraph.style is not None:
        style_ppr = paragraph.style._element.pPr
        if style_ppr is not None:
            return style_ppr.find(qn("w:spacing"))
    return None


def spacing_before_after_lines(paragraph) -> tuple[float | None, float | None]:
    spacing = _spacing_element_from_paragraph(paragraph)
    if spacing is None:
        return None, None
    before = spacing.get(qn("w:beforeLines"))
    after = spacing.get(qn("w:afterLines"))
    return (
        float(before) / 100 if before is not None else None,
        float(after) / 100 if after is not None else None,
    )


def _set_line_spacing_units(paragraph, before: float | None, after: float | None) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    for attr in ("w:before", "w:after"):
        qattr = qn(attr)
        if qattr in spacing.attrib:
            del spacing.attrib[qattr]
    if before is not None:
        spacing.set(qn("w:beforeLines"), str(int(round(before * 100))))
    if after is not None:
        spacing.set(qn("w:afterLines"), str(int(round(after * 100))))


def paragraph_format_differences(paragraph, rule: Rule | None) -> list[str]:
    if rule is None:
        return []
    fmt = common_run_format(paragraph)
    style_pf = paragraph.style.paragraph_format if paragraph.style is not None else None
    differences: list[str] = []
    if not approx(fmt["size"], rule.size_pt):
        differences.append(f"字号：当前 {value_label(fmt['size'], 'pt')}，应为 {rule.size_pt}pt")
    if rule.bold is not None and fmt["bold"] is not rule.bold:
        differences.append(f"加粗：当前 {bold_label(fmt['bold'])}，应为 {bold_label(rule.bold)}")
    alignment = paragraph.alignment if paragraph.alignment is not None else (style_pf.alignment if style_pf else None)
    if rule.alignment is not None and alignment != rule.alignment:
        if not (rule.key in {"abstract_body_zh", "abstract_body_en", "acknowledgement_body"} and alignment in {WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.JUSTIFY}):
            differences.append(
                f"对齐方式：当前 {ALIGN_LABELS_ZH.get(alignment, '继承/未直接设置' if alignment is None else str(alignment))}，"
                f"应为 {ALIGN_LABELS_ZH.get(rule.alignment, str(rule.alignment))}"
            )
    actual_line_spacing = line_spacing_pt(paragraph)
    if rule.line_spacing_pt is not None and not approx(actual_line_spacing, rule.line_spacing_pt):
        differences.append(f"固定行距：当前 {value_label(round(actual_line_spacing, 2) if actual_line_spacing is not None else None, '磅')}，应为 {rule.line_spacing_pt}磅")
    actual_indent = first_line_indent_cm(paragraph)
    if rule.first_line_indent_cm is not None and not approx(actual_indent, rule.first_line_indent_cm, tolerance=0.15):
        differences.append(f"首行缩进：当前 {value_label(round(actual_indent, 2) if actual_indent is not None else None, 'cm')}，应为 {rule.first_line_indent_cm}cm")
    actual_indent_chars = first_line_indent_chars(paragraph)
    if rule.first_line_indent_chars is not None and not approx(actual_indent_chars, rule.first_line_indent_chars, tolerance=0.02):
        differences.append(f"首行缩进：当前 {value_label(actual_indent_chars, '字符')}，应为 {rule.first_line_indent_chars}字符")
    actual_hanging_chars = hanging_indent_chars(paragraph)
    if rule.hanging_indent_chars is not None and not approx(actual_hanging_chars, rule.hanging_indent_chars, tolerance=0.02):
        if has_twips_hanging_indent(paragraph):
            differences.append(f"悬挂缩进：当前检测到长度单位悬挂缩进（w:hanging），应为 Word 字符单位 {rule.hanging_indent_chars} 字符（w:hangingChars）。")
        else:
            differences.append(f"悬挂缩进：当前 {value_label(actual_hanging_chars, '字符')}，应为 {rule.hanging_indent_chars}字符（w:hangingChars）。")
    before, after = spacing_before_after_pt(paragraph)
    if rule.space_before_pt is not None and not approx(before, rule.space_before_pt):
        differences.append(f"段前：当前 {value_label(before, 'pt')}，应为 {rule.space_before_pt}pt")
    if rule.space_after_pt is not None and not approx(after, rule.space_after_pt):
        differences.append(f"段后：当前 {value_label(after, 'pt')}，应为 {rule.space_after_pt}pt")
    before_lines, after_lines = spacing_before_after_lines(paragraph)
    if rule.space_before_lines is not None and not approx(before_lines, rule.space_before_lines, tolerance=0.02):
        differences.append(f"段前：当前 {value_label(before_lines, '行')}，应为 {rule.space_before_lines}行")
    if rule.space_after_lines is not None and not approx(after_lines, rule.space_after_lines, tolerance=0.02):
        differences.append(f"段后：当前 {value_label(after_lines, '行')}，应为 {rule.space_after_lines}行")
    return differences


def current_format(paragraph, rule: Rule | None = None) -> str:
    fmt = common_run_format(paragraph)
    pf = paragraph.paragraph_format
    style_pf = paragraph.style.paragraph_format if paragraph.style is not None else None
    alignment = paragraph.alignment if paragraph.alignment is not None else (style_pf.alignment if style_pf else None)
    align = ALIGN_LABELS_ZH.get(alignment, "继承/未直接设置" if alignment is None else str(alignment))
    line_spacing_value = pf.line_spacing if pf.line_spacing is not None else (style_pf.line_spacing if style_pf else None)
    line_spacing = line_spacing_value.pt if hasattr(line_spacing_value, "pt") else None
    first_indent = first_line_indent_cm(paragraph)
    first_indent_chars = first_line_indent_chars(paragraph)
    hanging_chars = hanging_indent_chars(paragraph)
    before, after = spacing_before_after_pt(paragraph)
    before_lines, after_lines = spacing_before_after_lines(paragraph)
    parts = [
        f"段落样式：{paragraph.style.name if paragraph.style else '未知'}",
        f"中文字体：{fmt['eastAsia'] or '继承/未直接设置'}",
        f"西文字体：{fmt['ascii'] or '继承/未直接设置'}",
        f"字号：{value_label(fmt['size'], 'pt')}",
        f"加粗：{bold_label(fmt['bold'])}",
        f"对齐方式：{align}",
        f"固定行距：{value_label(round(line_spacing, 2) if line_spacing is not None else None, '磅')}",
        f"首行缩进：{value_label(first_indent_chars, '字符') if first_indent_chars is not None else value_label(round(first_indent, 2) if first_indent is not None else None, 'cm')}",
        f"悬挂缩进：{value_label(hanging_chars, '字符')}",
        f"段前：{value_label(before_lines, '行') if before_lines is not None else value_label(before, 'pt')}",
        f"段后：{value_label(after_lines, '行') if after_lines is not None else value_label(after, 'pt')}",
    ]
    differences = paragraph_format_differences(paragraph, rule)
    if differences:
        parts.insert(0, "疑似不符合项：" + "；".join(differences))
    return "; ".join(parts)


def approx(actual: float | None, expected: float | None, tolerance: float = 0.4) -> bool:
    if expected is None:
        return True
    if actual is None:
        return False
    return abs(actual - expected) <= tolerance


def paragraph_matches(paragraph, rule: Rule) -> bool:
    fmt = common_run_format(paragraph)
    pf = paragraph.paragraph_format
    style_pf = paragraph.style.paragraph_format if paragraph.style is not None else None
    size = fmt["size"]
    if not approx(size, rule.size_pt):
        return False
    if rule.bold is not None and fmt["bold"] is not rule.bold:
        return False
    alignment = paragraph.alignment if paragraph.alignment is not None else (style_pf.alignment if style_pf else None)
    if rule.alignment is not None and alignment != rule.alignment:
        if not (rule.key in {"abstract_body_zh", "abstract_body_en", "acknowledgement_body"} and alignment in {WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.JUSTIFY}):
            return False
    if rule.line_spacing_pt is not None:
        line_spacing_value = pf.line_spacing if pf.line_spacing is not None else (style_pf.line_spacing if style_pf else None)
        line_spacing = line_spacing_value.pt if hasattr(line_spacing_value, "pt") else None
        if not approx(line_spacing, rule.line_spacing_pt):
            return False
    if rule.first_line_indent_cm is not None:
        indent_value = pf.first_line_indent if pf.first_line_indent is not None else (style_pf.first_line_indent if style_pf else None)
        actual_indent = indent_value.cm if indent_value is not None else None
        if not approx(actual_indent, rule.first_line_indent_cm, tolerance=0.15):
            return False
    if rule.first_line_indent_chars is not None:
        actual_indent_chars = first_line_indent_chars(paragraph)
        if not approx(actual_indent_chars, rule.first_line_indent_chars, tolerance=0.02):
            return False
    if rule.hanging_indent_chars is not None:
        actual_hanging_chars = hanging_indent_chars(paragraph)
        if not approx(actual_hanging_chars, rule.hanging_indent_chars, tolerance=0.02):
            return False
        if has_twips_hanging_indent(paragraph):
            return False
    before, after = spacing_before_after_pt(paragraph)
    if rule.space_before_pt is not None and not approx(before, rule.space_before_pt):
        return False
    if rule.space_after_pt is not None and not approx(after, rule.space_after_pt):
        return False
    before_lines, after_lines = spacing_before_after_lines(paragraph)
    if rule.space_before_lines is not None and not approx(before_lines, rule.space_before_lines, tolerance=0.02):
        return False
    if rule.space_after_lines is not None and not approx(after_lines, rule.space_after_lines, tolerance=0.02):
        return False
    return True


def set_run_font(run, rule: Rule) -> None:
    run.font.name = rule.font_ascii
    run._element.rPr.rFonts.set(qn("w:eastAsia"), rule.font_east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), rule.font_ascii)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), rule.font_ascii)
    run.font.size = Pt(rule.size_pt)
    if rule.bold is not None:
        run.bold = rule.bold


def apply_rule(paragraph, rule: Rule) -> None:
    if rule.alignment is not None:
        paragraph.alignment = rule.alignment
    pf = paragraph.paragraph_format
    if rule.line_spacing_pt is not None:
        pf.line_spacing = Pt(rule.line_spacing_pt)
    if rule.first_line_indent_cm is not None:
        pf.first_line_indent = Cm(rule.first_line_indent_cm)
    if rule.first_line_indent_chars is not None:
        _set_first_line_indent_chars(paragraph, rule.first_line_indent_chars)
    if rule.hanging_indent_chars is not None:
        _set_hanging_indent_chars(paragraph, rule.hanging_indent_chars)
    if rule.space_before_pt is not None:
        pf.space_before = Pt(rule.space_before_pt)
    if rule.space_after_pt is not None:
        pf.space_after = Pt(rule.space_after_pt)
    if rule.space_before_lines is not None or rule.space_after_lines is not None:
        _set_line_spacing_units(paragraph, rule.space_before_lines, rule.space_after_lines)
    for run in paragraph.runs:
        set_run_font(run, rule)
    if not paragraph.runs and paragraph.text:
        set_run_font(paragraph.add_run(), rule)


def section_paragraph_ranges(document: Document) -> list[tuple[int, int]]:
    breaks: list[int] = []
    for idx, paragraph in enumerate(document.paragraphs):
        ppr = paragraph._p.pPr
        if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
            breaks.append(idx)
    ranges: list[tuple[int, int]] = []
    start = 0
    for end in breaks:
        ranges.append((start, end))
        start = end + 1
    if start <= len(document.paragraphs) - 1:
        ranges.append((start, len(document.paragraphs) - 1))
    while len(ranges) < len(document.sections):
        ranges.append((len(document.paragraphs), len(document.paragraphs) - 1))
    return ranges[: len(document.sections)]


def section_text(document: Document, start: int, end: int) -> str:
    if start > end:
        return ""
    return "\n".join(paragraph_text(document.paragraphs[idx]) for idx in range(start, end + 1))


def document_has_toc_field(document: Document) -> bool:
    xml = document._element.xml
    return "TOC" in xml and "instrText" in xml


def first_body_paragraph_index(document: Document) -> int | None:
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    for idx, region in enumerate(regions):
        if region == "body":
            return idx
    return None


def ensure_body_starts_new_page_number_section(document: Document) -> bool:
    body_idx = first_body_paragraph_index(document)
    if body_idx is None or body_idx <= 0:
        return False
    ranges = section_paragraph_ranges(document)
    for section_idx, (start, end) in enumerate(ranges):
        if start <= body_idx <= end and start < body_idx:
            previous = document.paragraphs[body_idx - 1]
            ppr = previous._p.get_or_add_pPr()
            if ppr.find(qn("w:sectPr")) is not None:
                return False
            front_sect_pr = deepcopy(document.sections[section_idx]._sectPr)
            for child in list(front_sect_pr):
                if child.tag == qn("w:footerReference"):
                    front_sect_pr.remove(child)
            ppr.append(front_sect_pr)
            return True
    return False


def find_header_start_section(document: Document) -> tuple[int | None, str]:
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    ranges = section_paragraph_ranges(document)
    target_regions = {
        "declaration",
        "abstract_zh",
        "abstract_en",
        "toc",
        "body",
        "references",
        "appendix",
        "acknowledgement",
    }
    for idx, (start, end) in enumerate(ranges):
        if start > end:
            continue
        section_regions = set(regions[start : end + 1])
        if section_regions & target_regions:
            return idx, "declaration-or-later"
    if len(document.sections) >= 2:
        return 1, "fallback-second-section"
    return None, "not-found"


def clear_header(header) -> None:
    for paragraph in header.paragraphs:
        paragraph.clear()


def _write_cuit_header(header) -> None:
    clear_header(header)
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.clear()
    run = para.add_run("成都信息工程大学学士学位论文")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.name = "Times New Roman"
    run.bold = False
    run.italic = False
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)


def apply_header(section, enabled: bool) -> None:
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False
    targets = [section.header, section.first_page_header, section.even_page_header]
    if not enabled:
        for header in targets:
            clear_header(header)
        return
    for header in targets:
        _write_cuit_header(header)


def apply_cuit_page_headers(document: Document, allow_layout_fixes: bool) -> str:
    header_start, header_reason = find_header_start_section(document)
    if header_start is None:
        return header_reason
    if header_reason == "fallback-second-section" and not allow_layout_fixes:
        return "manual-review-needed"
    for section in document.sections:
        section.header.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        section.even_page_header.is_linked_to_previous = False
    for idx, section in enumerate(document.sections):
        apply_header(section, enabled=idx >= header_start)
    return header_reason


def _set_section_page_number_format(section, fmt: str, start: int | None = None, clear_start: bool = False) -> None:
    if OxmlElement is None:
        return
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        cols = sect_pr.find(qn("w:cols"))
        if cols is not None:
            sect_pr.insert(list(sect_pr).index(cols), pg_num_type)
        else:
            sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))
    elif clear_start and qn("w:start") in pg_num_type.attrib:
        del pg_num_type.attrib[qn("w:start")]


def _roman_lower(number: int) -> str:
    values = [
        (1000, "m"),
        (900, "cm"),
        (500, "d"),
        (400, "cd"),
        (100, "c"),
        (90, "xc"),
        (50, "l"),
        (40, "xl"),
        (10, "x"),
        (9, "ix"),
        (5, "v"),
        (4, "iv"),
        (1, "i"),
    ]
    remaining = max(1, number)
    result: list[str] = []
    for value, token in values:
        while remaining >= value:
            result.append(token)
            remaining -= value
    return "".join(result)


def _add_simple_field(paragraph, instr: str, display: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)
    run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    run._r.append(instr_text)
    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)
    paragraph.add_run(display)
    run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def _reset_footer_to_page_field(section, field_kind: str, display_number: int = 1) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    for child in list(footer._element):
        footer._element.remove(child)
    paragraph = footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if field_kind == "front":
        display = _roman_lower(display_number)
        instr = " PAGE  \\* roman  \\* MERGEFORMAT "
        _add_simple_field(paragraph, instr, display)
    else:
        paragraph.add_run("第")
        _add_simple_field(paragraph, " PAGE  \\* Arabic  \\* MERGEFORMAT ", str(display_number))
        paragraph.add_run("页 共")
        _add_simple_field(paragraph, " NUMPAGES  \\* Arabic  \\* MERGEFORMAT ", str(display_number))
        paragraph.add_run("页")
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(9)


def _rewrite_footer_page_field(section, field_kind: str, start: int = 1) -> None:
    _reset_footer_to_page_field(section, field_kind, display_number=start)
    return
    field_kind = field_kind.lower()
    visible = _roman_lower(start) if field_kind == "front" else str(start)
    instr = " PAGE  \\* roman  \\* MERGEFORMAT " if field_kind == "front" else " PAGE  \\* Arabic  \\* MERGEFORMAT "
    in_page_field = False
    seen_separator = False
    for node in section.footer._element.iter():
        if node.tag == qn("w:fldChar"):
            fld_type = node.get(qn("w:fldCharType"))
            if fld_type == "begin":
                in_page_field = True
                seen_separator = False
            elif fld_type == "separate" and in_page_field:
                seen_separator = True
            elif fld_type == "end":
                in_page_field = False
                seen_separator = False
        elif node.tag == qn("w:instrText") and in_page_field and node.text and "PAGE" in node.text.upper():
            node.text = instr
        elif node.tag == qn("w:t") and in_page_field and seen_separator:
            if node.text and re.fullmatch(r"\s*[\dIVXLCDMivxlcdm]+\s*", node.text):
                node.text = visible


def apply_page_number_formats(document: Document) -> None:
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    ranges = section_paragraph_ranges(document)
    front_started = False
    main_started = False
    decisions: list[tuple[int, str, int | None]] = []
    abstract_or_toc_sections: list[int] = []
    body_or_later_sections: list[int] = []
    for section_idx, _section in enumerate(document.sections):
        start, end = ranges[section_idx] if section_idx < len(ranges) else (0, -1)
        section_regions = set(regions[start : end + 1]) if start <= end else set()
        sec_text = section_text(document, start, end)
        sec_compact = compact_text(sec_text)
        looks_like_toc_section = (
            ("目录" in sec_compact)
            or ("目次" in sec_compact)
            or bool(re.search(r"\.{6,}", sec_text))
            or ("论文总页数" in sec_text)
        )
        if {"declaration", "abstract_zh", "abstract_en", "toc"} & section_regions:
            abstract_or_toc_sections.append(section_idx)
        elif looks_like_toc_section:
            abstract_or_toc_sections.append(section_idx)
        if {"body", "references", "appendix", "acknowledgement"} & section_regions:
            body_or_later_sections.append(section_idx)

    first_front = min(abstract_or_toc_sections) if abstract_or_toc_sections else None
    last_known_front = max(abstract_or_toc_sections) if abstract_or_toc_sections else None
    main_candidates = [idx for idx in body_or_later_sections if last_known_front is None or idx > last_known_front]
    first_main = min(main_candidates) if main_candidates else None
    last_front = None
    if first_front is not None and first_main is not None and first_main >= first_front:
        last_front = first_main - 1
    elif first_front is not None:
        last_front = max(abstract_or_toc_sections)

    for section_idx, section in enumerate(document.sections):
        is_front = first_front is not None and last_front is not None and first_front <= section_idx <= last_front
        is_main = first_main is not None and section_idx >= first_main
        if is_front:
            _set_section_page_number_format(section, "lowerRoman", start=1 if not front_started else None, clear_start=front_started)
            decisions.append((section_idx, "front", None))
            front_started = True
        elif is_main:
            _set_section_page_number_format(section, "decimal", start=1 if not main_started else None, clear_start=main_started)
            decisions.append((section_idx, "main", 1 if not main_started else None))
            main_started = True
        else:
            footer = section.footer
            footer.is_linked_to_previous = False
            for child in list(footer._element):
                footer._element.remove(child)
    for section_idx, kind, display in decisions:
        if kind == "main":
            _rewrite_footer_page_field(document.sections[section_idx], kind, start=display or 1)
    for section_idx, kind, display in decisions:
        if kind == "front":
            _rewrite_footer_page_field(document.sections[section_idx], kind, start=display or 1)


def _xml_para_text(para_xml) -> str:
    return re.sub(r"\s+", " ", "".join(para_xml.itertext()) if para_xml is not None else "").strip()


def ensure_abstract_heading_split(document: Document, regions: list[str] | None = None) -> None:
    if regions is None:
        texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
        regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    # Iterate over a snapshot because we may insert heading paragraphs before current ones.
    for idx, paragraph in enumerate(list(document.paragraphs)):
        if idx >= len(regions) or regions[idx] != "abstract_zh":
            continue
        text = paragraph_text(paragraph)
        if not text:
            continue
        zh_match = re.match(r"^\s*摘要\s*[:：]\s*(.+)$", text)
        if zh_match:
            prev_text_compact = compact_text(_xml_para_text(paragraph._p.getprevious()))
            if prev_text_compact != "摘要":
                heading = paragraph.insert_paragraph_before("摘 要")
                apply_rule(heading, RULES["abstract_title_zh"])
            paragraph.text = zh_match.group(1).strip()
            apply_rule(paragraph, RULES["abstract_body_zh"])
            continue


def normalize_english_abstract_title(document: Document, regions: list[str] | None = None) -> None:
    if regions is None:
        texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
        regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    for idx, paragraph in enumerate(document.paragraphs):
        if idx >= len(regions) or regions[idx] != "abstract_en":
            continue
        stripped = paragraph_text(paragraph).strip()
        if not stripped:
            continue
        split_match = re.match(r"^abstract\s*[:：]\s*(.*)$", stripped, re.I)
        if split_match:
            tail = split_match.group(1).strip()
            if tail:
                heading = paragraph.insert_paragraph_before("ABSTRACT")
                apply_rule(heading, RULES["abstract_title_en"])
                _replace_paragraph_runs(paragraph, [(tail, False)])
                apply_rule(paragraph, RULES["abstract_body_en"])
            else:
                _replace_paragraph_runs(paragraph, [("ABSTRACT", True)])
                apply_rule(paragraph, RULES["abstract_title_en"])
            continue
        if re.match(r"^abstract\s*$", stripped, re.I):
            _replace_paragraph_runs(paragraph, [("ABSTRACT", True)])


def normalize_keywords_label_runs(document: Document, regions: list[str] | None = None) -> None:
    if regions is None:
        texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
        regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    for idx, paragraph in enumerate(document.paragraphs):
        if idx >= len(regions) or regions[idx] not in {"abstract_zh", "abstract_en"}:
            continue
        stripped = paragraph_text(paragraph).strip()
        if not stripped:
            continue
        zh = re.match(r"^\s*(关键词)\s*[:：]\s*(.*)$", stripped)
        en = re.match(r"^\s*(Key\s*words)\s*[:：]\s*(.*)$", stripped, re.I)
        if zh:
            tail = zh.group(2).strip()
            label = "关键词："
            body = f"{tail}" if tail else ""
            _replace_paragraph_runs(paragraph, [(label, True), (body, False)])
            continue
        if en:
            tail = en.group(2).strip()
            label = "Key words: "
            body = f"{tail}" if tail else ""
            _replace_paragraph_runs(paragraph, [(label, True), (body, False)])


def _is_plain_blank_paragraph(paragraph) -> bool:
    if paragraph_text(paragraph):
        return False
    p = paragraph._p
    ppr = p.pPr
    if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
        return False
    for node in p.iter():
        if node.tag == qn("w:br"):
            return False
        if node.tag == qn("w:drawing"):
            return False
        if node.tag == qn("w:object"):
            return False
        if node.tag == qn("w:pict"):
            return False
    return True


def ensure_english_abstract_starts_new_page(document: Document) -> bool:
    changed = False
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    title_overrides = detect_abstract_thesis_titles(texts)
    for idx, paragraph in enumerate(document.paragraphs):
        text = texts[idx]
        if not text:
            continue
        key = title_overrides.get(idx) or classify_paragraph(text, in_body=False, region="abstract_en")
        if key in {"thesis_title_en", "abstract_title_en"}:
            prev = document.paragraphs[idx - 1] if idx > 0 else None
            prev_has_sect_break = False
            if prev is not None and prev._p.pPr is not None and prev._p.pPr.find(qn("w:sectPr")) is not None:
                prev_has_sect_break = True
                # Keep section break for pagination, but normalize this placeholder paragraph.
                prev_pfmt = prev.paragraph_format
                if prev_pfmt.first_line_indent is not None:
                    prev_pfmt.first_line_indent = None
                    changed = True
                if prev_pfmt.line_spacing is not None:
                    prev_pfmt.line_spacing = None
                    changed = True
                if prev_pfmt.space_before is not None:
                    prev_pfmt.space_before = None
                    changed = True
                if prev_pfmt.space_after is not None:
                    prev_pfmt.space_after = None
                    changed = True
            if prev_has_sect_break:
                if paragraph.paragraph_format.page_break_before:
                    paragraph.paragraph_format.page_break_before = False
                    changed = True
            elif paragraph.paragraph_format.page_break_before is not True:
                paragraph.paragraph_format.page_break_before = True
                changed = True
            break
    return changed


def remove_redundant_blank_paragraphs_by_region(
    document: Document,
    regions: list[str],
    allowed_regions: set[str],
) -> bool:
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    removed = False
    for idx in range(len(document.paragraphs) - 2, 0, -1):
        if idx >= len(regions) or regions[idx] not in allowed_regions:
            continue
        if texts[idx]:
            continue
        paragraph = document.paragraphs[idx]
        if not _is_plain_blank_paragraph(paragraph):
            continue
        prev_region = regions[idx - 1] if idx - 1 < len(regions) else None
        next_region = regions[idx + 1] if idx + 1 < len(regions) else None
        if prev_region in allowed_regions and next_region in allowed_regions:
            paragraph._element.getparent().remove(paragraph._element)
            removed = True
    return removed


def remove_sectpr_only_blank_placeholders(
    document: Document,
    regions: list[str],
    allowed_regions: set[str],
) -> bool:
    removed = False
    for idx in range(len(document.paragraphs) - 1, 0, -1):
        if idx >= len(regions) or regions[idx] not in allowed_regions:
            continue
        paragraph = document.paragraphs[idx]
        if paragraph_text(paragraph):
            continue
        ppr = paragraph._p.pPr
        if ppr is None:
            continue
        sect_pr = ppr.find(qn("w:sectPr"))
        if sect_pr is None:
            continue
        # Keep true contentful blank lines untouched; only strip placeholder lines.
        has_content_signal = any(
            node.tag in {qn("w:br"), qn("w:drawing"), qn("w:object"), qn("w:pict")}
            for node in paragraph._p.iter()
        )
        if has_content_signal:
            continue
        # Move section properties to previous paragraph so sectioning remains valid.
        prev = document.paragraphs[idx - 1]
        prev_ppr = prev._p.get_or_add_pPr()
        old_prev_sect = prev_ppr.find(qn("w:sectPr"))
        if old_prev_sect is not None:
            prev_ppr.remove(old_prev_sect)
        prev_ppr.append(deepcopy(sect_pr))
        paragraph._element.getparent().remove(paragraph._element)
        removed = True
    return removed


def _is_target_empty_paragraph_text(text: str) -> bool:
    return re.sub(r"[\s\u3000\xa0]+", "", text or "") == ""


def _paragraph_has_image_payload(paragraph) -> bool:
    suffixes = (
        "}drawing",
        "}pict",
        "}object",
        "}OLEObject",
        "}inline",
        "}anchor",
        "}blip",
        "}shape",
        "}imagedata",
    )
    for node in paragraph._p.iter():
        tag = str(node.tag)
        if tag.endswith(suffixes):
            return True
    return False


def _blank_paragraph_has_risky_nodes(paragraph) -> bool:
    if _paragraph_has_image_payload(paragraph):
        return True
    ppr = paragraph._p.pPr
    if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
        return True
    risky_tags = {
        qn("w:fldChar"),
        qn("w:instrText"),
        qn("w:drawing"),
        qn("w:object"),
        qn("w:pict"),
    }
    for node in paragraph._p.iter():
        if node.tag in risky_tags:
            return True
        if node.tag == qn("w:br"):
            br_type = node.get(qn("w:type"))
            if br_type in {"page", "column"}:
                return True
    return False


def _is_effectively_blank_paragraph(paragraph) -> bool:
    if _paragraph_has_toc_field(paragraph):
        return False
    if not _is_target_empty_paragraph_text(paragraph.text):
        return False
    if _paragraph_has_image_payload(paragraph):
        return False
    for node in paragraph._p.iter():
        if node.tag in {qn("w:drawing"), qn("w:object"), qn("w:pict")}:
            return False
        if node.tag == qn("w:t") and _is_target_empty_paragraph_text(node.text or "") is False:
            return False
    return True


def _blank_paragraph_has_only_safe_boundary_nodes(paragraph) -> bool:
    allowed = {
        qn("w:p"),
        qn("w:pPr"),
        qn("w:r"),
        qn("w:rPr"),
        qn("w:br"),
        qn("w:bookmarkStart"),
        qn("w:bookmarkEnd"),
        qn("w:commentRangeStart"),
        qn("w:commentRangeEnd"),
        qn("w:sectPr"),
        qn("w:pStyle"),
        qn("w:spacing"),
        qn("w:ind"),
        qn("w:jc"),
        qn("w:rFonts"),
        qn("w:sz"),
        qn("w:szCs"),
        qn("w:highlight"),
        qn("w:proofErr"),
        qn("w:t"),
        qn("w:tab"),
    }
    for node in paragraph._p.iter():
        parent = node.getparent() if hasattr(node, "getparent") else None
        while parent is not None:
            if parent.tag == qn("w:sectPr"):
                parent = None
                break
            parent = parent.getparent() if hasattr(parent, "getparent") else None
        if parent is None and node.tag != qn("w:sectPr"):
            # Nodes under sectPr are considered safe boundary metadata.
            sect_ancestor = False
            cur = node.getparent() if hasattr(node, "getparent") else None
            while cur is not None:
                if cur.tag == qn("w:sectPr"):
                    sect_ancestor = True
                    break
                cur = cur.getparent() if hasattr(cur, "getparent") else None
            if sect_ancestor:
                continue
        if node.tag not in allowed:
            return False
    return True


def _migrate_sectpr_to_previous_paragraph(document: Document, blank_idx: int, prev_idx: int) -> bool:
    if blank_idx <= 0 or prev_idx < 0 or blank_idx >= len(document.paragraphs) or prev_idx >= len(document.paragraphs):
        return False
    blank = document.paragraphs[blank_idx]
    ppr = blank._p.pPr
    if ppr is None:
        return False
    sect_pr = ppr.find(qn("w:sectPr"))
    if sect_pr is None:
        return False
    prev = document.paragraphs[prev_idx]
    prev_ppr = prev._p.get_or_add_pPr()
    old_prev_sect = prev_ppr.find(qn("w:sectPr"))
    if old_prev_sect is not None:
        prev_ppr.remove(old_prev_sect)
    prev_ppr.append(deepcopy(sect_pr))
    return True


def _blank_paragraph_is_toc_field_placeholder(paragraph) -> bool:
    if not _is_target_empty_paragraph_text(paragraph.text):
        return False
    if _paragraph_has_image_payload(paragraph):
        return False
    has_toc_instr = False
    has_field_char = False
    for node in paragraph._p.iter():
        if node.tag in {qn("w:drawing"), qn("w:object"), qn("w:pict")}:
            return False
        if node.tag == qn("w:instrText"):
            text = (node.text or "").upper()
            if "TOC" in text:
                has_toc_instr = True
            else:
                return False
        elif node.tag == qn("w:fldChar"):
            has_field_char = True
    return has_toc_instr or has_field_char


def _paragraph_has_toc_field(paragraph) -> bool:
    for node in paragraph._p.iter():
        if node.tag == qn("w:instrText") and "TOC" in (node.text or "").upper():
            return True
    return False


def _is_body_heading_text(text: str) -> bool:
    compact = (text or "").strip()
    if not compact:
        return False
    return bool(
        re.match(r"^\s*第[一二三四五六七八九十百0-9]+章\s+\S+", compact)
        or re.match(r"^\s*\d+\s+\S+", compact)
        or re.match(r"^\s*\d+(?:\.\d+)+\s+\S+", compact)
    )


def cleanup_module_boundary_blank_paragraphs(
    document: Document,
    regions: list[str] | None = None,
    apply_changes: bool = True,
) -> tuple[list[Issue], int]:
    issues: list[Issue] = []
    removed = 0
    if regions is None:
        texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
        regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    n = len(document.paragraphs)
    if n == 0:
        return issues, removed

    def _next_non_empty(start_idx: int) -> int | None:
        for i in range(start_idx, len(document.paragraphs)):
            if not _is_target_empty_paragraph_text(paragraph_text(document.paragraphs[i])):
                return i
        return None

    def _next_toc_target(start_idx: int) -> int | None:
        for i in range(start_idx, len(document.paragraphs)):
            text = paragraph_text(document.paragraphs[i]).strip()
            region = regions[i] if i < len(regions) else ""
            if region == "body" or _is_body_heading_text(text):
                return None
            if region == "toc" or bool(re.match(r"^\s*目\s*录\s*$", text)) or _paragraph_has_toc_field(document.paragraphs[i]):
                return i
        return None

    boundary_specs: list[tuple[re.Pattern[str], callable]] = [
        (
            re.compile(r"^\s*Key\s*words\s*[:：]", re.I),
            lambda idx, region, text: (
                region == "toc"
                or bool(re.match(r"^\s*目\s*录\s*$", text))
                or _paragraph_has_toc_field(document.paragraphs[idx])
            ),
        ),
        (
            re.compile(r"^\s*关键词\s*[:：]"),
            lambda idx, region, text: region == "abstract_en" or bool(re.match(r"^\s*ABSTRACT\s*$", text, re.I)),
        ),
    ]

    i = 0
    while i < len(document.paragraphs):
        anchor_text = paragraph_text(document.paragraphs[i])
        anchor_region = regions[i] if i < len(regions) else ""
        matched = None
        for pat, target_match in boundary_specs:
            if pat.search(anchor_text):
                matched = (pat, target_match)
                break
        if matched is None or anchor_region == "cover":
            i += 1
            continue
        _pat, target_match = matched
        next_non_empty = _next_non_empty(i + 1)
        if next_non_empty is None:
            i += 1
            continue
        if _pat.pattern.lower().find("key\\s*words") >= 0:
            toc_target = _next_toc_target(i + 1)
            if toc_target is not None and _is_effectively_blank_paragraph(document.paragraphs[toc_target]):
                after_toc_target = _next_non_empty(toc_target + 1)
                if after_toc_target is not None:
                    aft_region = regions[after_toc_target] if after_toc_target < len(regions) else ""
                    aft_text = paragraph_text(document.paragraphs[after_toc_target]).strip()
                    if aft_region == "body" or _is_body_heading_text(aft_text):
                        toc_target = None
            if toc_target is None:
                if apply_changes:
                    for blank_idx in range(next_non_empty - 1, i, -1):
                        if blank_idx >= len(document.paragraphs):
                            continue
                        blank = document.paragraphs[blank_idx]
                        if not _is_effectively_blank_paragraph(blank):
                            continue
                        if _paragraph_has_image_payload(blank):
                            continue
                        is_toc_placeholder = _blank_paragraph_is_toc_field_placeholder(blank)
                        has_sectpr = blank._p.pPr is not None and blank._p.pPr.find(qn("w:sectPr")) is not None
                        has_page_break = any(
                            node.tag == qn("w:br") and node.get(qn("w:type")) == "page"
                            for node in blank._p.iter()
                        )
                        if has_sectpr and not _migrate_sectpr_to_previous_paragraph(document, blank_idx, i):
                            continue
                        if is_toc_placeholder or has_sectpr or has_page_break or _blank_paragraph_has_only_safe_boundary_nodes(blank):
                            blank._element.getparent().remove(blank._element)
                            removed += 1
                issues.append(
                    Issue(
                        paragraph_index=i,
                        rule_key="abstract_boundary_cleanup_crossed_toc",
                        text_type="abstract boundary cleanup",
                        text_excerpt=anchor_text[:160],
                        current="英文关键词后未能可靠定位目录起点，已停止清理以避免破坏正文分页。",
                        expected="英文关键词边界清理只允许定位到目录区域或目录域段落。",
                        message="英文关键词后未能可靠定位目录起点，已停止清理以避免破坏正文分页。",
                        category="page",
                        location="abstract_en to toc boundary",
                    )
                )
                i = next_non_empty
                continue
            next_non_empty = toc_target
        next_text = paragraph_text(document.paragraphs[next_non_empty]).strip()
        next_region = regions[next_non_empty] if next_non_empty < len(regions) else ""
        if not target_match(next_non_empty, next_region, next_text):
            i += 1
            continue
        failed = False
        for blank_idx in range(next_non_empty - 1, i, -1):
            if blank_idx >= len(document.paragraphs):
                continue
            blank = document.paragraphs[blank_idx]
            blank_region = regions[blank_idx] if blank_idx < len(regions) else ""
            if blank_region == "cover":
                continue
            if not _is_effectively_blank_paragraph(blank):
                continue
            is_toc_placeholder = _blank_paragraph_is_toc_field_placeholder(blank)
            if is_toc_placeholder and _pat.pattern.lower().find("key\\s*words") >= 0:
                # Keep TOC field paragraph itself, but allow cleaning ordinary boundary blanks around it.
                continue
            if not _blank_paragraph_has_only_safe_boundary_nodes(blank) and not is_toc_placeholder:
                failed = True
                continue
            has_sectpr = blank._p.pPr is not None and blank._p.pPr.find(qn("w:sectPr")) is not None
            has_page_break = any(
                node.tag == qn("w:br") and node.get(qn("w:type")) == "page"
                for node in blank._p.iter()
            )
            if has_page_break and apply_changes and next_non_empty < len(document.paragraphs) and next_region == "toc":
                document.paragraphs[next_non_empty].paragraph_format.page_break_before = True
            if has_sectpr and apply_changes:
                if not _migrate_sectpr_to_previous_paragraph(document, blank_idx, i):
                    failed = True
                    continue
            if apply_changes:
                blank._element.getparent().remove(blank._element)
                removed += 1
        if failed:
            issues.append(
                Issue(
                    paragraph_index=i,
                    rule_key="abstract_toc_boundary_blank_manual_review",
                    text_type="abstract/toc boundary blank paragraph",
                    text_excerpt=anchor_text[:160],
                    current="英文摘要与目录之间存在承载分节/分页属性的空白段，无法安全自动删除。",
                    expected="英文摘要与目录边界不应出现可见空白段；如含分节/分页属性需人工确认。",
                    message="英文摘要与目录之间存在承载分节/分页属性的空白段，无法安全自动删除，需要人工确认。",
                    category="body-spacing",
                    location="abstract_en to toc boundary",
                )
            )
        i = next_non_empty
    return issues, removed


def cleanup_body_heading_following_blank_paragraphs(
    document: Document,
    regions: list[str] | None = None,
    apply_changes: bool = True,
) -> tuple[list[Issue], int]:
    issues: list[Issue] = []
    removed = 0
    if regions is None:
        texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
        regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    i = 0
    while i < len(document.paragraphs) - 1:
        text = paragraph_text(document.paragraphs[i]).strip()
        region = regions[i] if i < len(regions) else ""
        if region != "body" or not _is_body_heading_text(text):
            i += 1
            continue
        j = i + 1
        while j < len(document.paragraphs) and _is_target_empty_paragraph_text(paragraph_text(document.paragraphs[j])):
            blank = document.paragraphs[j]
            if _paragraph_has_image_payload(blank):
                break
            has_sectpr = blank._p.pPr is not None and blank._p.pPr.find(qn("w:sectPr")) is not None
            if has_sectpr:
                issues.append(
                    Issue(
                        paragraph_index=i,
                        rule_key="body_heading_following_blank_manual_review",
                        text_type="body heading following blank",
                        text_excerpt=text[:160],
                        current="章标题与正文之间存在承载分节属性的空白段，无法安全自动删除。",
                        expected="章标题后不应存在导致标题与正文分离的空白段。",
                        message="章标题与正文之间存在承载分节属性的空白段，无法安全自动删除，需要人工确认。",
                        category="page",
                        location=f"paragraph {i}",
                    )
                )
                break
            if apply_changes:
                blank._element.getparent().remove(blank._element)
                removed += 1
            else:
                removed += 1
            # do not j += 1 after delete; list shifts
        if i + 1 < len(document.paragraphs):
            nxt = document.paragraphs[i + 1]
            if nxt.paragraph_format.page_break_before:
                nxt.paragraph_format.page_break_before = False
        i += 1
    return issues, removed


def cleanup_visible_blank_paragraphs_after_layout(document: Document) -> tuple[list[Issue], bool]:
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    issues1, removed1 = cleanup_module_boundary_blank_paragraphs(document, regions=regions, apply_changes=True)
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    issues2, removed2 = cleanup_body_heading_following_blank_paragraphs(document, regions=regions, apply_changes=True)
    return issues1 + issues2, (removed1 + removed2) > 0


def collect_empty_paragraph_issues(
    document: Document,
    regions: list[str] | None = None,
    allowed_regions: set[str] | None = None,
) -> tuple[list[Issue], list[int]]:
    issues: list[Issue] = []
    removable: list[int] = []
    if regions is None:
        texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
        regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    target_regions = allowed_regions or {"abstract_zh", "abstract_en", "body", "acknowledgement"}
    for idx, paragraph in enumerate(document.paragraphs):
        if idx >= len(regions) or regions[idx] not in target_regions:
            continue
        if not _is_target_empty_paragraph_text(paragraph.text):
            continue
        risky = _blank_paragraph_has_risky_nodes(paragraph)
        current = "检测到风险空段，需要人工确认，未自动删除。" if risky else "检测到安全空段，将在修复文档中删除。"
        issues.append(
            Issue(
                paragraph_index=idx,
                rule_key="empty_paragraph",
                text_type="empty paragraph in body/abstract/acknowledgement",
                text_excerpt="[blank paragraph]",
                current=current,
                expected="摘要正文、正文、致谢正文中不应保留空白段落，应删除。",
                message=f"文本类型：empty paragraph in body/abstract/acknowledgement\n当前问题：{current}\n应改为：摘要正文、正文、致谢正文中不应保留空白段落，应删除。",
                category="structure",
                location=f"paragraph {idx}",
            )
        )
        if not risky:
            removable.append(idx)
    return issues, removable


def remove_target_empty_paragraphs(document: Document, removable: list[int]) -> int:
    removed = 0
    for idx in sorted(set(removable), reverse=True):
        if idx < 0 or idx >= len(document.paragraphs):
            continue
        paragraph = document.paragraphs[idx]
        if _is_target_empty_paragraph_text(paragraph.text) and not _blank_paragraph_has_risky_nodes(paragraph):
            paragraph._element.getparent().remove(paragraph._element)
            removed += 1
    return removed


def normalize_document_structure(document: Document) -> None:
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    allowed_regions = {"abstract_zh", "abstract_en", "body", "acknowledgement"}
    ensure_abstract_heading_split(document, regions)
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    normalize_english_abstract_title(document, regions)
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    normalize_keywords_label_runs(document, regions)
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    remove_sectpr_only_blank_placeholders(document, regions, allowed_regions)
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    remove_redundant_blank_paragraphs_by_region(document, regions, {"abstract_zh", "abstract_en"})
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    _empty_issues, to_remove = collect_empty_paragraph_issues(document, regions=regions, allowed_regions=allowed_regions)
    remove_target_empty_paragraphs(document, to_remove)
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    cleanup_module_boundary_blank_paragraphs(document, regions=regions, apply_changes=True)
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    cleanup_body_heading_following_blank_paragraphs(document, regions=regions, apply_changes=True)


def apply_page_setup(document: Document, allow_layout_fixes: bool) -> None:
    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)
        if allow_layout_fixes:
            section.header_distance = Cm(1.5)
    apply_cuit_page_headers(document, allow_layout_fixes=True)
    if allow_layout_fixes:
        # TOC new-page handling is report-only to avoid page-break regressions on hidden TOC placeholders.
        pass


def _section_format(section) -> str:
    return (
        f"page={round(section.page_width.cm, 2)}cm x {round(section.page_height.cm, 2)}cm; "
        f"margins top={round(section.top_margin.cm, 2)}cm, bottom={round(section.bottom_margin.cm, 2)}cm, "
        f"left={round(section.left_margin.cm, 2)}cm, right={round(section.right_margin.cm, 2)}cm"
    )


def _section_matches(section) -> bool:
    return (
        approx(section.page_width.cm, 21, 0.1)
        and approx(section.page_height.cm, 29.7, 0.1)
        and approx(section.top_margin.cm, 2.5, 0.1)
        and approx(section.bottom_margin.cm, 2.5, 0.1)
        and approx(section.left_margin.cm, 3, 0.1)
        and approx(section.right_margin.cm, 3, 0.1)
    )


def _section_page_kinds(document: Document, regions: list[str]) -> list[str | None]:
    kinds: list[str | None] = []
    ranges = section_paragraph_ranges(document)
    for section_idx, _section in enumerate(document.sections):
        start, end = ranges[section_idx] if section_idx < len(ranges) else (0, -1)
        section_regions = set(regions[start : end + 1]) if start <= end else set()
        if {"declaration", "abstract_zh", "abstract_en", "toc"} & section_regions:
            kinds.append("front")
        elif {"body", "references", "appendix", "acknowledgement"} & section_regions:
            kinds.append("main")
        else:
            kinds.append(None)
    return kinds


def collect_section_issues(document: Document) -> list[Issue]:
    issues: list[Issue] = []
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    structure = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))
    regions = structure["regions"]
    boundary_errors, _boundary_warnings = compute_section_boundary_findings(texts, regions)
    for item in boundary_errors:
        issues.append(
            _issue_global(
                "section_sequence",
                "thesis section sequence",
                str(item),
                "九大模块必须按固定顺序且不重复：cover->declaration->abstract->toc->symbols->body->references->appendix->acknowledgement。",
                "structure",
                "document structure",
            )
        )
    section_page_kinds = _section_page_kinds(document, regions)
    header_start, header_reason = find_header_start_section(document)
    if header_reason in {"not-found", "fallback-second-section"}:
        rule = RULES["page_header"]
        issues.append(
            Issue(
                paragraph_index=-1,
                rule_key="page_header_missing",
                text_type=rule.label,
                text_excerpt="document header start",
                current="未能可靠定位从摘要/目录/正文开始的页眉边界，需要人工复核。",
                expected=rule.expected,
                message=f"文本类型：{rule.label}\n当前情况：未能可靠定位从摘要/目录/正文开始的页眉边界，需要人工复核。\n应改为：{rule.expected}",
                category=rule.category,
                location="document header start",
            )
        )
    for sec_idx, section in enumerate(document.sections, start=1):
        if not _section_matches(section):
            current = _section_format(section)
            issues.append(
                Issue(
                    paragraph_index=-1,
                    rule_key="page_setup",
                    text_type="page setup",
                    text_excerpt=f"section {sec_idx}",
                    current=current,
                    expected="A4 21cm x 29.7cm; margins top/bottom 2.5cm, left/right 3cm.",
                    message=f"文本类型：page setup\n当前格式：{current}\n应改为：A4，页边距上/下2.5cm，左/右3cm。",
                    category="page",
                    location=f"section {sec_idx}",
                )
            )
        zero_idx = sec_idx - 1
        header_text = " ".join(paragraph_text(p) for p in section.header.paragraphs).strip()
        rule = RULES["page_header"]
        if header_start is not None and zero_idx < header_start and header_text:
            issues.append(
                Issue(
                    paragraph_index=-1,
                    rule_key="page_header_missing",
                    text_type=rule.label,
                    text_excerpt=f"section {sec_idx} header",
                    current=header_text,
                    expected="封面 section 不应添加该页眉。",
                    message=f"文本类型：{rule.label}\n当前格式：封面 section 已存在页眉：{header_text}\n应改为：封面 section 不添加该页眉。",
                    category=rule.category,
                    location=f"section {sec_idx} header",
                )
            )
        if header_start is not None and zero_idx >= header_start and "成都信息工程大学学士学位论文" not in header_text:
            issues.append(
                Issue(
                    paragraph_index=-1,
                    rule_key="page_header_missing",
                    text_type=rule.label,
                    text_excerpt=f"section {sec_idx} header",
                    current=header_text or "[empty header]",
                    expected=rule.expected,
                    message=f"文本类型：{rule.label}\n当前格式：{header_text or '[empty header]'}\n应改为：{rule.expected}",
                    category=rule.category,
                    location=f"section {sec_idx} header",
                )
            )
        if header_start is not None and zero_idx >= header_start:
            para = section.header.paragraphs[0] if section.header.paragraphs else None
            header_ok = True
            if para is None or para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                header_ok = False
            if para is not None and para.runs:
                run = para.runs[0]
                size_ok = run.font.size is not None and approx(run.font.size.pt, 9.0, 0.1)
                east = None
                ascii_font = None
                if run._element.rPr is not None and run._element.rPr.rFonts is not None:
                    east = run._element.rPr.rFonts.get(qn("w:eastAsia"))
                    ascii_font = run._element.rPr.rFonts.get(qn("w:ascii")) or run._element.rPr.rFonts.get(qn("w:hAnsi"))
                if not size_ok or east != "宋体" or ascii_font != "Times New Roman":
                    header_ok = False
            if not header_ok:
                issues.append(
                    Issue(
                        paragraph_index=-1,
                        rule_key="page_header_format",
                        text_type=rule.label,
                        text_excerpt=f"section {sec_idx} header",
                        current=header_text or "[empty header]",
                        expected=rule.expected,
                        message=f"文本类型：{rule.label}\n当前格式：页眉字体/字号/对齐不符合要求。\n应改为：{rule.expected}",
                        category=rule.category,
                        location=f"section {sec_idx} header",
                    )
                )
        page_kind = section_page_kinds[zero_idx] if zero_idx < len(section_page_kinds) else None
        issues.extend(collect_page_number_issues(section, sec_idx, page_kind))
    return issues


def _section_page_number_format(section) -> tuple[str | None, str | None]:
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        return None, None
    return pg_num_type.get(qn("w:fmt")), pg_num_type.get(qn("w:start"))


def _footer_field_info(section) -> tuple[str, str]:
    instr: list[str] = []
    texts: list[str] = []
    for node in section.footer._element.iter():
        if node.tag == qn("w:instrText") and node.text:
            instr.append(node.text)
        elif node.tag == qn("w:t") and node.text:
            texts.append(node.text)
    return "".join(texts).strip(), " ".join(instr).strip()


def validate_front_page_number(fmt: str | None, visible_text: str, field_text: str) -> str | None:
    if "PAGE" not in field_text.upper() and not visible_text:
        return None
    if fmt != "lowerRoman":
        return f"页码编号格式为 {fmt or '未设置'}，应为 lowerRoman 小写罗马数字。"
    compact = re.sub(r"\s+", "", visible_text)
    if compact and re.search(r"\d|[IVXLCDM]", compact):
        return f"页码显示为 {visible_text}，应为小写罗马数字 i, ii, iii, iv, v, ...。"
    return None


def validate_main_page_number(fmt: str | None, visible_text: str, field_text: str) -> str | None:
    if "PAGE" not in field_text.upper() and not visible_text:
        return None
    compact = re.sub(r"\s+", "", visible_text)
    field_upper = field_text.upper()
    has_main_pattern = bool(re.search(r"第.+页共.+页", compact)) or ("PAGE" in field_upper and "NUMPAGES" in field_upper)
    if fmt not in {None, "decimal"}:
        return f"页码编号格式为 {fmt}，正文至致谢应使用阿拉伯数字。"
    if not has_main_pattern:
        return f"页码显示为 {visible_text or field_text}，应为“第*页 共*页”。"
    return None


def collect_page_number_issues(section, sec_idx: int, page_kind: str | None) -> list[Issue]:
    visible_text, field_text = _footer_field_info(section)
    if "PAGE" not in field_text.upper() and not visible_text:
        return []
    if page_kind is None:
        return []
    fmt, _start = _section_page_number_format(section)
    if page_kind == "main":
        problem = validate_main_page_number(fmt, visible_text, field_text)
        rule_key = "main_page_number"
    else:
        problem = validate_front_page_number(fmt, visible_text, field_text)
        rule_key = "front_page_number"
    if problem is None:
        return []
    rule = RULES[rule_key]
    return [
        Issue(
            paragraph_index=-1,
            rule_key=rule_key,
            text_type=rule.label,
            text_excerpt=f"section {sec_idx} footer",
            current=problem,
            expected=rule.expected,
            message=f"文本类型：{rule.label}\n当前格式：{problem}\n应改为：{rule.expected}",
            category=rule.category,
            location=f"section {sec_idx} footer",
        )
    ]


def collect_abstract_title_issues(document: Document) -> list[Issue]:
    issues: list[Issue] = []
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    has_zh_title = any(compact_text(text) == "摘要" or text.strip() == "摘 要" for text in texts)
    has_en_title = any(
        idx < len(regions) and regions[idx] == "abstract_en" and compact_text(text) == "ABSTRACT"
        for idx, text in enumerate(texts)
    )
    for idx, text in enumerate(texts):
        stripped = text.strip()
        if not stripped:
            continue
        if not has_zh_title and re.match(r"^摘要\s*[:：]", stripped):
            rule = RULES["abstract_title_zh"]
            issues.append(
                Issue(
                    paragraph_index=idx,
                    rule_key="abstract_title_zh_missing",
                    text_type=rule.label,
                    text_excerpt=stripped[:160],
                    current="摘要标题与摘要正文合并在同一段，且未检测到独立“摘 要”标题。",
                    expected=rule.expected,
                    message=f"文本类型：{rule.label}\n当前格式：摘要标题与正文合并\n应改为：{rule.expected}",
                    category=rule.category,
                    location=f"paragraph {idx}",
                )
            )
            break
    for idx, text in enumerate(texts):
        if idx >= len(regions) or regions[idx] != "abstract_en":
            continue
        stripped = text.strip()
        if not stripped:
            continue
        match = re.match(r"^(Abstract)\s*([:：])?\s*(.*)$", stripped, re.I)
        if not match:
            continue
        tail = match.group(3).strip()
        rule = RULES["abstract_title_en"]
        if tail:
            issues.append(
                Issue(
                    paragraph_index=idx,
                    rule_key="abstract_title_en_split",
                    text_type=rule.label,
                    text_excerpt=stripped[:160],
                    current=f"当前标题文本为“{stripped}”，标题与正文混在同一段。",
                    expected="英文摘要标题必须写作“ABSTRACT”，所有字母大写，并单独成段。",
                    message=(
                        f"文本类型：{rule.label}\n当前格式：{stripped}\n"
                        "应改为：ABSTRACT（标题单独成段，正文另起段）。"
                    ),
                    after="已拆分为：ABSTRACT + 摘要正文。",
                    category=rule.category,
                    location=f"paragraph {idx}",
                )
            )
            if not has_en_title:
                issues.append(
                    Issue(
                        paragraph_index=idx,
                        rule_key="abstract_title_en_missing",
                        text_type=rule.label,
                        text_excerpt=stripped[:160],
                        current="英文摘要标题与摘要正文合并在同一段，且未检测到独立“ABSTRACT”标题。",
                        expected=rule.expected,
                        message=f"文本类型：{rule.label}\n当前格式：英文摘要标题与正文合并\n应改为：{rule.expected}",
                        category=rule.category,
                        location=f"paragraph {idx}",
                    )
                )
                break
            continue
        if stripped != "ABSTRACT":
            issues.append(
                Issue(
                    paragraph_index=idx,
                    rule_key="abstract_title_en_text",
                    text_type=rule.label,
                    text_excerpt=stripped[:160],
                    current=f"当前标题文本为“{stripped}”。",
                    expected="英文摘要标题必须写作“ABSTRACT”，所有字母大写，并单独成段。",
                    message=f"文本类型：{rule.label}\n当前格式：{stripped}\n应改为：ABSTRACT",
                    category=rule.category,
                    location=f"paragraph {idx}",
                )
            )
    return issues


def _issue_global(rule_key: str, text_type: str, current: str, expected: str, category: str, location: str, excerpt: str = "") -> Issue:
    return Issue(
        paragraph_index=-1,
        rule_key=rule_key,
        text_type=text_type,
        text_excerpt=excerpt or location,
        current=current,
        expected=expected,
        message=f"文本类型：{text_type}\n当前位置：{location}\n当前情况：{current}\n应改为：{expected}",
        category=category,
        location=location,
    )


def _document_xml_text(document: Document) -> str:
    return ET.tostring(document.element.body, encoding="unicode")


def paragraph_has_right_tab(paragraph) -> bool:
    ppr = paragraph._p.pPr
    if ppr is None:
        return False
    tabs = ppr.find(qn("w:tabs"))
    if tabs is None:
        return False
    for tab in tabs.findall(qn("w:tab")):
        if tab.get(qn("w:val")) == "right":
            return True
    return False


def find_toc_start_paragraph_index(document: Document) -> int | None:
    for idx, paragraph in enumerate(document.paragraphs):
        if re.sub(r"\s+", "", paragraph_text(paragraph)) in {"目录", "目次"}:
            return idx
    for idx, paragraph in enumerate(document.paragraphs):
        for node in paragraph._p.iter():
            if node.tag == qn("w:instrText") and "TOC" in (node.text or "").upper():
                return idx
            if node.tag == qn("w:fldSimple") and "TOC" in (node.get(qn("w:instr")) or "").upper():
                return idx
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    for idx, region in enumerate(regions):
        if region == "toc":
            return idx
    for idx, text in enumerate(texts):
        stripped = text.strip()
        if re.search(r"(\.{2,}|…+)\s*\d+\s*$", stripped) and ("引言" in stripped or "章" in stripped):
            return idx
    return None


def _paragraph_has_page_break_node(paragraph) -> bool:
    for node in paragraph._p.iter():
        if node.tag == qn("w:br") and node.get(qn("w:type")) == "page":
            return True
        if node.tag == qn("w:lastRenderedPageBreak"):
            return True
    return False


def ensure_toc_starts_new_page(document: Document) -> bool:
    toc_idx = find_toc_start_paragraph_index(document)
    if toc_idx is None:
        return False
    changed = False
    toc_paragraph = document.paragraphs[toc_idx]
    if toc_paragraph.paragraph_format.page_break_before is not True:
        toc_paragraph.paragraph_format.page_break_before = True
        changed = True
    remove_indexes: list[int] = []
    scan = toc_idx - 1
    while scan >= 0:
        prev = document.paragraphs[scan]
        if not _is_effectively_blank_paragraph(prev):
            break
        ppr = prev._p.pPr
        if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
            break
        if _blank_paragraph_is_toc_field_placeholder(prev):
            break
        if _paragraph_has_page_break_node(prev) and _blank_paragraph_has_only_safe_boundary_nodes(prev):
            remove_indexes.append(scan)
            scan -= 1
            continue
        break
    for idx in remove_indexes:
        paragraph = document.paragraphs[idx]
        paragraph._element.getparent().remove(paragraph._element)
        changed = True
    return changed


def collect_toc_issues(document: Document) -> list[Issue]:
    issues: list[Issue] = []
    xml = _document_xml_text(document)
    has_toc_field = "TOC" in xml and "instrText" in xml
    has_toc_title = any(re.sub(r"\s+", "", paragraph_text(p)) in {"目录", "目次"} for p in document.paragraphs)
    if has_toc_title and not has_toc_field:
        issues.append(
            _issue_global(
                "toc_field",
                "TOC field",
                "目录标题存在，但未检测到 Word TOC 域代码。",
                "目录应由 Word/WPS 目录域或等效可更新结构生成，便于页码刷新。",
                "toc",
                "document TOC",
            )
        )
    in_toc = False
    for idx, paragraph in enumerate(document.paragraphs):
        text = paragraph_text(paragraph)
        compact = text.replace(" ", "")
        if compact == "目录":
            in_toc = True
            continue
        if in_toc and re.match(r"^第[一二三四五六七八九十百\d]+章\s*\S+", text):
            in_toc = False
        if not in_toc or not text:
            continue
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        looks_toc_entry = style_name.startswith("toc") or re.search(r"(\.{2,}|…+|\t)\s*\d+$", text) or re.match(r"^(\d+\.\d+|第[一二三四五六七八九十百\d]+章)", text)
        if looks_toc_entry and not paragraph_has_right_tab(paragraph):
            issues.append(
                Issue(
                    paragraph_index=idx,
                    rule_key="toc_tab_stop",
                    text_type="TOC page-number right tab",
                    text_excerpt=text[:160],
                    current="未检测到右对齐制表位。",
                    expected="目录页码应通过右对齐制表位或等效目录域右对齐。",
                    message=f"文本类型：TOC page-number right tab\n当前格式：未检测到右对齐制表位\n应改为：目录页码右对齐。",
                    category="toc",
                    location=f"paragraph {idx}",
                )
            )
    if has_toc_title or has_toc_field:
        toc_idx = find_toc_start_paragraph_index(document)
        issues.append(
            Issue(
                paragraph_index=toc_idx if toc_idx is not None else -1,
                rule_key="toc_start_new_page_manual_review",
                text_type="目录另起页人工复核",
                text_excerpt=(paragraph_text(document.paragraphs[toc_idx])[:160] if toc_idx is not None else "TOC start"),
                current="检测到目录/TOC 自动目录，工具不自动设置段前分页。",
                expected="目录应另起一页。",
                message="目录为 Word/WPS 自动目录或 TOC 域时，自动设置 page_break_before 可能作用到隐藏 TOC 占位段并生成空白页。本工具仅提示，请人工在 Word/WPS 中确认目录标题“目 录”另起一页。",
                category="toc",
                location=f"paragraph {toc_idx}" if toc_idx is not None else "document TOC",
            )
        )
    return issues


def iter_body_blocks(document: Document):
    para_by_el = {p._p: ("p", i, p) for i, p in enumerate(document.paragraphs)}
    table_by_el = {t._tbl: ("tbl", i, t) for i, t in enumerate(document.tables)}
    for child in document.element.body:
        if child in para_by_el:
            yield para_by_el[child]
        elif child in table_by_el:
            yield table_by_el[child]


def table_has_vertical_borders(table) -> bool:
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        return False
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        return False
    for name in ("insideV", "left", "right"):
        node = borders.find(qn(f"w:{name}"))
        if node is not None and node.get(qn("w:val")) not in {None, "nil", "none"}:
            return True
    return False


def collect_table_issues(document: Document) -> list[Issue]:
    issues: list[Issue] = []
    blocks = list(iter_body_blocks(document))
    caption_ok_re = re.compile(r"^(?:表|Table)\s*\d+(?:[-.]\d+)?\s+\S.+$", re.I)
    caption_no_space_re = re.compile(r"^(?:表|Table)\s*\d+(?:[-.]\d+)?\S.+$", re.I)
    caption_only_re = re.compile(r"^(?:表|Table)\s*\d+(?:[-.]\d+)?\s*$", re.I)
    caption_any_re = re.compile(r"^(?:表|Table)\s*\d+(?:[-.]\d+)?", re.I)
    continued_re = re.compile(r"(续表|[(（]续[)）])")
    for pos, (kind, table_idx, table) in enumerate(blocks):
        if kind != "tbl":
            continue
        prev_text = ""
        next_text = ""
        if pos > 0 and blocks[pos - 1][0] == "p":
            prev_text = paragraph_text(blocks[pos - 1][2])
        if pos + 1 < len(blocks) and blocks[pos + 1][0] == "p":
            next_text = paragraph_text(blocks[pos + 1][2])

        caption_above_ok = caption_ok_re.match(prev_text) is not None
        if not caption_above_ok:
            issues.append(_issue_global("table_caption_missing", "table caption", f"表{table_idx + 1} 上方未检测到规范表题。", "表应有编号和表题，且置于表上方。", "table", f"table {table_idx + 1}", prev_text or next_text))
        if caption_any_re.match(next_text):
            issues.append(_issue_global("table_caption_position", "table caption", f"表{table_idx + 1} 下方检测到表题，位置错误。", "表题应置于表上方。", "table", f"table {table_idx + 1}", next_text))
        if prev_text and caption_no_space_re.match(prev_text) and not caption_ok_re.match(prev_text):
            issues.append(_issue_global("table_caption_number_format", "table caption", f"表{table_idx + 1} 表序与表名之间缺少空格。", "表序与表名之间应空一个空格。", "table", f"table {table_idx + 1}", prev_text))
        if prev_text and caption_only_re.match(prev_text):
            issues.append(_issue_global("table_caption_missing", "table caption", f"表{table_idx + 1} 仅检测到表序，缺少表名。", "表题应置于编号之后，表序与表名之间空一个空格。", "table", f"table {table_idx + 1}", prev_text))

        style_name = table.style.name if table.style is not None else ""
        if "grid" in style_name.lower() or table_has_vertical_borders(table):
            issues.append(_issue_global("table_three_line_style", "three-line table", f"表{table_idx + 1} 疑似全网格线/竖线。", "表格宜采用三线表；OOXML 边框检查为启发式，需要人工确认。", "table", f"table {table_idx + 1}"))

        if continued_re.search(prev_text) or continued_re.search(next_text):
            if continued_re.search(next_text):
                issues.append(_issue_global("table_caption_position", "continued table caption", f"续表 {table_idx + 1} 标记出现在表格下方。", "续表标记应置于表上方。", "table", f"table {table_idx + 1}", next_text))
            if not caption_any_re.match(prev_text):
                issues.append(_issue_global("table_continued_header", "continued table", f"续表 {table_idx + 1} 未检测到重复表编号。", "续表编号应重复原表编号，并在编号后跟“（续）”。", "table", f"table {table_idx + 1}", prev_text or next_text))
            first_row = " ".join(cell.text.strip() for cell in table.rows[0].cells) if table.rows else ""
            if not first_row:
                issues.append(_issue_global("table_continued_header", "continued table", f"续表 {table_idx + 1} 未检测到重复表头。", "续表应重复编号和表头；跨页续表头重复情况需人工确认。", "table", f"table {table_idx + 1}"))
            issues.append(_issue_global("table_continued_header", "continued table", f"续表 {table_idx + 1} 跨页续表与表头重复情况无法可靠自动判断。", "续表应重复表编号和表头；若 OOXML 无法可靠判断，需人工确认。", "table", f"table {table_idx + 1}"))

        issues.append(_issue_global("table_manual_review", "table manual review", f"表{table_idx + 1} 的自明性、随文出现、横读竖读与 CY/T170 符合性需人工复核。", "表的自明性、随文出现、横读竖读和 CY/T170 需人工复核。", "table", f"table {table_idx + 1}"))
    return issues


REFERENCE_272_FORMATS: dict[str, dict[str, object]] = {
    "M": {
        "type_marker": "M",
        "label": "专著",
        "template_text": "[序号] 主要责任者. 题名[M]. 出版地: 出版者, 出版年: 起止页码.",
        "required_elements": ["责任者", "题名", "类型标识[M]", "出版地", "出版者", "出版年"],
    },
    "C": {
        "type_marker": "C",
        "label": "论文集/会议录",
        "template_text": "[序号] 主要责任者. 题名[C]. 出版地: 出版者, 出版年: 起止页码.",
        "required_elements": ["责任者", "题名", "类型标识[C]", "出版地", "出版者", "出版年"],
    },
    "R": {
        "type_marker": "R",
        "label": "报告",
        "template_text": "[序号] 主要责任者. 题名[R]. 出版地: 出版者或机构, 年份.",
        "required_elements": ["责任者", "题名", "类型标识[R]", "出版地或机构", "年份"],
    },
    "D": {
        "type_marker": "D",
        "label": "学位论文",
        "template_text": "[序号] 主要责任者. 题名[D]. 保存地点: 保存单位, 年份.",
        "required_elements": ["责任者", "题名", "类型标识[D]", "保存地点", "保存单位", "年份"],
    },
    "P": {
        "type_marker": "P",
        "label": "专利",
        "template_text": "[序号] 专利申请者. 专利题名: 专利国别, 专利号[P]. 公告日期或公开日期.",
        "required_elements": ["责任者", "题名", "类型标识[P]", "专利号", "日期"],
    },
    "S": {
        "type_marker": "S",
        "label": "标准",
        "template_text": "[序号] 标准编号, 标准名称[S]. 出版地: 出版者, 出版年.",
        "required_elements": ["标准编号", "标准名称", "类型标识[S]", "年份"],
    },
    "J": {
        "type_marker": "J",
        "label": "期刊文章",
        "template_text": "[序号] 主要责任者. 题名[J]. 刊名, 年, 卷(期): 起止页码.",
        "required_elements": ["责任者", "题名", "类型标识[J]", "刊名", "年份", "卷期", "页码"],
    },
    "N": {
        "type_marker": "N",
        "label": "报纸文章",
        "template_text": "[序号] 主要责任者. 题名[N]. 报纸名, 出版日期(版次).",
        "required_elements": ["责任者", "题名", "类型标识[N]", "报纸名", "出版日期", "版次"],
    },
    "EB/OL": {
        "type_marker": "EB/OL",
        "label": "电子资源",
        "template_text": "[序号] 主要责任者. 题名[EB/OL]. 更新或修改日期[引用日期]. 获取和访问路径. DOI.",
        "required_elements": ["责任者", "题名", "类型标识[EB/OL]", "引用日期", "访问路径"],
    },
    "extracted": {
        "type_marker": "extracted",
        "label": "析出文献",
        "template_text": "[序号] 析出文献责任者. 析出题名[类型]//原文献责任者. 原文献题名. 出版地: 出版者, 出版年: 页码.",
        "required_elements": ["//", "前后要素基本完整", "出版信息要素"],
    },
}


def _reference_type_marker(text: str) -> str | None:
    marker_match = re.search(r"\[\s*([A-Z]+(?:/[A-Z]+)?)\s*\]", text)
    return marker_match.group(1) if marker_match else None


def _reference_text_number(text: str) -> int | None:
    match = _reference_entry_start_match(text)
    if not match:
        return None
    value = match.group(1) or match.group(2)
    if not value:
        return None
    try:
        return int(value)
    except ValueError:
        return None


def _numpr_val(num_pr, child_tag: str) -> str | None:
    child = num_pr.find(qn(child_tag))
    if child is None:
        return None
    return child.get(qn("w:val"))


def get_paragraph_numbering(paragraph) -> dict[str, object] | None:
    ppr = paragraph._p.pPr
    num_pr = None
    if ppr is not None:
        num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        style = getattr(paragraph, "style", None)
        style_elem = getattr(style, "_element", None)
        if style_elem is not None:
            style_ppr = style_elem.find(qn("w:pPr"))
            if style_ppr is not None:
                num_pr = style_ppr.find(qn("w:numPr"))
    if num_pr is None:
        return None
    num_id = _numpr_val(num_pr, "w:numId")
    ilvl = _numpr_val(num_pr, "w:ilvl") or "0"
    info: dict[str, object] = {
        "num_id": num_id,
        "ilvl": ilvl,
        "num_fmt": None,
        "lvl_text": None,
        "start": None,
    }
    numbering_part = getattr(paragraph.part, "numbering_part", None)
    if numbering_part is None or num_id is None:
        return info
    numbering = getattr(numbering_part, "_element", None)
    if numbering is None:
        return info

    num_elem = None
    for elem in numbering.findall(qn("w:num")):
        if elem.get(qn("w:numId")) == str(num_id):
            num_elem = elem
            break
    if num_elem is None:
        return info

    abstract_num_id = None
    abstract_node = num_elem.find(qn("w:abstractNumId"))
    if abstract_node is not None:
        abstract_num_id = abstract_node.get(qn("w:val"))

    lvl_elem = None
    override_level = None
    for lvl_override in num_elem.findall(qn("w:lvlOverride")):
        if lvl_override.get(qn("w:ilvl")) != str(ilvl):
            continue
        start_override = lvl_override.find(qn("w:startOverride"))
        if start_override is not None:
            try:
                override_level = int(start_override.get(qn("w:val")))
            except Exception:
                override_level = None
        lvl_elem = lvl_override.find(qn("w:lvl"))
        if lvl_elem is not None:
            break

    if lvl_elem is None and abstract_num_id is not None:
        abstract_elem = None
        for elem in numbering.findall(qn("w:abstractNum")):
            if elem.get(qn("w:abstractNumId")) == str(abstract_num_id):
                abstract_elem = elem
                break
        if abstract_elem is not None:
            for lvl in abstract_elem.findall(qn("w:lvl")):
                if lvl.get(qn("w:ilvl")) == str(ilvl):
                    lvl_elem = lvl
                    break

    if lvl_elem is None:
        return info

    num_fmt = lvl_elem.find(qn("w:numFmt"))
    lvl_text = lvl_elem.find(qn("w:lvlText"))
    start = lvl_elem.find(qn("w:start"))
    info["num_fmt"] = num_fmt.get(qn("w:val")) if num_fmt is not None else None
    info["lvl_text"] = lvl_text.get(qn("w:val")) if lvl_text is not None else None
    if override_level is not None:
        info["start"] = override_level
    elif start is not None:
        try:
            info["start"] = int(start.get(qn("w:val")))
        except Exception:
            info["start"] = None
    return info


def has_word_auto_numbering(paragraph) -> bool:
    info = get_paragraph_numbering(paragraph)
    if not info:
        return False
    num_fmt = str(info.get("num_fmt") or "").lower()
    if num_fmt in {"", "bullet", "none"}:
        return False
    lvl_text = str(info.get("lvl_text") or "")
    return bool(re.search(r"%\d+", lvl_text))


def get_reference_visual_index(
    paragraph,
    fallback_index: int,
    numbering_counters: dict[tuple[str, str], int],
) -> tuple[int, str]:
    text_no = _reference_text_number(paragraph_text(paragraph))
    if text_no is not None:
        return text_no, "text"
    if has_word_auto_numbering(paragraph):
        info = get_paragraph_numbering(paragraph) or {}
        key = (str(info.get("num_id") or ""), str(info.get("ilvl") or "0"))
        if key not in numbering_counters:
            start = info.get("start")
            try:
                numbering_counters[key] = int(start) - 1 if start is not None else 0
            except Exception:
                numbering_counters[key] = 0
        numbering_counters[key] += 1
        return numbering_counters[key], "word_auto_numbering"
    return fallback_index, "inferred"


def _clean_ref_field(value: str) -> str:
    return value.strip(" \t\r\n.,;:，；：")


def parse_reference_entry(text: str) -> dict[str, str]:
    marker = _reference_type_marker(text) or ""
    marker_match = re.search(r"\[\s*[A-Z]+(?:/[A-Z]+)?\s*\]", text)
    before_marker = text[: marker_match.start()] if marker_match else text
    after_marker = text[marker_match.end() :] if marker_match else ""

    before_marker = _clean_ref_field(before_marker)
    after_marker = _clean_ref_field(after_marker)
    parts = before_marker.split(".", 1)
    authors = _clean_ref_field(parts[0]) if parts else ""
    title = _clean_ref_field(parts[1]) if len(parts) > 1 else ""
    if not title and marker_match:
        title = _clean_ref_field(before_marker)

    source_title = ""
    year = ""
    tail_after_year = ""
    patent_no = ""
    url = ""
    doi = ""
    publisher = ""
    place = ""

    year_match = re.search(r"(19|20)\d{2}", after_marker)
    if year_match:
        year = year_match.group(0)
        source_title = _clean_ref_field(after_marker[: year_match.start()])
        tail_after_year = _clean_ref_field(after_marker[year_match.end() :])
    else:
        source_title = _clean_ref_field(after_marker)

    patent_match = re.search(r"\b(?:CN|US|EP|JP|WO)\s*[0-9A-Z.]+\b", text, re.I)
    if patent_match:
        patent_no = patent_match.group(0).replace(" ", "")

    path_match = re.search(r"(https?://\S+|www\.\S+)", text, re.I)
    if path_match:
        url = path_match.group(1).rstrip(".,;)")

    doi_match = re.search(r"\bDOI\s*[:：]?\s*([^\s,;]+)|\bdoi\.org/\S+", text, re.I)
    if doi_match:
        doi = (doi_match.group(1) or doi_match.group(0)).rstrip(".,;)")

    pub_match = re.search(r"(?:(?P<place>[^,:，：]+)\s*[:：]\s*)?(?P<publisher>[^,，]+)\s*[,，]\s*(?P<year>(?:19|20)\d{2})", after_marker)
    if pub_match:
        place = _clean_ref_field(pub_match.group("place") or "")
        publisher = _clean_ref_field(pub_match.group("publisher") or "")
        if not year:
            year = pub_match.group("year")

    return {
        "authors": authors,
        "title": title,
        "type_marker": marker,
        "source_title": source_title,
        "year": year,
        "tail_after_year": tail_after_year,
        "patent_no": patent_no,
        "url": url,
        "doi": doi,
        "publisher": publisher,
        "place": place,
    }


def _has_citation_date(text: str) -> bool:
    return bool(re.search(r"\[[0-9]{4}[-/年][^\]]*\]", text)) or ("[引用日期]" in text)


def _has_online_path(text: str) -> bool:
    return bool(re.search(r"https?://|www\.|doi\.org|DOI\s*[:：]|doi\s*[:：]", text, re.I))


def _has_doi_or_url(text: str) -> bool:
    return _has_online_path(text)


def _is_online_reference_marker(marker: str | None) -> bool:
    if not marker:
        return False
    marker_u = marker.upper()
    return marker_u in {"EB/OL", "DB/OL", "CP/OL", "J/OL"} or marker_u.endswith("/OL")


def _has_year(text: str) -> bool:
    return bool(re.search(r"(19|20)\d{2}", text))


def _has_page_range(text: str) -> bool:
    return bool(re.search(r"\d+\s*[-–—]\s*\d+", text))


def _has_volume_issue(text: str) -> bool:
    return bool(re.search(r"\d+\s*[（(]\s*\d+\s*[)）]", text))

def _has_journal_volume_or_page(text: str) -> bool:
    patterns = [
        r"\d+\s*\(\s*\d+\s*\)\s*:\s*\d+(?:\s*[-–—]\s*\d+)?",
        r"\d+\s*,\s*\d+\s*\(\s*\d+\s*\)\s*:\s*\d+",
        r"\d+\s*[-–—]\s*\d+",
        r":\s*\d+",
    ]
    return any(re.search(p, text) for p in patterns)


def _reference_missing_elements(marker: str, text: str) -> list[str]:
    parsed = parse_reference_entry(text)
    missing: list[str] = []
    if not parsed["authors"]:
        missing.append("责任者")
    if not parsed["title"]:
        missing.append("题名")
    if marker == "J":
        if not parsed["source_title"]:
            missing.append("刊名")
        if not parsed["year"]:
            missing.append("年份")
        if not _has_journal_volume_or_page(parsed["tail_after_year"] or text):
            missing.append("卷期或页码")
    elif marker in {"M", "C"}:
        if marker == "M" and not parsed["place"]:
            missing.append("出版地")
        if not parsed["publisher"]:
            missing.append("出版者")
        if not parsed["year"]:
            missing.append("出版年")
    elif marker == "R":
        if not parsed["year"]:
            missing.append("年份")
        if not parsed["source_title"] and not parsed["publisher"]:
            missing.append("出版地或机构")
    elif marker == "D":
        if not re.search(r"大学|学院|研究院|研究所|学校", text):
            missing.append("保存单位")
        if not (":" in text or "：" in text):
            missing.append("保存地")
        if not parsed["year"]:
            missing.append("年份")
    elif marker == "P":
        if not parsed["patent_no"]:
            missing.append("专利号")
        if not parsed["year"] and not re.search(r"(19|20)\d{2}[-/.年]\d{1,2}", text):
            missing.append("公告/公开日期")
    elif marker == "S":
        if not re.search(r"\b(?:GB|GB/T|ISO|IEC|IEEE|CY/T)[-\sA-Z0-9.]*", text, re.I):
            missing.append("标准编号")
        if not parsed["year"]:
            missing.append("年份")
    elif marker == "N":
        if not re.search(r"(19|20)\d{2}[-/.年]", text):
            missing.append("出版日期")
        if not re.search(r"[（(][A-Za-z0-9一二三四五六七八九十]+[)）]|第?\d+版", text):
            missing.append("版次")
    elif marker == "EB/OL":
        if not _has_citation_date(text):
            missing.append("引用日期")
        if not _has_online_path(text):
            missing.append("访问路径")
    return missing


def _matches_any_272_format(text: str, marker: str | None) -> tuple[list[str], dict[str, list[str]]]:
    matched: list[str] = []
    missing_map: dict[str, list[str]] = {}
    markers = [marker] if marker in REFERENCE_272_FORMATS else [m for m in REFERENCE_272_FORMATS if m != "extracted"]
    if "//" in text:
        markers = list(dict.fromkeys(markers + ["extracted"]))
    for mk in markers:
        if mk == "extracted":
            miss = []
            parts = text.split("//", 1)
            if len(parts) != 2:
                miss.append("//")
            else:
                if _reference_type_marker(parts[0]) is None:
                    miss.append("析出部分类型标识")
                if ":" not in parts[1] and "：" not in parts[1]:
                    miss.append("原文献出版信息")
                if not _has_year(parts[1]):
                    miss.append("原文献年份")
            if not miss:
                matched.append(mk)
            missing_map[mk] = miss
            continue
        miss = _reference_missing_elements(mk, text)
        if not miss:
            matched.append(mk)
        missing_map[mk] = miss
    return matched, missing_map




def _reference_issue(paragraph_index: int, rule_key: str, text_type: str, current: str, expected: str, excerpt: str = "", category: str = "references") -> Issue:
    location = f"paragraph {paragraph_index}" if paragraph_index >= 0 else "references"
    return Issue(
        paragraph_index=paragraph_index,
        rule_key=rule_key,
        text_type=text_type,
        text_excerpt=excerpt or location,
        current=current,
        expected=expected,
        message=f"文本类型：{text_type}\n当前位置：{location}\n当前情况：{current}\n应改为：{expected}",
        category=category,
        location=location,
    )

def _reference_entries_from_regions(paragraphs, texts: list[str], regions: list[str]) -> tuple[list[dict[str, object]], list[Issue]]:
    entries: list[dict[str, object]] = []
    issues: list[Issue] = []
    current: dict[str, object] | None = None
    inferred_no = 1
    numbering_counters: dict[tuple[str, str], int] = {}
    for idx, paragraph in enumerate(paragraphs):
        if idx >= len(regions) or regions[idx] != "references":
            continue
        text = paragraph_text(paragraph).strip()
        if not text:
            continue
        if compact_text(text) == "参考文献":
            continue

        text_no = _reference_text_number(text)
        has_auto_no = has_word_auto_numbering(paragraph)
        looks_like_entry = bool(text_no is not None or has_auto_no or _looks_like_reference_entry_text(text))

        if looks_like_entry:
            visual_index, numbering_source = get_reference_visual_index(paragraph, inferred_no, numbering_counters)
            numbering_source_for_summary = numbering_source
            if numbering_source == "word_auto_numbering":
                key = (str((get_paragraph_numbering(paragraph) or {}).get("num_id") or ""), str((get_paragraph_numbering(paragraph) or {}).get("ilvl") or "0"))
                if key in numbering_counters and numbering_counters[key] == visual_index:
                    numbering_source_for_summary = "word_auto_numbering_assumed_continuous"
            current = {
                "start_idx": idx,
                "number": visual_index,
                "visual_index": visual_index,
                "numbering_source": "text" if text_no is not None else numbering_source_for_summary,
                "parts": [text],
                "paragraph_indices": [idx],
            }
            entries.append(current)
            inferred_no = max(inferred_no, int(visual_index) + 1)
            if text_no is None and not has_auto_no:
                issues.append(_reference_issue(idx, "reference_entry_number_missing", "reference entry", "条目缺少序号。", "参考文献条目应以连续序号开头。", text[:120]))
            continue

        if current is None:
            issues.append(_reference_issue(idx, "reference_entry_number_missing", "reference entry", "条目缺少序号。", "参考文献条目应以连续序号开头。", text[:120]))
            continue

        current["parts"].append(text)
        current["paragraph_indices"].append(idx)
        issues.append(_reference_issue(idx, "reference_entry_continuation_review", "reference entry continuation", "疑似参考文献续行，需要人工确认。", "参考文献续行应与上一条合并核对著录格式。", text[:120]))
    return entries, issues


def _reference_entry_start_match(text: str):
    return re.match(r"^\s*(?:\[(\d+)\]|(\d+)(?:[\.\u3001]|\s+))\s*\S+", text)


def _looks_like_reference_entry_text(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if _reference_type_marker(stripped):
        return True
    has_year = bool(re.search(r"(19|20)\d{2}", stripped))
    has_ref_punc = "." in stripped or "：" in stripped or ":" in stripped
    return has_year and has_ref_punc and len(stripped) >= 12


def _reference_272_check_summary(entries: list[dict[str, object]]) -> dict[str, object]:
    summary_entries: list[dict[str, object]] = []
    matched = 0
    unmatched = 0
    mismatch = 0
    for entry in entries:
        text = " ".join(str(x).strip() for x in entry["parts"]).strip()
        marker = _reference_type_marker(text)
        matched_formats, missing_map = _matches_any_272_format(text, marker)
        status = "matched"
        if marker and marker in REFERENCE_272_FORMATS and marker not in matched_formats:
            status = "type_mismatch"
            mismatch += 1
        elif not matched_formats:
            status = "no_match"
            unmatched += 1
        else:
            matched += 1
        summary_entries.append(
            {
                "index": int(entry["number"]),
                "visual_index": int(entry.get("visual_index", entry["number"])),
                "numbering_source": str(entry.get("numbering_source", "inferred")),
                "paragraph_index": int(entry.get("start_idx", -1)),
                "text_excerpt": text[:160],
                "detected_type": marker or "",
                "matched_template": matched_formats[0] if matched_formats else "",
                "status": status,
                "missing_fields": missing_map.get(marker or "", []),
            }
        )
    return {
        "total_entries": len(entries),
        "matched_entries": matched,
        "unmatched_entries": unmatched,
        "type_mismatch_entries": mismatch,
        "entries": summary_entries,
    }


def _reference_273_check_summary(entries: list[dict[str, object]]) -> dict[str, object]:
    summary_entries: list[dict[str, object]] = []
    online_entries = 0
    author_et_al_review_entries = 0
    citation_date_missing_entries = 0
    access_path_missing_entries = 0
    doi_or_url_missing_entries = 0
    for entry in entries:
        text = " ".join(str(x).strip() for x in entry["parts"]).strip()
        marker = _reference_type_marker(text)
        head_part = text.split(".", 1)[0]
        author_items = [x.strip() for x in re.split(r"[,，、;；]|(?:\band\b)|&", head_part) if x.strip()]
        needs_et_al_review = (len(author_items) >= 4 and ("等" not in head_part and "et al" not in text.lower())) or ("等" in head_part and len(author_items) < 3)

        is_online = _is_online_reference_marker(marker) or ("http" in text.lower()) or ("doi" in text.lower())
        citation_date_missing = False
        access_path_missing = False
        doi_or_url_missing = False
        if is_online:
            online_entries += 1
            citation_date_missing = not _has_citation_date(text)
            access_path_missing = not _has_online_path(text)
            doi_or_url_missing = not _has_doi_or_url(text)
            citation_date_missing_entries += 1 if citation_date_missing else 0
            access_path_missing_entries += 1 if access_path_missing else 0
            doi_or_url_missing_entries += 1 if doi_or_url_missing else 0

        if needs_et_al_review:
            author_et_al_review_entries += 1

        summary_entries.append(
            {
                "index": int(entry["number"]),
                "visual_index": int(entry.get("visual_index", entry["number"])),
                "paragraph_index": int(entry.get("start_idx", -1)),
                "detected_type": marker or "",
                "text_excerpt": text[:160],
                "is_online_entry": is_online,
                "author_et_al_review": needs_et_al_review,
                "citation_date_missing": citation_date_missing,
                "access_path_missing": access_path_missing,
                "doi_or_url_missing": doi_or_url_missing,
            }
        )

    return {
        "total_entries": len(entries),
        "online_entries": online_entries,
        "author_et_al_review_entries": author_et_al_review_entries,
        "citation_date_missing_entries": citation_date_missing_entries,
        "access_path_missing_entries": access_path_missing_entries,
        "doi_or_url_missing_entries": doi_or_url_missing_entries,
        "entries": summary_entries,
    }


def collect_reference_issues(document: Document) -> list[Issue]:
    issues: list[Issue] = []
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    entries, entry_issues = _reference_entries_from_regions(document.paragraphs, texts, regions)
    issues.extend(entry_issues)
    if not entries:
        reference_title_idx = next((i for i, t in enumerate(texts) if re.match(r"^\s*参考文献\s*$", t.strip())), -1)
        if reference_title_idx >= 0:
            fallback_regions = ["other"] * len(texts)
            for i in range(reference_title_idx + 1, len(texts)):
                fallback_regions[i] = "references"
            entries, entry_issues = _reference_entries_from_regions(document.paragraphs, texts, fallback_regions)
            issues.extend(entry_issues)
    has_reference_title = any(idx < len(regions) and regions[idx] == "references" and compact_text(paragraph_text(p)) == "参考文献" for idx, p in enumerate(document.paragraphs))
    if not has_reference_title:
        has_reference_title = any(re.match(r"^\s*参考文献\s*$", paragraph_text(p).strip()) for p in document.paragraphs)
    if has_reference_title and not entries:
        issues.append(_issue_global("reference_entries_not_detected", "reference entries detection", "已检测到参考文献标题，但未能识别参考文献条目。", "请检查分段或编号格式，使条目能被识别并检查。", "references", "references section"))

    expected_no = 1
    for entry in entries:
        idx = int(entry["start_idx"])
        text = " ".join(str(x).strip() for x in entry["parts"]).strip()
        current_no = int(entry.get("visual_index", entry["number"]))
        numbering_source = str(entry.get("numbering_source", "inferred"))
        if numbering_source != "inferred" and current_no != expected_no:
            issues.append(_reference_issue(idx, "reference_sequence", "reference entry", f"序号不连续：期望[{expected_no}]，实际[{current_no}]。", "应使用连续的序号，并符合规范2.7.2和2.7.3。", text[:120]))
            expected_no = current_no
        expected_no += 1

        marker = _reference_type_marker(text)
        if marker is None:
            issues.append(_reference_issue(idx, "reference_type_marker", "reference entry", "缺少文献类型标识（如[M]/[J]/[D]/[EB/OL]）。", "应包含类型标识并符合规范2.7.2和2.7.3。", text[:120]))
        matched_formats, missing_map = _matches_any_272_format(text, marker)
        if marker and marker in REFERENCE_272_FORMATS and marker not in matched_formats:
            template = str(REFERENCE_272_FORMATS[marker]["template_text"])
            missing_elements = "、".join(missing_map.get(marker, [])) or "未识别完整要素"
            issues.append(_reference_issue(idx, "reference_272_type_format_mismatch", "reference entry 2.7.2 type format", f"当前条目类型标识为[{marker}]，但缺失要素：{missing_elements}。", "该条目应符合其文献类型标识对应的2.7.2著录格式。", text[:200]))
            issues.append(_reference_issue(idx, "reference_entry_format", "reference entry", f"类型[{marker}]建议模板：{template}", "该检查为启发式检查，需要人工确认。", text[:200]))

        if not matched_formats and not (marker and marker in REFERENCE_272_FORMATS):
            tried = "/".join([k for k in REFERENCE_272_FORMATS.keys()])
            missing_summary = "; ".join(f"{k}: {('、'.join(v) if v else '要素未完整匹配')}" for k, v in missing_map.items())
            issues.append(_reference_issue(idx, "reference_272_format_no_match", "reference entry 2.7.2 format", f"当前条目未匹配{tried}任一种格式；类型标识={marker or '无'}；缺失要素：{missing_summary}", "参考文献条目应符合2.7.2中至少一种主要参考文献著录格式。", text[:220]))

        is_online = _is_online_reference_marker(marker) or ("http" in text.lower()) or ("doi" in text.lower())
        if is_online and not _has_online_path(text):
            issues.append(_reference_issue(idx, "reference_273_online_access_missing", "reference entry 2.7.3 online access", "在线文献缺少访问路径。", "网上获取文献应补充获取和访问路径。", text[:200]))
        if is_online and not _has_citation_date(text):
            issues.append(_reference_issue(idx, "reference_273_citation_date_missing", "reference entry 2.7.3 citation date", "在线文献缺少引用日期。", "网上获取文献应补充[引用日期]。", text[:200]))
        if is_online and not _has_doi_or_url(text):
            issues.append(_reference_issue(idx, "reference_273_doi_or_url_missing", "reference entry 2.7.3 DOI/URL", "在线文献缺少 DOI 或 URL。", "网上获取文献应包含 DOI 或 URL。", text[:200]))

        head_part = text.split(".", 1)[0]
        author_items = [x.strip() for x in re.split(r"[,，、;；]|(?:\band\b)|&", head_part) if x.strip()]
        if len(author_items) >= 4 and ("等" not in head_part and "et al" not in text.lower()):
            issues.append(_reference_issue(idx, "reference_273_author_et_al_review", "reference entry 2.7.3 author list", "责任者疑似超过3人但未标注“等”或“et al.”。", "多人作者建议列前3位后加“等”；该检查为启发式，需要人工确认。", text[:200]))
        if "等" in head_part and len(author_items) < 3:
            issues.append(_reference_issue(idx, "reference_273_author_et_al_review", "reference entry 2.7.3 author list", "责任者包含“等”，但前置作者数量疑似不足3位。", "作者数量与“等”用法需人工确认；该检查为启发式。", text[:200]))
    return issues

def _shape_cm(value) -> float | None:
    if value is None:
        return None
    try:
        return float(value.cm)
    except Exception:
        return None


def collect_image_size_issues(document: Document) -> list[Issue]:
    issues: list[Issue] = []
    for idx, shape in enumerate(getattr(document, "inline_shapes", []), start=1):
        h_cm = _shape_cm(getattr(shape, "height", None))
        w_cm = _shape_cm(getattr(shape, "width", None))
        if h_cm is None or w_cm is None:
            issues.append(
                _issue_global(
                    "image_size_advisory",
                    "image size",
                    "无法通过 OOXML 获取图片尺寸，需人工确认。",
                    "图片一般高6cm×宽8cm；高度不得超过16cm。",
                    "image",
                    f"image {idx}",
                )
            )
            continue
        if h_cm > 16:
            new_h = 16.0
            new_w = w_cm * (new_h / h_cm) if h_cm else w_cm
            issues.append(
                _issue_global(
                    "image_height_limit",
                    "image size",
                    f"图片高度超过16cm，当前尺寸 {h_cm:.2f}cm×{w_cm:.2f}cm。",
                    f"图片高度超过16cm，已按比例缩放至16cm高（缩放后约 {new_h:.2f}cm×{new_w:.2f}cm），需人工复核排版。",
                    "image",
                    f"image {idx}",
                )
            )
        elif abs(h_cm - 6.0) > 0.2 or abs(w_cm - 8.0) > 0.2:
            issues.append(
                _issue_global(
                    "image_size_advisory",
                    "image size",
                    f"当前尺寸 {h_cm:.2f}cm×{w_cm:.2f}cm。",
                    "图片一般建议高6cm×宽8cm；当前尺寸如因图片量或排版需要缩放，请人工确认。",
                    "image",
                    f"image {idx}",
                )
            )
    return issues


def count_document_images(document: Document) -> int:
    rel_ids: set[str] = set()
    for node in document._element.iter():
        tag = str(node.tag)
        if tag.endswith("}blip"):
            rid = node.get(qn("r:embed")) or node.get(qn("r:link"))
            if rid:
                rel_ids.add(rid)
        elif tag.endswith("}imagedata"):
            rid = node.get(qn("r:id"))
            if rid:
                rel_ids.add(rid)
    return len(rel_ids)


def count_figure_captions(document: Document) -> int:
    pattern = re.compile(r"^\s*图\s*\d+(?:[-.]\d+)?\s+")
    return sum(1 for paragraph in document.paragraphs if pattern.match(paragraph_text(paragraph).strip()))


def collect_figure_caption_image_issues(document: Document) -> list[Issue]:
    issues: list[Issue] = []
    pattern = re.compile(r"^\s*图\s*\d+(?:[-.]\d+)?\s+")
    for idx, paragraph in enumerate(document.paragraphs):
        text = paragraph_text(paragraph).strip()
        if not pattern.match(text):
            continue
        has_nearby_image = False
        for lookback in range(max(0, idx - 3), idx):
            if _paragraph_has_image_payload(document.paragraphs[lookback]):
                has_nearby_image = True
                break
        if not has_nearby_image:
            issues.append(
                Issue(
                    paragraph_index=idx,
                    rule_key="figure_caption_without_nearby_image",
                    text_type="figure caption",
                    text_excerpt=text[:160],
                    current="检测到图题附近没有对应图片。",
                    expected="图题前 1-3 个段落内应存在对应图片或图片锚点。",
                    message="检测到图题附近没有对应图片，可能图片丢失或图题位置错误。",
                    category="image",
                    location=f"paragraph {idx}",
                )
            )
    return issues


def enforce_image_height_limit(document: Document) -> int:
    scaled = 0
    for shape in getattr(document, "inline_shapes", []):
        h_cm = _shape_cm(getattr(shape, "height", None))
        w_cm = _shape_cm(getattr(shape, "width", None))
        if h_cm is None or w_cm is None or h_cm <= 16:
            continue
        new_h = 16.0
        new_w = w_cm * (new_h / h_cm) if h_cm else w_cm
        shape.height = Cm(new_h)
        shape.width = Cm(new_w)
        scaled += 1
    return scaled


def collect_issues(document: Document) -> list[Issue]:
    issues: list[Issue] = []
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    structure = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))
    regions = structure["regions"]
    title_overrides = detect_abstract_thesis_titles(texts)
    for idx, paragraph in enumerate(document.paragraphs):
        text = texts[idx]
        if not text:
            continue
        key = classify_checked_paragraph(paragraph, idx, text, regions[idx], title_overrides)
        if key is None:
            continue
        rule = RULES[key]
        if not paragraph_matches(paragraph, rule):
            current = current_format(paragraph, rule)
            issues.append(
                Issue(
                    paragraph_index=idx,
                    rule_key=rule.key,
                    text_type=rule.label,
                    text_excerpt=text[:160],
                    current=current,
                    expected=rule.expected,
                    message=f"文本类型：{rule.label}\n当前格式：{current}\n应改为：{rule.expected}",
                    category=rule.category,
                    location=f"paragraph {idx}",
                )
            )
        if key == "body":
            normalized_text = normalize_body_cjk_spacing(text)
            if normalized_text != text:
                issues.append(
                    Issue(
                        paragraph_index=idx,
                        rule_key="body_cjk_spacing",
                        text_type="body paragraph CJK spacing",
                        text_excerpt=text[:160],
                        current=text,
                        expected=BODY_CJK_SPACING_EXPECTED,
                        message=(
                            "文本类型：body paragraph CJK spacing\n"
                            f"当前问题：{text}\n"
                            f"应改为：{normalized_text}"
                        ),
                        category="body-spacing",
                        location=f"paragraph {idx}",
                    )
                )
        issues.extend(collect_run_format_issues(paragraph, idx, rule, text))
    issues.extend(collect_section_issues(document))
    issues.extend(collect_abstract_title_issues(document))
    issues.extend(collect_toc_issues(document))
    issues.extend(detect_toc_manual_review_issues(document))
    issues.extend(collect_table_issues(document))
    issues.extend(collect_reference_issues(document))
    issues.extend(collect_image_size_issues(document))
    issues.extend(collect_figure_caption_image_issues(document))
    empty_issues, _to_remove = collect_empty_paragraph_issues(document)
    issues.extend(empty_issues)
    boundary_issues, _removed_preview = cleanup_module_boundary_blank_paragraphs(document, apply_changes=False)
    issues.extend(boundary_issues)
    heading_issues, _removed_preview2 = cleanup_body_heading_following_blank_paragraphs(document, apply_changes=False)
    issues.extend(heading_issues)
    return issues


def apply_supported_rules(document: Document, allow_layout_fixes: bool) -> None:
    apply_page_setup(document, allow_layout_fixes=allow_layout_fixes)
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    structure = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))
    regions = structure["regions"]
    title_overrides = detect_abstract_thesis_titles(texts)
    for idx, paragraph in enumerate(document.paragraphs):
        text = texts[idx]
        if not text:
            continue
        key = classify_checked_paragraph(paragraph, idx, text, regions[idx], title_overrides)
        if key is not None:
            apply_rule(paragraph, RULES[key])
            if key == "body":
                normalized_text = normalize_body_cjk_spacing(paragraph.text)
                if normalized_text != paragraph.text:
                    _apply_text_to_runs(paragraph, normalized_text)
    texts_after = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions_after = analyze_section_sequence(texts_after, has_toc_field=document_has_toc_field(document))["regions"]
    ref_entries, _ref_issues = _reference_entries_from_regions(document.paragraphs, texts_after, regions_after)
    for entry in ref_entries:
        for pidx in entry.get("paragraph_indices", []):
            if 0 <= pidx < len(document.paragraphs):
                apply_rule(document.paragraphs[pidx], RULES["reference_entry"])
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    cleanup_visible_blank_paragraphs_after_layout(document)
    enforce_image_height_limit(document)


def qname(ns: str, tag: str) -> str:
    return f"{{{ns}}}{tag}"


def ensure_comments_parts(workdir: Path) -> ET.ElementTree:
    comments_path = workdir / "word" / "comments.xml"
    if comments_path.exists():
        return ET.parse(comments_path)
    root = ET.Element(qname(W_NS, "comments"))
    tree = ET.ElementTree(root)
    comments_path.parent.mkdir(parents=True, exist_ok=True)
    tree.write(comments_path, encoding="utf-8", xml_declaration=True)
    return tree


def ensure_content_type(workdir: Path) -> None:
    path = workdir / "[Content_Types].xml"
    tree = ET.parse(path)
    root = tree.getroot()
    exists = any(child.attrib.get("PartName") == "/word/comments.xml" for child in root)
    if not exists:
        override = ET.SubElement(root, qname(CONTENT_NS, "Override"))
        override.set("PartName", "/word/comments.xml")
        override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
    tree.write(path, encoding="utf-8", xml_declaration=True)


def ensure_comments_relationship(workdir: Path) -> None:
    rels = workdir / "word" / "_rels" / "document.xml.rels"
    tree = ET.parse(rels)
    root = tree.getroot()
    rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
    if any(child.attrib.get("Type") == rel_type for child in root):
        tree.write(rels, encoding="utf-8", xml_declaration=True)
        return
    used = {child.attrib.get("Id", "") for child in root}
    n = 1
    while f"rId{n}" in used:
        n += 1
    rel = ET.SubElement(root, qname(PKG_REL_NS, "Relationship"))
    rel.set("Id", f"rId{n}")
    rel.set("Type", rel_type)
    rel.set("Target", "comments.xml")
    tree.write(rels, encoding="utf-8", xml_declaration=True)


def add_comment_node(comments_root: ET.Element, comment_id: int, text: str) -> None:
    comment = ET.SubElement(comments_root, qname(W_NS, "comment"))
    comment.set(qname(W_NS, "id"), str(comment_id))
    comment.set(qname(W_NS, "author"), "thesis format skill")
    comment.set(qname(W_NS, "date"), datetime.now(timezone.utc).isoformat())
    p = ET.SubElement(comment, qname(W_NS, "p"))
    for line_no, line in enumerate(text.splitlines()):
        if line_no:
            ET.SubElement(p, qname(W_NS, "br"))
        r = ET.SubElement(p, qname(W_NS, "r"))
        t = ET.SubElement(r, qname(W_NS, "t"))
        t.text = line


def add_comment_marker(paragraph: ET.Element, comment_id: int) -> None:
    start = ET.Element(qname(W_NS, "commentRangeStart"))
    start.set(qname(W_NS, "id"), str(comment_id))
    end = ET.Element(qname(W_NS, "commentRangeEnd"))
    end.set(qname(W_NS, "id"), str(comment_id))
    ref_run = ET.Element(qname(W_NS, "r"))
    ref = ET.SubElement(ref_run, qname(W_NS, "commentReference"))
    ref.set(qname(W_NS, "id"), str(comment_id))
    paragraph.insert(0, start)
    paragraph.append(end)
    paragraph.append(ref_run)


def _comment_title(issue: Issue) -> str:
    labels = {
        "keywords": "关键词段落",
        "keywords_zh": "中文关键词段落",
        "keywords_en": "英文关键词段落",
        "keywords_runs": "关键词段落",
        "keywords_zh_runs": "中文关键词段落",
        "keywords_en_runs": "英文关键词段落",
        "thesis_title_zh": "中文论文题目",
        "thesis_title_en": "英文论文题目",
        "abstract_title_zh": "中文摘要标题",
        "abstract_body_zh": "中文摘要正文",
        "abstract_title_en": "英文摘要标题",
        "abstract_body_en": "英文摘要正文",
        "chapter": "章标题",
        "heading2": "二级标题",
        "heading3": "三级标题",
        "body": "正文段落",
        "figure_caption": "图题",
        "table_caption": "表题",
        "reference_title": "参考文献标题",
        "reference_entry": "参考文献条目",
        "appendix_title": "附录标题",
        "appendix_body": "附录正文",
    }
    key = issue.rule_key.removesuffix("_runs")
    return labels.get(issue.rule_key) or labels.get(key) or issue.text_type


def _short_expected_items(issue: Issue) -> list[str]:
    key = issue.rule_key.removesuffix("_runs")
    if key in {"keywords", "keywords_zh", "keywords_en"}:
        return [
            "中文用宋体小四，英文/数字用 Times New Roman 小四",
            "关键词标签加粗",
            "关键词之间用分号分隔",
            "段落行距为固定 20 磅",
        ]
    if key == "thesis_title_zh":
        return [
            "中文论文题目使用宋体三号 16pt，加粗，居中",
            "行距为固定 20 磅",
        ]
    if key == "thesis_title_en":
        return [
            "英文论文题目使用 Times New Roman 三号 16pt，加粗，居中",
            "行距为固定 20 磅",
        ]
    body_items = {
        "body": [
            "中文用宋体小四 12pt，英文/数字用 Times New Roman 12pt",
            "左对齐",
            "首行缩进 2 个汉字符，段前 0 磅，段后 0 磅",
            "行距为固定 20 磅",
        ],
        "abstract_body_zh": [
            "中文用宋体小四 12pt",
            "行距为固定 20 磅",
        ],
        "abstract_body_en": [
            "英文用 Times New Roman 小四 12pt",
            "行距为固定 20 磅",
        ],
        "appendix_body": [
            "中文用宋体小四 12pt，英文/数字用 Times New Roman 12pt",
            "两端对齐",
            "首行缩进 2 个汉字符",
            "行距为固定 20 磅",
        ],
        "acknowledgement_body": [
            "中文用宋体小四 12pt",
            "行距为固定 20 磅",
        ],
    }
    if key in body_items:
        return body_items[key]
    heading_items = {
        "chapter": [
            "宋体三号 16pt，加粗，居中",
            "行距为固定 20 磅，段前 0.5 行，段后 0.5 行",
            "章序号与章名之间空一格",
        ],
        "heading2": [
            "宋体四号 14pt，加粗，左对齐",
            "行距为固定 20 磅，段前 0.5 行，段后 0.5 行",
            "序号与题名之间空一格",
        ],
        "heading3": [
            "宋体小四 12pt，加粗，左对齐",
            "行距为固定 20 磅，段前 0.5 行，段后 0.5 行",
            "序号与题名之间空一格",
        ],
        "abstract_title_zh": [
            "中文摘要标题使用宋体三号 16pt，加粗，居中",
            "行距为固定 20 磅",
        ],
        "abstract_title_en": [
            "英文摘要标题使用 Times New Roman 三号 16pt，加粗，居中",
            "行距为固定 20 磅",
        ],
        "toc_title": [
            "目录标题使用宋体三号 16pt，加粗，居中",
            "行距为固定 20 磅",
        ],
    }
    if key in heading_items:
        return heading_items[key]
    if key in {"figure_caption", "table_caption"}:
        return [
            "图表题使用宋体五号",
            "居中，行距为固定 20 磅",
            "编号和标题之间保留一个空格",
        ]
    if key.startswith("reference"):
        return [
            "“参考文献”标题使用宋体三号 16pt，加粗，按章标题格式排版",
            "条目中文用宋体小四 12pt，英文用 Times New Roman 12pt",
            "条目行距为固定 20 磅，续行缩进 2 个汉字符左对齐",
            "不在自动检查中判断 GB/T 7714 著录内容细节",
        ]
    return [issue.expected]


def _extract_summary_items(text: str) -> list[str]:
    return [p.strip(" ；;。") for p in re.split(r"[;\n]|；", text) if p.strip()]


def _clean_expected_numeric(value: str) -> str:
    text = value.strip()
    text = re.sub(r"\.0(?=[^\d]|$)", "", text)
    return text


def humanize_problem_item(item: str) -> str:
    text = re.sub(r"^\d+\.\s*", "", (item or "").strip())
    if not text:
        return "检测到该段可能不符合规范，建议人工复核。"
    if "继承/未直接设置" in text or "未直接设置" in text or "继承" in text:
        text = text.replace("继承/未直接设置", "").replace("未直接设置", "").replace("继承", "").strip(" ：:，,")
    patterns: list[tuple[re.Pattern[str], str]] = [
        (re.compile(r"^固定行距[:：].*应为\s*([0-9.]+)\s*磅"), "行距不符合要求，应为固定 {v} 磅。"),
        (re.compile(r"^中文字体[:：].*应为\s*([^\s，。;；]+)"), "中文字体不符合要求，应使用{v}。"),
        (re.compile(r"^西文字体[:：].*应为\s*([^\s，。;；]+(?:\s+[^\s，。;；]+)*)"), "西文字体不符合要求，应使用{v}。"),
        (re.compile(r"^字号[:：].*应为\s*([0-9.]+)\s*pt", re.IGNORECASE), "字号不符合要求，应为 {v}pt。"),
        (re.compile(r"^加粗[:：].*应为\s*是"), "文字应加粗。"),
        (re.compile(r"^加粗[:：].*应为\s*否"), "文字不应加粗。"),
        (re.compile(r"^对齐方式[:：].*应为\s*([^\s，。;；]+)"), "对齐方式不符合要求，应{v}。"),
        (re.compile(r"^段前[:：].*应为\s*([0-9.]+)\s*行"), "段前间距不符合要求，应为 {v} 行。"),
        (re.compile(r"^段后[:：].*应为\s*([0-9.]+)\s*行"), "段后间距不符合要求，应为 {v} 行。"),
        (re.compile(r"^段前[:：].*应为\s*([0-9.]+)\s*pt", re.IGNORECASE), "段前间距不符合要求，应为 {v}pt。"),
        (re.compile(r"^段后[:：].*应为\s*([0-9.]+)\s*pt", re.IGNORECASE), "段后间距不符合要求，应为 {v}pt。"),
        (re.compile(r"^首行缩进[:：].*应为\s*([0-9.]+)\s*字符"), "首行缩进不符合要求，应缩进 {v} 个汉字符。"),
        (re.compile(r"^悬挂缩进[:：].*应为\s*([0-9.]+)\s*字符"), "参考文献续行缩进不符合要求，应使用 Word 的“悬挂缩进 {v} 字符”。"),
    ]
    for pattern, template in patterns:
        match = pattern.search(text)
        if match:
            value = _clean_expected_numeric(match.group(1)) if match.groups() else ""
            return template.format(v=value)
    if text.startswith("中文字体当前为 "):
        current_font = text.removeprefix("中文字体当前为 ").split("，", 1)[0].strip()
        if current_font:
            return f"中文字体使用了{current_font}，应改为宋体。"
        return "中文字体不符合要求，应使用宋体。"
    if text.startswith("西文字体当前为 "):
        current_font = text.removeprefix("西文字体当前为 ").split("，", 1)[0].strip()
        norm_font = normalize_font_name(current_font)
        if norm_font == "Times New Roman":
            return ""
        if current_font:
            return f"西文字体使用了{current_font}，应改为 Times New Roman。"
        return "西文字体不符合要求，应使用 Times New Roman。"
    if "目录" in text and "人工复核" in text:
        return "目录分页或页码需要人工复核，工具不自动修改。"
    if "风险空段" in text:
        return "检测到可能影响分页的空白段，建议人工复核。"
    if text.startswith("文字片段=") or text.startswith("片段="):
        return ""
    if "段落样式" in text:
        return ""
    return text


def humanize_issue_current(current: str | None) -> str:
    text = (current or "").strip()
    if not text:
        return "检测到该段可能不符合规范。"
    if "疑似不符合项：" in text:
        items = _issue_problem_lines(text)
        if items:
            return "；".join(items)
        return "检测到该段可能不符合规范。"
    if "文字片段=" in text:
        return "检测到该段存在文字级格式问题。"
    cleaned = text.replace("继承/未直接设置", "").replace("未直接设置", "").replace("段落样式：Normal", "").strip(" ;；")
    return cleaned or "检测到该段可能不符合规范。"


def _extract_run_problem_items(current: str) -> list[str]:
    items: list[str] = []
    parts = [part.strip() for part in current.split(" | ") if part.strip()]
    for part in parts:
        if part.startswith("... +"):
            continue
        segments = _extract_summary_items(part)
        segment_map: dict[str, str] = {}
        for seg in segments:
            if "=" in seg:
                key, value = seg.split("=", 1)
                segment_map[key.strip()] = value.strip().strip("'")
        east = segment_map.get("中文字体")
        if east and east != "继承/未直接设置":
            items.append(f"中文字体当前为 {east}，建议核对是否符合规则要求。")
        ascii_font = segment_map.get("西文字体")
        norm_ascii = normalize_font_name(ascii_font)
        if ascii_font and ascii_font != "继承/未直接设置" and norm_ascii != "Times New Roman":
            items.append(f"西文字体当前为 {ascii_font}，建议核对是否符合规则要求。")
    if not items:
        return ["检测到该段可能不符合规范，建议人工复核。"]
    # 去重并保持顺序
    unique: list[str] = []
    seen: set[str] = set()
    for item in items:
        if item not in seen:
            unique.append(item)
            seen.add(item)
    return unique


def _issue_problem_lines(current: str | None) -> list[str]:
    text = (current or "").strip()
    if not text:
        return ["检测到该段可能不符合规范，建议人工复核。"]
    if "疑似不符合项：" in text:
        tail = text.split("疑似不符合项：", 1)[1]
        stop_markers = ["; 段落样式：", "；段落样式：", ";段落样式：", "； 段落样式："]
        cut = tail
        for marker in stop_markers:
            if marker in cut:
                cut = cut.split(marker, 1)[0]
        diffs = _extract_summary_items(cut)
        humanized = [humanize_problem_item(x) for x in diffs]
        filtered = [x for x in humanized if x]
        return filtered or ["检测到该段可能不符合规范，建议人工复核。"]
    if "文字片段=" in text:
        humanized = [humanize_problem_item(x) for x in _extract_run_problem_items(text)]
        filtered = [x for x in humanized if x]
        return filtered or ["检测到该段可能不符合规范，建议人工复核。"]
    plain = _extract_summary_items(text.replace("\r", "\n"))
    humanized = [humanize_problem_item(x) for x in plain]
    filtered = [x for x in humanized if x]
    return filtered if filtered else ["检测到该段可能不符合规范，建议人工复核。"]


def comment_message_for_issue(issue: Issue) -> str:
    title = _comment_title(issue)
    lines = [f"{title}格式不符合要求。", "", "正确格式：", issue.expected or "请按规范要求设置。", "", "本处问题："]
    for index, item in enumerate(_issue_problem_lines(issue.current), start=1):
        lines.append(f"{index}. {item}")
    return "\n".join(lines)


def _merge_comment_message(issues: list[Issue]) -> str:
    if not issues:
        return ""
    title = _comment_title(issues[0])
    expected_items: list[str] = []
    seen_expected: set[str] = set()
    problem_items: list[str] = []
    seen_problem: set[str] = set()
    for issue in issues:
        expected = (issue.expected or "").strip()
        if expected and expected not in seen_expected:
            expected_items.append(expected)
            seen_expected.add(expected)
        for item in _issue_problem_lines(issue.current):
            cleaned = item.strip()
            if not cleaned or cleaned in seen_problem:
                continue
            problem_items.append(cleaned)
            seen_problem.add(cleaned)
    if not expected_items:
        expected_items = ["请按规范要求设置。"]
    if not problem_items:
        problem_items = ["检测到该段可能不符合规范，建议人工复核。"]
    lines = [f"{title}格式不符合要求。", "", "正确格式："]
    for idx, item in enumerate(expected_items, start=1):
        lines.append(f"{idx}. {item}")
    lines.extend(["", "本处问题："])
    for idx, item in enumerate(problem_items, start=1):
        lines.append(f"{idx}. {item}")
    return "\n".join(lines)


def issues_for_comments(issues: Iterable[Issue]) -> list[Issue]:
    selected: list[Issue] = []
    seen: set[tuple[int, str, str]] = set()
    for issue in issues:
        key = (issue.paragraph_index, issue.rule_key, issue.current or "")
        if key in seen:
            continue
        seen.add(key)
        selected.append(issue)
    return selected


def write_zip_from_dir(source_dir: Path, out_path: Path) -> None:
    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for file in source_dir.rglob("*"):
            if file.is_file():
                zf.write(file, file.relative_to(source_dir).as_posix())


def create_annotated_docx(source: Path, target: Path, issues: Iterable[Issue]) -> None:
    issues = issues_for_comments(issues)
    if not issues:
        shutil.copyfile(source, target)
        return
    if target.exists():
        target.unlink()
    with TemporaryDirectory() as tmp:
        tmpdir = Path(tmp)
        with zipfile.ZipFile(source, "r") as zf:
            zf.extractall(tmpdir)
        comments_tree = ensure_comments_parts(tmpdir)
        comments_root = comments_tree.getroot()
        existing_ids = [
            int(node.attrib.get(qname(W_NS, "id"), "0"))
            for node in comments_root.findall(qname(W_NS, "comment"))
            if node.attrib.get(qname(W_NS, "id"), "0").isdigit()
        ]
        next_id = (max(existing_ids) + 1) if existing_ids else 0

        document_xml = tmpdir / "word" / "document.xml"
        doc_tree = ET.parse(document_xml)
        doc_root = doc_tree.getroot()
        paragraphs = doc_root.findall(f".//{qname(W_NS, 'body')}/{qname(W_NS, 'p')}")
        issues_by_paragraph: dict[int, list[Issue]] = {}
        for issue in issues:
            if issue.paragraph_index < 0:
                continue
            if issue.paragraph_index >= len(paragraphs):
                continue
            issues_by_paragraph.setdefault(issue.paragraph_index, []).append(issue)

        for paragraph_index in sorted(issues_by_paragraph.keys()):
            merged_issues = issues_by_paragraph[paragraph_index]
            comment_id = next_id
            next_id += 1
            add_comment_node(comments_root, comment_id, _merge_comment_message(merged_issues))
            add_comment_marker(paragraphs[paragraph_index], comment_id)

        comments_tree.write(tmpdir / "word" / "comments.xml", encoding="utf-8", xml_declaration=True)
        doc_tree.write(document_xml, encoding="utf-8", xml_declaration=True)
        ensure_content_type(tmpdir)
        ensure_comments_relationship(tmpdir)
        write_zip_from_dir(tmpdir, target)


def try_com_save(path: Path, renderer: str) -> str:
    if renderer == "ooxml":
        return "ooxml"
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        if renderer in {"office", "wps"}:
            raise RuntimeError(f"{renderer} renderer requested but win32com is unavailable: {exc}") from exc
        return "ooxml"

    attempts = []
    if renderer in {"auto", "office"}:
        attempts.append(("office", "Word.Application"))
    if renderer in {"auto", "wps"}:
        attempts.extend([("wps", "KWPS.Application"), ("wps", "WPS.Application")])

    last_error: Exception | None = None
    for name, prog_id in attempts:
        app = None
        doc = None
        try:
            app = win32com.client.DispatchEx(prog_id)
            app.Visible = False
            doc = app.Documents.Open(str(path.resolve()))
            doc.SaveAs2(str(path.resolve()), FileFormat=16)
            doc.Close(False)
            app.Quit()
            return name
        except Exception as exc:
            last_error = exc
            try:
                if doc is not None:
                    doc.Close(False)
                if app is not None:
                    app.Quit()
            except Exception:
                pass
            if renderer in {"office", "wps"}:
                raise RuntimeError(f"{renderer} renderer failed via {prog_id}: {exc}") from exc
    if renderer == "auto":
        return "ooxml"
    raise RuntimeError(f"Renderer failed: {last_error}")


def probe_com_renderer(renderer: str) -> str:
    if renderer == "ooxml":
        return "ooxml"
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        if renderer in {"office", "wps"}:
            raise RuntimeError(f"{renderer} renderer requested but win32com is unavailable: {exc}") from exc
        return "ooxml"
    attempts = []
    if renderer in {"auto", "office"}:
        attempts.append(("office", "Word.Application"))
    if renderer in {"auto", "wps"}:
        attempts.extend([("wps", "KWPS.Application"), ("wps", "WPS.Application")])
    last_error: Exception | None = None
    for name, prog_id in attempts:
        app = None
        try:
            app = win32com.client.DispatchEx(prog_id)
            app.Visible = False
            app.Quit()
            return name
        except Exception as exc:
            last_error = exc
            try:
                if app is not None:
                    app.Quit()
            except Exception:
                pass
            if renderer in {"office", "wps"}:
                raise RuntimeError(f"{renderer} renderer failed via {prog_id}: {exc}") from exc
    if renderer == "auto":
        return "ooxml"
    raise RuntimeError(f"Renderer probe failed: {last_error}")


def _dispatch_app(renderer_used: str):
    import win32com.client  # type: ignore

    prog_ids = {
        "office": ["Word.Application"],
        "wps": ["KWPS.Application", "WPS.Application"],
    }.get(renderer_used, [])
    last_error: Exception | None = None
    for prog_id in prog_ids:
        try:
            app = win32com.client.DispatchEx(prog_id)
            app.Visible = False
            return app
        except Exception as exc:
            last_error = exc
    raise RuntimeError(f"Cannot start {renderer_used} renderer: {last_error}")


def _collect_pages_with_com(docx_path: Path, issues: list[Issue], renderer_used: str) -> dict[int, int]:
    if renderer_used not in {"office", "wps"}:
        return {}
    app = None
    doc = None
    pages: dict[int, int] = {}
    try:
        app = _dispatch_app(renderer_used)
        doc = app.Documents.Open(str(docx_path.resolve()))
        # Word constant wdActiveEndPageNumber = 3.
        for issue in issues:
            if issue.paragraph_index < 0:
                continue
            paragraph_no = issue.paragraph_index + 1
            if paragraph_no <= doc.Paragraphs.Count:
                pages[issue.paragraph_index] = int(doc.Paragraphs(paragraph_no).Range.Information(3))
        doc.Close(False)
        app.Quit()
    except Exception:
        try:
            if doc is not None:
                doc.Close(False)
            if app is not None:
                app.Quit()
        except Exception:
            pass
        return {}
    return pages


def _export_pdf_with_com(docx_path: Path, pdf_path: Path, renderer_used: str) -> None:
    app = None
    doc = None
    try:
        app = _dispatch_app(renderer_used)
        doc = app.Documents.Open(str(docx_path.resolve()))
        # Prefer ExportAsFixedFormat when available; fall back to SaveAs2 PDF.
        try:
            doc.ExportAsFixedFormat(str(pdf_path.resolve()), 17)
        except Exception:
            doc.SaveAs2(str(pdf_path.resolve()), FileFormat=17)
        doc.Close(False)
        app.Quit()
    except Exception:
        try:
            if doc is not None:
                doc.Close(False)
            if app is not None:
                app.Quit()
        except Exception:
            pass
        raise


def _render_pdf_pages(pdf_path: Path, pages: set[int], image_dir: Path, prefix: str) -> dict[int, Path]:
    import pymupdf  # type: ignore

    image_dir.mkdir(parents=True, exist_ok=True)
    rendered: dict[int, Path] = {}
    with pymupdf.open(pdf_path) as doc:
        for page_no in sorted(pages):
            if page_no < 1 or page_no > len(doc):
                continue
            page = doc[page_no - 1]
            pix = page.get_pixmap(matrix=pymupdf.Matrix(1.5, 1.5), alpha=False)
            out = image_dir / f"{prefix}_page_{page_no}.png"
            pix.save(out)
            rendered[page_no] = out
    return rendered


def _find_codex_render_docx() -> Path | None:
    candidates = sorted(
        (Path.home() / ".codex" / "plugins" / "cache" / "openai-primary-runtime" / "documents").glob(
            "*/skills/documents/render_docx.py"
        ),
        reverse=True,
    )
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def _render_docx_pages_with_artifact_tool(docx_path: Path, out_dir: Path) -> list[Path]:
    render_script = _find_codex_render_docx()
    if render_script is None:
        raise RuntimeError("documents render_docx.py was not found.")
    out_dir.mkdir(parents=True, exist_ok=True)
    command = [
        sys.executable,
        str(render_script),
        str(docx_path.resolve()),
        "--output_dir",
        str(out_dir.resolve()),
        "--renderer",
        "artifact-tool",
    ]
    proc = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="replace")
    if proc.returncode != 0:
        raise RuntimeError(
            "artifact-tool DOCX render failed. "
            f"stdout={proc.stdout.strip()} stderr={proc.stderr.strip()}"
        )
    pages = sorted(
        out_dir.glob("page-*.png"),
        key=lambda p: int(re.search(r"page-(\d+)\.png$", p.name).group(1))
        if re.search(r"page-(\d+)\.png$", p.name)
        else 0,
    )
    if not pages:
        raise RuntimeError(f"artifact-tool did not produce page PNGs in {out_dir}")
    return pages


def _png_dimensions(path: Path) -> tuple[int | None, int | None]:
    try:
        with path.open("rb") as fh:
            header = fh.read(24)
        if header[:8] != b"\x89PNG\r\n\x1a\n":
            return None, None
        return int.from_bytes(header[16:20], "big"), int.from_bytes(header[20:24], "big")
    except Exception:
        return None, None


def _png_visual_stats(path: Path | None) -> dict[str, float | int | None]:
    if path is None or not path.exists():
        return {"ink_ratio": None, "dark_ratio": None, "width": None, "height": None}
    try:
        from PIL import Image
    except Exception:
        return {"ink_ratio": None, "dark_ratio": None, "width": None, "height": None}
    try:
        with Image.open(path) as image:
            gray = image.convert("L")
            width, height = gray.size
            histogram = gray.resize((max(1, width // 4), max(1, height // 4))).histogram()
            total = sum(histogram)
            if total == 0:
                return {"ink_ratio": None, "dark_ratio": None, "width": width, "height": height}
            ink = sum(histogram[:245])
            dark = sum(histogram[:80])
            return {
                "ink_ratio": round(ink / total, 4),
                "dark_ratio": round(dark / total, 4),
                "width": width,
                "height": height,
            }
    except Exception:
        return {"ink_ratio": None, "dark_ratio": None, "width": None, "height": None}


def _qa_page_records(before_pages: list[Path], after_pages: list[Path]) -> list[dict[str, object]]:
    records: list[dict[str, object]] = []
    total = max(len(before_pages), len(after_pages))
    for idx in range(total):
        before = before_pages[idx] if idx < len(before_pages) else None
        after = after_pages[idx] if idx < len(after_pages) else None
        before_size = before.stat().st_size if before and before.exists() else 0
        after_size = after.stat().st_size if after and after.exists() else 0
        bw, bh = _png_dimensions(before) if before else (None, None)
        aw, ah = _png_dimensions(after) if after else (None, None)
        before_stats = _png_visual_stats(before)
        after_stats = _png_visual_stats(after)
        findings: list[str] = []
        if before is None:
            findings.append("missing before-page image")
        if after is None:
            findings.append("missing after-page image")
        if before_size and before_size < 1500:
            findings.append("before-page image is suspiciously small")
        if after_size and after_size < 1500:
            findings.append("after-page image is suspiciously small")
        if bw and aw and (bw, bh) != (aw, ah):
            findings.append("before/after rendered page dimensions differ")
        before_ink = before_stats.get("ink_ratio")
        after_ink = after_stats.get("ink_ratio")
        if isinstance(after_ink, float) and after_ink < 0.005:
            findings.append("after-page image appears blank or nearly blank")
        if isinstance(before_ink, float) and isinstance(after_ink, float) and abs(after_ink - before_ink) > 0.18:
            findings.append("before/after page ink density changed substantially")
        after_dark = after_stats.get("dark_ratio")
        if isinstance(after_dark, float) and after_dark > 0.25:
            findings.append("after-page has unusually large dark regions; inspect for overlap or rendering blocks")
        records.append(
            {
                "page": idx + 1,
                "before": str(before.resolve()) if before else None,
                "after": str(after.resolve()) if after else None,
                "before_size_bytes": before_size,
                "after_size_bytes": after_size,
                "before_dimensions": [bw, bh],
                "after_dimensions": [aw, ah],
                "before_visual_stats": before_stats,
                "after_visual_stats": after_stats,
                "findings": findings,
            }
        )
    return records


def render_full_document_qa(
    input_path: Path,
    fixed_path: Path,
    output_dir: Path,
    mode: str,
) -> tuple[str, dict[str, object] | None]:
    try:
        qa_dir = output_dir / "render_qa"
        before_dir = qa_dir / "before"
        after_dir = qa_dir / "after"
        before_pages = _render_docx_pages_with_artifact_tool(input_path, before_dir)
        after_pages = _render_docx_pages_with_artifact_tool(fixed_path, after_dir)
        records = _qa_page_records(before_pages, after_pages)
        finding_count = sum(len(record["findings"]) for record in records)
        status = (
            "Full-document render QA generated with artifact-tool. "
            f"Rendered {len(before_pages)} before pages and {len(after_pages)} after pages; "
            f"basic image checks found {finding_count} warning(s). "
            "When the fixed DOCX is produced through OOXML fallback, these page images are QA previews only and may not match Word/WPS rendering exactly."
        )
        return status, {
            "renderer": "artifact-tool",
            "before_page_count": len(before_pages),
            "after_page_count": len(after_pages),
            "warning_count": finding_count,
            "pages": records,
            "note": (
                "Use these page images to inspect layout drift, overlap, clipping, and header/footer issues when Office/WPS COM is unavailable. "
                "If the fixed DOCX was produced through OOXML fallback, preview images may be inaccurate because Word/WPS pagination, fields, fonts, and compatibility layout are not refreshed by COM."
            ),
        }
    except Exception as exc:
        status = f"Full-document render QA unavailable: {exc}"
        if mode == "require":
            raise RuntimeError(status) from exc
        return status, None


def attach_screenshots(
    input_path: Path,
    fixed_path: Path,
    issues: list[Issue],
    output_dir: Path,
    renderer_used: str,
    mode: str,
) -> tuple[str, dict[str, object] | None]:
    if mode == "never":
        note = "Screenshots disabled by --screenshots never."
        for issue in issues:
            issue.screenshot_note = note
        return note, None
    if renderer_used not in {"office", "wps"}:
        note, qa = render_full_document_qa(input_path, fixed_path, output_dir, mode)
        issue_note = (
            "Paragraph-level screenshots require Microsoft Word or WPS COM rendering. "
            "Full-document before/after page render QA was generated instead."
            if qa
            else note
        )
        for issue in issues:
            issue.screenshot_note = issue_note
        return note, qa
    try:
        pages = _collect_pages_with_com(input_path, issues, renderer_used)
        for issue in issues:
            issue.page = pages.get(issue.paragraph_index)
        needed_pages = {page for page in pages.values() if page}
        if not needed_pages:
            note = "Screenshots skipped because paragraph page numbers could not be resolved."
            if mode == "require":
                raise RuntimeError(note)
            for issue in issues:
                issue.screenshot_note = note
            return note, None

        screenshot_dir = output_dir / "screenshots"
        before_pdf = output_dir / "_before_render.pdf"
        after_pdf = output_dir / "_after_render.pdf"
        _export_pdf_with_com(input_path, before_pdf, renderer_used)
        _export_pdf_with_com(fixed_path, after_pdf, renderer_used)
        before_images = _render_pdf_pages(before_pdf, needed_pages, screenshot_dir, "before")
        after_images = _render_pdf_pages(after_pdf, needed_pages, screenshot_dir, "after")
        before_pdf.unlink(missing_ok=True)
        after_pdf.unlink(missing_ok=True)
        for issue in issues:
            if issue.page in before_images:
                issue.before_screenshot = str(before_images[issue.page].resolve())
            if issue.page in after_images:
                issue.after_screenshot = str(after_images[issue.page].resolve())
        status, qa = render_full_document_qa(input_path, fixed_path, output_dir, "auto")
        return "Paragraph-level screenshots generated from Office/WPS-rendered DOCX pages. " + status, qa
    except Exception as exc:
        note = f"Screenshots unavailable: {exc}"
        if mode == "require":
            raise RuntimeError(note) from exc
        for issue in issues:
            issue.screenshot_note = note
        return note, None


def attach_after_formats(fixed_path: Path, issues: list[Issue]) -> None:
    fixed_doc = Document(str(fixed_path))
    for issue in issues:
        if 0 <= issue.paragraph_index < len(fixed_doc.paragraphs):
            rule_key = issue.rule_key.removesuffix("_runs")
            issue.after = current_format(fixed_doc.paragraphs[issue.paragraph_index], RULES.get(rule_key))
        elif issue.paragraph_index >= 0:
            issue.after = "结构修改后段落位置改变，无法可靠定位修改后状态。"


def issue_dict(issue: Issue) -> dict[str, object]:
    raw_after = issue.after
    after_text = "已按规则修复格式。"
    if issue.rule_key in {"reference_272_type_format_mismatch", "reference_entry_format"}:
        after_text = "著录内容不自动改写，请人工按模板修改。"
    elif issue.rule_key in {"toc_page_number_manual_review", "toc_body_section_manual_review", "toc_start_new_page_manual_review"}:
        after_text = "该问题需要人工复核，工具不自动修改。"
    elif issue.after and "无法可靠定位" in issue.after:
        after_text = "结构调整后位置变化，建议人工复核。"
    problem_text = "\n".join(f"{i}. {line}" for i, line in enumerate(_issue_problem_lines(issue.current), start=1))
    return {
        "paragraph_index": issue.paragraph_index,
        "rule_key": issue.rule_key,
        "category": issue.category,
        "text_type": issue.text_type,
        "text_excerpt": issue.text_excerpt,
        "location": issue.location,
        "page": issue.page,
        "current": humanize_issue_current(issue.current),
        "expected": issue.expected,
        "after": after_text,
        "message": f"本处问题：\n{problem_text}",
        "raw_current": issue.current,
        "raw_after": raw_after,
        "before_screenshot": issue.before_screenshot,
        "after_screenshot": issue.after_screenshot,
        "screenshot_note": issue.screenshot_note,
    }


def summarize_issues(issues: list[Issue]) -> dict[str, int]:
    summary: dict[str, int] = {}
    for issue in issues:
        summary[issue.category] = summary.get(issue.category, 0) + 1
    return dict(sorted(summary.items()))


def write_html_report(report: dict[str, object], path: Path) -> None:
    issues = report.get("issues", [])
    assert isinstance(issues, list)
    category_summary = report.get("issue_summary_by_category") or {}
    category_html = ""
    if isinstance(category_summary, dict):
        rows_summary = "".join(
            f"<tr><td>{html.escape(str(category))}</td><td>{html.escape(str(count))}</td></tr>"
            for category, count in category_summary.items()
        )
        category_html = (
            "<h2>问题分类汇总</h2>"
            "<table class='summary-table'><thead><tr><th>类别</th><th>数量</th></tr></thead>"
            f"<tbody>{rows_summary}</tbody></table>"
        )
    ref272_summary = report.get("reference_272_check_summary") or {}
    ref272_html = ""
    if isinstance(ref272_summary, dict):
        bad_rows = []
        for entry in ref272_summary.get("entries", []):
            if not isinstance(entry, dict):
                continue
            if entry.get("status") == "matched":
                continue
            missing = entry.get("missing_fields") or []
            missing_text = "、".join(str(x) for x in missing) if isinstance(missing, list) else str(missing)
            template = str(entry.get("matched_template") or "请按对应文献类型模板补全字段")
            bad_rows.append(
                "<tr>"
                f"<td>{html.escape(str(entry.get('visual_index', entry.get('index', ''))))}</td>"
                f"<td>{html.escape(str(entry.get('detected_type', '')))}</td>"
                f"<td>{html.escape(str(entry.get('text_excerpt', '')))}</td>"
                f"<td>{html.escape(missing_text)}</td>"
                f"<td>{html.escape(template)}</td>"
                "</tr>"
            )
        ref272_html = (
            "<h2>参考文献著录格式检查汇总（2.7.2）</h2>"
            "<table class='summary-table'><tbody>"
            f"<tr><td>total_entries</td><td>{html.escape(str(ref272_summary.get('total_entries', 0)))}</td></tr>"
            f"<tr><td>matched_entries</td><td>{html.escape(str(ref272_summary.get('matched_entries', 0)))}</td></tr>"
            f"<tr><td>type_mismatch_entries</td><td>{html.escape(str(ref272_summary.get('type_mismatch_entries', 0)))}</td></tr>"
            f"<tr><td>unmatched_entries</td><td>{html.escape(str(ref272_summary.get('unmatched_entries', 0)))}</td></tr>"
            "</tbody></table>"
        )
        if bad_rows:
            ref272_html += (
                "<table class='summary-table'><thead><tr><th>序号</th><th>类型</th><th>原文摘要</th><th>缺失字段</th><th>建议模板</th></tr></thead>"
                f"<tbody>{''.join(bad_rows)}</tbody></table>"
            )
    ref273_summary = report.get("reference_273_check_summary") or {}
    ref273_html = ""
    if isinstance(ref273_summary, dict):
        ref273_html = (
            "<h2>2.7.3 补充规则检查</h2>"
            "<table class='summary-table'><tbody>"
            f"<tr><td>online_entries</td><td>{html.escape(str(ref273_summary.get('online_entries', 0)))}</td></tr>"
            f"<tr><td>author_et_al_review_entries</td><td>{html.escape(str(ref273_summary.get('author_et_al_review_entries', 0)))}</td></tr>"
            f"<tr><td>citation_date_missing_entries</td><td>{html.escape(str(ref273_summary.get('citation_date_missing_entries', 0)))}</td></tr>"
            f"<tr><td>access_path_missing_entries</td><td>{html.escape(str(ref273_summary.get('access_path_missing_entries', 0)))}</td></tr>"
            f"<tr><td>doi_or_url_missing_entries</td><td>{html.escape(str(ref273_summary.get('doi_or_url_missing_entries', 0)))}</td></tr>"
            "</tbody></table>"
        )
    render_qa = report.get("render_qa")
    qa_html = ""
    if isinstance(render_qa, dict):
        qa_rows = []
        for page_obj in render_qa.get("pages", []):
            page = page_obj if isinstance(page_obj, dict) else {}
            before = page.get("before")
            after = page.get("after")
            warnings = page.get("findings") or []
            warning_text = ", ".join(str(item) for item in warnings) if isinstance(warnings, list) else str(warnings)
            before_html = ""
            after_html = ""
            if before:
                before_html = f'<figure><figcaption>修改前 第 {html.escape(str(page.get("page")))} 页</figcaption><img src="{html.escape(Path(str(before)).resolve().as_uri())}"></figure>'
            if after:
                after_html = f'<figure><figcaption>修改后 第 {html.escape(str(page.get("page")))} 页</figcaption><img src="{html.escape(Path(str(after)).resolve().as_uri())}"></figure>'
            qa_rows.append(
                "<section class='qa-page'>"
                f"<h3>页面 {html.escape(str(page.get('page')))}</h3>"
                f"<p><strong>本处问题：</strong>{html.escape(warning_text or '未发现图片尺寸/文件异常')}</p>"
                f"<div class='shots'>{before_html}{after_html}</div>"
                "</section>"
            )
        qa_html = (
            "<h1>渲染截图 QA</h1>"
            "<div class='summary'>"
            f"<p><strong>渲染器：</strong>{html.escape(str(render_qa.get('renderer', '')))}</p>"
            f"<p><strong>修改前页数：</strong>{html.escape(str(render_qa.get('before_page_count', '')))}</p>"
            f"<p><strong>修改后页数：</strong>{html.escape(str(render_qa.get('after_page_count', '')))}</p>"
            f"<p><strong>基础警告数：</strong>{html.escape(str(render_qa.get('warning_count', '')))}</p>"
            f"<p>{html.escape(str(render_qa.get('note', '')))}</p>"
            "</div>"
            + "".join(qa_rows)
        )
    structure_warnings = report.get("structure_warnings") or []
    if structure_warnings:
        warning_items = "".join(f"<li>{html.escape(str(item))}</li>" for item in structure_warnings)
        structure_html = f"<section class='structure'><h2>组成部分提醒</h2><ul>{warning_items}</ul></section>"
    else:
        structure_html = ""
    layout_fix_policy = report.get("layout_fix_policy") or {}
    toc_manual_review_html = ""
    if isinstance(layout_fix_policy, dict) and layout_fix_policy.get("toc_page_numbering") == "report_only":
        toc_manual_review_html = (
            "<section class='structure'>"
            "<h2>目录页码人工复核说明</h2>"
            "<ol>"
            "<li>目录是 Word/WPS 自动目录或 TOC 域时，页码格式不应由工具自动改。</li>"
            "<li>目录条目右侧页码应保持阿拉伯数字。</li>"
            "<li>目录页底部页码属于前置部分，应为小写罗马数字。</li>"
            "<li>请人工检查目录所在节的页码格式。</li>"
            "<li>请人工检查正文第一章所在节是否从阿拉伯数字 1 开始。</li>"
            "<li>工具不会自动插入/移动分节符，以避免再次生成空白页。</li>"
            "</ol>"
            "</section>"
        )
    toc_new_page_review_html = ""
    if isinstance(layout_fix_policy, dict) and layout_fix_policy.get("toc_start_new_page") == "report_only":
        toc_new_page_review_html = (
            "<section class='structure'>"
            "<h2>目录另起页人工复核说明</h2>"
            "<ol>"
            "<li>目录应另起一页。</li>"
            "<li>如果目录与英文摘要 Key words 在同一页，请在 Word/WPS 中将目录标题“目 录”设置为段前分页，或在其前插入分页符。</li>"
            "<li>如果目录是自动目录/TOC 域，不建议工具自动设置 page_break_before，因为可能作用到隐藏 TOC 占位段并生成空白页。</li>"
            "<li>工具不会自动插入分页符、分节符或修改 TOC 域。</li>"
            "</ol>"
            "</section>"
        )
    llm_html = ""
    rows = []
    for n, issue_obj in enumerate(issues, start=1):
        issue = issue_obj if isinstance(issue_obj, dict) else {}
        before_img = issue.get("before_screenshot")
        after_img = issue.get("after_screenshot")
        screenshots = html.escape(str(issue.get("screenshot_note") or ""))
        if before_img and after_img:
            before_uri = Path(str(before_img)).resolve().as_uri()
            after_uri = Path(str(after_img)).resolve().as_uri()
            screenshots = (
                f'<div class="shots"><figure><figcaption>修改前</figcaption>'
                f'<img src="{html.escape(before_uri)}"></figure>'
                f'<figure><figcaption>修改后</figcaption>'
                f'<img src="{html.escape(after_uri)}"></figure></div>'
            )
        problem_text = str(issue.get("message") or "")
        if problem_text.startswith("本处问题："):
            problem_text = problem_text[len("本处问题：") :].lstrip("\r\n ")
        rows.append(
            "<section class='issue'>"
            f"<h2>#{n} {html.escape(str(issue.get('text_type', 'unknown')))}</h2>"
            f"<p><strong>位置：</strong>段落 {html.escape(str(issue.get('paragraph_index')))}"
            f"{'，页 ' + html.escape(str(issue.get('page'))) if issue.get('page') else ''}</p>"
            f"<p><strong>文本：</strong>{html.escape(str(issue.get('text_excerpt') or ''))}</p>"
            f"<p><strong>修改前：</strong>{html.escape(str(issue.get('current') or ''))}</p>"
            f"<p><strong>正确格式：</strong>{html.escape(str(issue.get('expected') or ''))}</p>"
            f"<p><strong>本处问题：</strong>{html.escape(problem_text)}</p>"
            f"<p><strong>修改后：</strong>{html.escape(str(issue.get('after') or ''))}</p>"
            f"{screenshots}"
            "</section>"
        )
    doc = f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>thesis format skill report</title>
<style>
body {{ font-family: "Microsoft YaHei", "Segoe UI", sans-serif; margin: 32px; line-height: 1.6; color: #1f2937; }}
h1 {{ margin-bottom: 8px; }}
.summary {{ padding: 16px; background: #f3f4f6; border-left: 4px solid #2563eb; margin: 16px 0 24px; }}
.summary-table {{ border-collapse: collapse; margin: 12px 0 24px; }}
.summary-table th, .summary-table td {{ border: 1px solid #d1d5db; padding: 6px 10px; text-align: left; }}
.issue {{ border: 1px solid #d1d5db; border-radius: 6px; padding: 16px; margin: 18px 0; }}
.structure {{ border: 1px solid #f59e0b; background: #fffbeb; border-radius: 6px; padding: 16px; margin: 18px 0; }}
.qa-page {{ border: 1px solid #bfdbfe; border-radius: 6px; padding: 16px; margin: 18px 0; }}
.issue h2 {{ margin-top: 0; font-size: 18px; }}
.shots {{ display: grid; grid-template-columns: repeat(2, minmax(0, 1fr)); gap: 16px; }}
figure {{ margin: 0; }}
img {{ max-width: 100%; border: 1px solid #d1d5db; }}
code {{ background: #f3f4f6; padding: 1px 4px; }}
</style>
</head>
<body>
<h1>学士学位论文 DOCX 格式检查报告</h1>
<div class="summary">
<p><strong>输入文件：</strong>{html.escape(str(report.get("input", "")))}</p>
<p><strong>批注版：</strong>{html.escape(str(report.get("annotated_docx", "")))}</p>
<p><strong>修正版：</strong>{html.escape(str(report.get("fixed_docx", "")))}</p>
<p><strong>修改/问题总数：</strong>{html.escape(str(report.get("issue_count", 0)))}</p>
<p><strong>渲染方式：</strong>fixed={html.escape(str(report.get("renderer_for_fixed", "")))}, comments={html.escape(str(report.get("renderer_for_comments", "")))}</p>
<p><strong>截图状态：</strong>{html.escape(str(report.get("screenshot_status", "")))}</p>
</div>
{category_html}
{ref272_html}
{ref273_html}
{structure_html}
{toc_manual_review_html}
{toc_new_page_review_html}
{llm_html}
{qa_html}
{''.join(rows)}
</body>
</html>
"""
    path.write_text(doc, encoding="utf-8")


def write_llm_review_separate_html(llm_review: dict[str, object], path: Path) -> None:
    def _normalize_llm_issue_row(item: object, index: int) -> dict[str, str]:
        if isinstance(item, dict):
            return {
                "id": str(item.get("id", index)),
                "page": str(item.get("page", "unknown")),
                "type": str(item.get("type", item.get("category", "other"))),
                "severity": str(item.get("severity", "")),
                "evidence": str(item.get("evidence", item.get("text", item.get("message", "")))),
                "spec_basis": str(item.get("spec_basis", item.get("basis", ""))),
                "suggestion": str(item.get("suggestion", item.get("advice", ""))),
            }
        return {
            "id": str(index),
            "page": "unknown",
            "type": "manual_review",
            "severity": "",
            "evidence": str(item),
            "spec_basis": "",
            "suggestion": "请人工复核该项。",
        }

    def _normalize_llm_manual_row(item: object, index: int) -> dict[str, str]:
        if isinstance(item, dict):
            return {
                "page": str(item.get("page", "unknown")),
                "reason": str(item.get("reason", item.get("text", item.get("message", "")))),
                "suggestion": str(item.get("suggestion", item.get("advice", ""))),
            }
        return {
            "page": "unknown",
            "reason": str(item),
            "suggestion": "请人工复核该项。",
        }

    review = llm_review.get("review") if isinstance(llm_review, dict) else {}
    review = review if isinstance(review, dict) else {}
    basis = review.get("basis") if isinstance(review.get("basis"), dict) else {}
    read_check = review.get("read_check") if isinstance(review.get("read_check"), dict) else {}
    issues = review.get("issues") if isinstance(review.get("issues"), list) else []
    manual_items = review.get("manual_review_items") if isinstance(review.get("manual_review_items"), list) else []
    uncertain_items = review.get("uncertain_items") if isinstance(review.get("uncertain_items"), list) else []
    safe_edit_plan = review.get("safe_edit_plan") if isinstance(review.get("safe_edit_plan"), list) else []
    rejected_claims = review.get("rejected_claims") if isinstance(review.get("rejected_claims"), list) else []
    validation_warnings = review.get("validation_warnings") if isinstance(review.get("validation_warnings"), list) else []
    issues_normalized = [_normalize_llm_issue_row(item, idx) for idx, item in enumerate(issues, start=1)]
    issues_sorted = sorted(issues_normalized, key=lambda x: x.get("page", "unknown"))
    issue_rows = "".join(
        "<tr>"
        f"<td>{html.escape(it.get('id', ''))}</td>"
        f"<td>{html.escape(it.get('page', 'unknown'))}</td>"
        f"<td>{html.escape(it.get('type', 'other'))}</td>"
        f"<td>{html.escape(it.get('severity', ''))}</td>"
        f"<td>{html.escape(it.get('evidence', ''))}</td>"
        f"<td>{html.escape(it.get('spec_basis', ''))}</td>"
        f"<td>{html.escape(it.get('suggestion', ''))}</td>"
        "</tr>"
        for it in issues_sorted
    )
    manual_normalized = [_normalize_llm_manual_row(item, idx) for idx, item in enumerate(manual_items, start=1)]
    manual_rows = "".join(
        "<tr>"
        f"<td>{html.escape(it.get('page', 'unknown'))}</td>"
        f"<td>{html.escape(it.get('reason', ''))}</td>"
        f"<td>{html.escape(it.get('suggestion', ''))}</td>"
        "</tr>"
        for it in manual_normalized
    )
    status = str(basis.get("document_upload_status", llm_review.get("document_upload_status", "failed")))
    fallback_used = bool(basis.get("fallback_used", llm_review.get("fallback_used", False)))
    fallback_reason = str(basis.get("fallback_reason", llm_review.get("fallback_reason", "")))
    review_mode = str(basis.get("review_mode", llm_review.get("mode", "")))
    doc_model_attempted = str(
        basis.get("doc_model_attempted", llm_review.get("doc_model_attempted", llm_review.get("model", "")))
    )
    fallback_model = str(basis.get("fallback_model", llm_review.get("fallback_model", "")))
    actual_review_model = str(
        llm_review.get("actual_review_model", fallback_model if fallback_used and fallback_model else llm_review.get("model", ""))
    )
    fallback_note = ""
    if review_mode == "qwen_long_fileid_docx" and status == "ok":
        fallback_note = "<p><strong>当前报告由 qwen-long 读取规范文件 file-id 和论文 file-id 后生成。</strong></p>"
    elif status != "ok" or fallback_used:
        fallback_note = "<p><strong>本次未能通过 API 直接提交 Word 文档，已回退为文本/结构化摘要审查。当前报告不是 LLM 独立读取 Word 后生成，而是 fallback_summary。</strong></p>"
    fallback_llm_note = ""
    if fallback_used:
        fallback_llm_note = (
            "<p><strong>说明：</strong>qwen-long file-id 双文档审查失败，系统已回退为文本/结构化摘要审查；"
            "以下问题列表仍由 LLM 根据抽取出的规范文本、fixed 文档结构和确定性检查摘要生成。</p>"
        )
    uncertain_rows = "".join(
        "<tr>"
        f"<td>{html.escape(str(idx))}</td>"
        f"<td>{html.escape(str(item if not isinstance(item, dict) else item.get('reason', item.get('text', item.get('message', '')))))}</td>"
        "</tr>"
        for idx, item in enumerate(uncertain_items, start=1)
    )
    plan_rows = "".join(
        "<tr>"
        f"<td>{html.escape(str((item.get('operation') if isinstance(item, dict) else '')))}</td>"
        f"<td>{html.escape(str((item.get('target_hint') if isinstance(item, dict) else '')))}</td>"
        f"<td>{html.escape(str((item.get('risk', 'unknown') if isinstance(item, dict) else 'unknown')))}</td>"
        f"<td>{html.escape(str((item.get('evidence_exact_quote', '') if isinstance(item, dict) else '')))}</td>"
        f"<td>{html.escape(json.dumps(item.get('parameters', {}), ensure_ascii=False) if isinstance(item, dict) else '{}')}</td>"
        "</tr>"
        for item in safe_edit_plan
    )
    rejected_rows = "".join(
        "<li>" + html.escape(str(item if not isinstance(item, dict) else item.get("reason", item.get("text", item)))) + "</li>"
        for item in rejected_claims
    )
    warning_rows = "".join(
        "<li>" + html.escape(str(item if not isinstance(item, dict) else item.get("warning", item.get("reason", item)))) + "</li>"
        for item in validation_warnings
    )
    raw_json = json.dumps(llm_review, ensure_ascii=False, indent=2)
    doc = f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>LLM 格式审查与修改建议报告</title>
<style>
body {{ font-family: "Microsoft YaHei", sans-serif; margin: 24px; line-height: 1.6; }}
table {{ border-collapse: collapse; width: 100%; }}
th, td {{ border: 1px solid #d1d5db; padding: 6px; vertical-align: top; text-align: left; }}
.warn {{ color: #b91c1c; font-weight: 700; }}
</style>
</head>
<body>
<h1>LLM 格式审查与修改建议报告</h1>
<p>该报告由千问根据格式指导书和目标论文生成，仅供人工复核；safe_edit_plan 只作为人工修改建议，不会自动执行。</p>
<p><strong>主链路尝试模型：</strong>{html.escape(str(llm_review.get("provider", "")))} / {html.escape(doc_model_attempted)}</p>
<p><strong>实际生成审查内容模型：</strong>{html.escape(str(llm_review.get("provider", "")))} / {html.escape(actual_review_model)}</p>
<p><strong>review_mode：</strong>{html.escape(str(basis.get("review_mode", llm_review.get("mode", ""))))}</p>
<p><strong>target_docx_source：</strong>{html.escape(str(basis.get("target_docx_source", "fixed")))}</p>
<p><strong>是否使用规范文件：</strong>{html.escape(str(basis.get("uses_spec_docx", llm_review.get("uses_spec_docx", False))))}</p>
<p><strong>是否使用目标论文：</strong>{html.escape(str(basis.get("uses_target_docx", basis.get("uses_fixed_docx", llm_review.get("uses_fixed_docx", False)))))}</p>
<p><strong>文档提交方式：</strong>{html.escape(status if status else "failed")}</p>
<p><strong>是否回退文本审查：</strong>{html.escape(str(fallback_used))}</p>
<p><strong>回退原因：</strong>{html.escape(fallback_reason)}</p>
<p><strong>overall_result：</strong>{html.escape(str(review.get("overall_result", "")))}</p>
<p><strong>overall_risk：</strong>{html.escape(str(review.get("overall_risk", "")))}</p>
<p><strong>issues 数量：</strong>{len(issues_sorted)}</p>
<p><strong>manual_review_items 数量：</strong>{len(manual_normalized)}</p>
<p><strong>validation_warnings 数量：</strong>{len(validation_warnings)}</p>
{fallback_note}
{fallback_llm_note}
<h2>读取验证 read_check</h2>
<table><tbody>
<tr><th>spec_file_read</th><td>{html.escape(str(read_check.get("spec_file_read", basis.get("spec_file_read", False))))}</td></tr>
<tr><th>target_file_read</th><td>{html.escape(str(read_check.get("target_file_read", basis.get("target_file_read", False))))}</td></tr>
<tr><th>title</th><td>{html.escape(str(read_check.get("title", "")))}</td></tr>
<tr><th>keywords</th><td>{html.escape(str(read_check.get("keywords", "")))}</td></tr>
<tr><th>toc_evidence</th><td>{html.escape(str(read_check.get("toc_evidence", "")))}</td></tr>
<tr><th>reference_numbering_observation</th><td>{html.escape(str(read_check.get("reference_numbering_observation", "")))}</td></tr>
</tbody></table>
<h2>问题列表</h2>
<table><thead><tr><th>ID</th><th>页码</th><th>类型</th><th>严重度</th><th>证据</th><th>规范依据</th><th>建议</th></tr></thead><tbody>{issue_rows}</tbody></table>
<h2>manual_review_items</h2>
<table><thead><tr><th>页码</th><th>原因</th><th>建议</th></tr></thead><tbody>{manual_rows}</tbody></table>
<h2>uncertain_items</h2>
<table><thead><tr><th>#</th><th>内容</th></tr></thead><tbody>{uncertain_rows}</tbody></table>
<h2>safe_edit_plan（仅建议，不自动执行）</h2>
<table><thead><tr><th>operation</th><th>target_hint</th><th>risk</th><th>evidence_exact_quote</th><th>parameters</th></tr></thead><tbody>{plan_rows}</tbody></table>
<h2>rejected_claims</h2>
<ul>{rejected_rows}</ul>
<h2 class="warn">validation_warnings</h2>
<ul class="warn">{warning_rows}</ul>
<details><summary>完整 JSON</summary><pre>{html.escape(raw_json)}</pre></details>
</body>
</html>"""
    path.write_text(doc, encoding="utf-8")


def stdout_json(report: dict[str, object]) -> str:
    return json.dumps(report, ensure_ascii=True, indent=2)


def build_run_timestamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def append_timestamp_to_name(path: Path, run_timestamp: str) -> Path:
    return path.with_name(f"{path.stem}_{run_timestamp}{path.suffix}")


def _has_toc_title(document: Document) -> bool:
    return any(re.sub(r"\s+", "", paragraph_text(p)) in {"目录", "目次"} for p in document.paragraphs)


def detect_toc_field_report_only(document: Document) -> bool:
    body = document.element.body
    for node in body.iter():
        if node.tag == qn("w:instrText") and "TOC" in (node.text or "").upper():
            return True
        if node.tag == qn("w:fldSimple") and "TOC" in (node.get(qn("w:instr")) or "").upper():
            return True
    return False


def detect_toc_manual_review_issues(document: Document) -> list[Issue]:
    issues: list[Issue] = []
    texts = [paragraph_text(paragraph) for paragraph in document.paragraphs]
    regions = analyze_section_sequence(texts, has_toc_field=document_has_toc_field(document))["regions"]
    has_toc = _has_toc_title(document) or detect_toc_field_report_only(document)
    if not has_toc:
        return issues

    issues.append(
        Issue(
            paragraph_index=-1,
            rule_key="toc_page_number_manual_review",
            text_type="目录页码人工复核",
            text_excerpt="目录/TOC",
            current="检测到目录/TOC 自动目录。由于目录页页码涉及 section、页脚 PAGE 域和 TOC 字段刷新，工具不自动修改。",
            expected="目录所在前置页页脚页码应为小写罗马数字 i, ii, iii...；正文从第 1 页开始使用阿拉伯数字或“第*页 共*页”。",
            message=(
                "请人工在 Word/WPS 中复核目录页页脚页码。若目录页底部显示 3、4 等阿拉伯数字，应将目录所在节页码格式设置为小写罗马数字；"
                "正文第一章所在节应设置为阿拉伯数字并从 1 开始。不要修改目录条目右侧的正文页码。"
            ),
            category="page",
            location="document TOC",
        )
    )

    first_toc_idx = next((idx for idx, region in enumerate(regions) if region == "toc"), None)
    first_body_idx = next((idx for idx, region in enumerate(regions) if region == "body"), None)
    if first_toc_idx is not None and first_body_idx is not None and first_body_idx > first_toc_idx:
        gap_texts = [texts[idx].strip() for idx in range(first_toc_idx, first_body_idx) if texts[idx].strip()]
        has_explicit_section_mark = any("分节符" in text for text in gap_texts)
        if not has_explicit_section_mark:
            issues.append(
                Issue(
                    paragraph_index=first_body_idx,
                    rule_key="toc_body_section_manual_review",
                    text_type="目录与正文分节人工复核",
                    text_excerpt=(texts[first_body_idx] or "正文起始段落")[:160],
                    current="工具检测到目录/TOC 与正文边界可能需要人工确认。",
                    expected="目录结束后、正文第一章前应有正确分节边界；前置部分使用小写罗马数字，正文部分从阿拉伯数字 1 开始。",
                    message="请在 Word/WPS 中显示编辑标记，检查目录结束后、正文第一章前是否有“下一页分节符”。本工具不自动插入或移动分节符，以避免产生空白页。",
                    category="page",
                    location="toc to body boundary",
                )
            )
    return issues


def _truncate_text(value: object, limit: int = 240) -> str:
    text = str(value or "")
    return text if len(text) <= limit else text[: limit - 3] + "..."


def extract_spec_docx_text(spec_path: Path) -> str:
    if not spec_path.exists():
        return ""
    doc = Document(str(spec_path))
    blocks: list[str] = []
    for p in doc.paragraphs:
        txt = re.sub(r"\s+", " ", paragraph_text(p)).strip()
        if txt:
            blocks.append(txt)
    for t_idx, table in enumerate(doc.tables, start=1):
        blocks.append(f"[TABLE {t_idx}]")
        for r_idx, row in enumerate(table.rows, start=1):
            cells = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells]
            blocks.append(f"row {r_idx}:\t" + "\t".join(cells))
    text = "\n".join(blocks)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def extract_fixed_docx_structure_for_llm(fixed_docx: Path) -> dict[str, object]:
    if not fixed_docx.exists():
        return {
            "paragraphs": [],
            "headings": [],
            "figure_captions": [],
            "table_captions": [],
            "references": [],
            "table_count": 0,
            "image_count": 0,
        }
    doc = Document(str(fixed_docx))
    paragraphs: list[dict[str, object]] = []
    heading_candidates: list[dict[str, object]] = []
    for i, p in enumerate(doc.paragraphs):
        txt = re.sub(r"\s+", " ", paragraph_text(p)).strip()
        if not txt:
            continue
        style_name = getattr(getattr(p, "style", None), "name", "") or ""
        item = {"index": i, "text_excerpt": txt[:120], "style": style_name}
        paragraphs.append(item)
        if re.match(r"^\d+(\.\d+)*\s*", txt) or "标题" in style_name.lower():
            heading_candidates.append(item)
    references = [x for x in paragraphs if re.match(r"^\[\d+\]", str(x["text_excerpt"]))]
    captions_figure = [x for x in paragraphs if str(x["text_excerpt"]).startswith("图")]
    captions_table = [x for x in paragraphs if str(x["text_excerpt"]).startswith("表")]
    return {
        "paragraphs": paragraphs[:200],
        "headings": heading_candidates[:80],
        "figure_captions": captions_figure[:80],
        "table_captions": captions_table[:80],
        "references": references[:200],
        "table_count": len(getattr(doc, "tables", [])),
        "image_count": count_document_images(doc),
    }


def apply_llm_safe_edit_plan(fixed_docx: Path, edit_plan: list[dict[str, object]], output_path: Path) -> dict[str, object]:
    if not fixed_docx.exists():
        if not output_path.exists():
            output_path.write_bytes(b"")
        return {"applied_count": 0, "skipped_count": len(edit_plan), "skipped_reasons": ["fixed_docx_missing"]}
    shutil.copyfile(fixed_docx, output_path)
    doc = Document(str(output_path))
    applied = 0
    skipped = 0
    skipped_reasons: list[str] = []
    high_risk_ops = {"section", "toc", "page_field", "page_number", "image_position", "table_border", "page_break", "header_footer"}
    for item in edit_plan:
        if not isinstance(item, dict):
            skipped += 1
            skipped_reasons.append("invalid_plan_item")
            continue
        op = str(item.get("operation", ""))
        if any(token in op for token in high_risk_ops):
            skipped += 1
            skipped_reasons.append(f"high_risk:{op}")
            continue
        if op == "manual_review":
            skipped += 1
            skipped_reasons.append("manual_review")
            continue
        hint = str(item.get("target_hint", ""))
        params = item.get("parameters", {}) if isinstance(item.get("parameters"), dict) else {}
        target = None
        for p in doc.paragraphs:
            if hint and hint in paragraph_text(p):
                target = p
                break
        if target is None:
            skipped += 1
            skipped_reasons.append(f"target_not_found:{hint[:20]}")
            continue
        if op in {"set_alignment", "set_paragraph_format"}:
            align = str(params.get("alignment", ""))
            if align == "left":
                target.paragraph_format.alignment = ALIGN_LEFT
            elif align == "center":
                target.paragraph_format.alignment = ALIGN_CENTER
            elif align == "right":
                target.paragraph_format.alignment = ALIGN_RIGHT
            elif align == "justified":
                target.paragraph_format.alignment = ALIGN_JUSTIFIED
            applied += 1
        elif op == "set_first_line_indent_chars":
            chars = float(params.get("value", 2))
            set_first_line_indent_chars(target, chars)
            applied += 1
        elif op == "set_hanging_indent_chars":
            chars = float(params.get("value", 2))
            set_hanging_indent_chars(target, chars)
            applied += 1
        elif op == "set_spacing":
            line_pt = params.get("line_spacing_pt")
            if line_pt is not None:
                target.paragraph_format.line_spacing = Pt(float(line_pt))
            applied += 1
        elif op == "set_run_font":
            for run in target.runs:
                run.font.name = str(params.get("font_ascii", run.font.name or "Times New Roman"))
            applied += 1
        else:
            skipped += 1
            skipped_reasons.append(f"unsupported:{op}")
    doc.save(str(output_path))
    return {"applied_count": applied, "skipped_count": skipped, "skipped_reasons": skipped_reasons}


def collect_after_render_images(output_dir: Path, max_pages: int | None, all_pages: bool) -> list[dict[str, object]]:
    patterns = [
        output_dir / "render_qa" / "after",
        output_dir / "render_qa" / "fixed",
        output_dir / "screenshots",
    ]
    imgs: list[Path] = []
    for folder in patterns:
        if folder.exists():
            imgs.extend(sorted([p for p in folder.glob("*.png")]))
    unique: list[Path] = []
    seen: set[str] = set()
    for img in imgs:
        key = str(img.resolve())
        if key not in seen:
            seen.add(key)
            unique.append(img)
    def page_no(path: Path) -> int:
        m = re.search(r"(\d+)", path.stem)
        return int(m.group(1)) if m else 10**9
    unique.sort(key=page_no)
    if not all_pages and max_pages is not None:
        unique = unique[:max_pages]
    result: list[dict[str, object]] = []
    for img in unique:
        n = page_no(img)
        result.append(
            {
                "page": n if n != 10**9 else "unknown",
                "reason": "render_qa_after",
                "checks": ["blank_page", "toc_page_number", "toc_start_page", "three_line_table", "figure_caption"],
                "image_path": str(img),
            }
        )
    return result


def _collect_reviewed_pages(report: dict[str, object], issues: list[Issue], max_pages: int) -> list[dict[str, object]]:
    pages: list[dict[str, object]] = []
    seen: set[str] = set()
    render_qa = report.get("render_qa") if isinstance(report.get("render_qa"), dict) else {}
    for page_obj in render_qa.get("pages", []) if isinstance(render_qa, dict) else []:
        if not isinstance(page_obj, dict):
            continue
        after_img = page_obj.get("after")
        if not after_img:
            continue
        page_no = page_obj.get("page")
        dedupe_key = f"render:{page_no}"
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        pages.append(
            {
                "page": page_no if isinstance(page_no, int) else "unknown",
                "reason": "fixed render page",
                "checks": ["blank_page", "toc_start", "toc_page_number", "table_three_line", "figure_near_caption"],
                "image_path": str(after_img),
            }
        )
        if len(pages) >= max_pages:
            return pages
    for issue in issues:
        if len(pages) >= max_pages:
            break
        issue_img = issue.after_screenshot or issue.before_screenshot
        if not issue_img:
            continue
        dedupe_key = f"issue:{issue_img}"
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        pages.append(
            {
                "page": issue.page if issue.page is not None else "unknown",
                "reason": issue.rule_key,
                "checks": ["table_three_line"] if issue.category == "table" else ["figure_near_caption"] if issue.category == "image" else ["blank_page"] if issue.category == "page" else ["format_risk"],
                "image_path": str(issue_img),
            }
        )
    return pages[:max_pages]


def _image_to_data_uri(path_str: str) -> str:
    path = Path(path_str)
    mime = mimetypes.guess_type(path.name)[0] or "image/png"
    return f"data:{mime};base64,{base64.b64encode(path.read_bytes()).decode('ascii')}"


def _is_vision_unsupported_error(exc: Exception) -> bool:
    text = str(exc).lower()
    return any(token in text for token in ["image_url", "multimodal", "vision", "unsupported", "content type", "invalid content"])


def _multipart_form_data(fields: dict[str, str], file_field: str, file_path: Path) -> tuple[bytes, str]:
    boundary = "----CodexBoundary" + datetime.now().strftime("%Y%m%d%H%M%S%f")
    chunks: list[bytes] = []
    for k, v in fields.items():
        chunks.append(f"--{boundary}\r\n".encode("utf-8"))
        chunks.append(f'Content-Disposition: form-data; name="{k}"\r\n\r\n{v}\r\n'.encode("utf-8"))
    mime = mimetypes.guess_type(file_path.name)[0] or "application/octet-stream"
    chunks.append(f"--{boundary}\r\n".encode("utf-8"))
    chunks.append(
        f'Content-Disposition: form-data; name="{file_field}"; filename="{file_path.name}"\r\nContent-Type: {mime}\r\n\r\n'.encode("utf-8")
    )
    chunks.append(file_path.read_bytes())
    chunks.append(b"\r\n")
    chunks.append(f"--{boundary}--\r\n".encode("utf-8"))
    return b"".join(chunks), boundary


def upload_llm_image_file(path: Path, config: LLMReviewConfig, api_key: str) -> dict[str, object]:
    body, boundary = _multipart_form_data({"purpose": "vision"}, "file", path)
    req = urllib.request.Request(
        f"{config.base_url.rstrip('/')}/files",
        data=body,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": f"multipart/form-data; boundary={boundary}",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=config.timeout) as resp:
        return json.loads(resp.read().decode("utf-8"))


def probe_deepseek_image_upload_support(config: LLMReviewConfig, api_key: str, first_image: Path) -> dict[str, object]:
    try:
        uploaded = upload_llm_image_file(first_image, config, api_key)
        return {"supported": True, "upload": uploaded}
    except Exception as exc:
        return {"supported": False, "reason": f"api_error:{exc.__class__.__name__}"}


def collect_llm_review_candidates(
    report: dict[str, object],
    issues: list[Issue],
    fixed_docx: Path,
    output_dir: Path,
    screenshots_dir: Path | None,
    config: LLMReviewConfig,
) -> dict[str, object]:
    spec_text = extract_spec_docx_text(config.spec_docx_path) if config.spec_docx_path else ""
    fixed_structure = extract_fixed_docx_structure_for_llm(fixed_docx)
    reviewed_pages = collect_after_render_images(output_dir, config.max_pages, config.all_pages)
    if not reviewed_pages:
        reviewed_pages = _collect_reviewed_pages(report, issues, config.max_pages)
    issue_dicts = [issue_dict(issue) for issue in issues]
    table_items = [item for item in issue_dicts if "table" in str(item.get("rule_key", "")) or str(item.get("category", "")) == "table"]
    figure_items = [item for item in issue_dicts if "figure" in str(item.get("rule_key", "")) or str(item.get("category", "")) == "image"]
    blank_items = [item for item in issue_dicts if "blank" in str(item.get("rule_key", "")) or str(item.get("category", "")) == "page"]
    ref272 = report.get("reference_272_check_summary") if isinstance(report.get("reference_272_check_summary"), dict) else {}
    ref273 = report.get("reference_273_check_summary") if isinstance(report.get("reference_273_check_summary"), dict) else {}
    ref272_entries = [e for e in ref272.get("entries", []) if isinstance(e, dict) and e.get("status") != "matched"] if isinstance(ref272, dict) else []
    ref273_entries = [e for e in ref273.get("entries", []) if isinstance(e, dict)] if isinstance(ref273, dict) else []
    return {
        "reviewed_pages": [{k: v for k, v in page.items() if k != "image_path"} for page in reviewed_pages],
        "review_image_paths": [str(page["image_path"]) for page in reviewed_pages if page.get("image_path")],
        "spec_docx_path": str(config.spec_docx_path) if config.spec_docx_path else "",
        "original_docx_path": str(report.get("input", "")),
        "spec_text": spec_text,
        "fixed_docx_structure": fixed_structure,
        "summary": {
            "issue_summary_by_category": report.get("issue_summary_by_category", {}),
            "layout_fix_policy": report.get("layout_fix_policy", {}),
            "reference_272_check_summary": {
                "total_entries": ref272.get("total_entries", 0) if isinstance(ref272, dict) else 0,
                "type_mismatch_entries": ref272.get("type_mismatch_entries", 0) if isinstance(ref272, dict) else 0,
                "unmatched_entries": ref272.get("unmatched_entries", 0) if isinstance(ref272, dict) else 0,
            },
            "reference_273_check_summary": {
                "online_entries": ref273.get("online_entries", 0) if isinstance(ref273, dict) else 0,
                "citation_date_missing_entries": ref273.get("citation_date_missing_entries", 0) if isinstance(ref273, dict) else 0,
                "access_path_missing_entries": ref273.get("access_path_missing_entries", 0) if isinstance(ref273, dict) else 0,
                "doi_or_url_missing_entries": ref273.get("doi_or_url_missing_entries", 0) if isinstance(ref273, dict) else 0,
            },
            "source_image_count": report.get("source_image_count"),
            "fixed_image_count": report.get("fixed_image_count"),
            "screenshot_status": report.get("screenshot_status"),
            "renderer_for_fixed": report.get("renderer_for_fixed"),
            "fixed_docx": str(fixed_docx),
            "original_docx": str(report.get("input", "")),
            "target_docx": str(fixed_docx) if config.doc_target == "fixed" else str(report.get("input", "")),
            "screenshots_dir": str(screenshots_dir) if screenshots_dir else "",
        },
        "toc_manual_review": {
            "toc_page_number_manual_review": any(item.get("rule_key") == "toc_page_number_manual_review" for item in issue_dicts),
            "toc_start_new_page_manual_review": any(item.get("rule_key") == "toc_start_new_page_manual_review" for item in issue_dicts),
            "toc_body_section_manual_review": any(item.get("rule_key") == "toc_body_section_manual_review" for item in issue_dicts),
            "note": "目录条目右侧页码应保持阿拉伯数字；目录页底部页码属于前置页，应为小写罗马数字。",
        },
        "blank_or_page_risks": [{"rule_key": x.get("rule_key"), "excerpt": _truncate_text(x.get("text_excerpt")), "message": _truncate_text(x.get("message"), 320)} for x in blank_items[: config.max_pages]],
        "table_risks": [{"rule_key": x.get("rule_key"), "excerpt": _truncate_text(x.get("text_excerpt")), "message": _truncate_text(x.get("message"), 320)} for x in table_items[: config.max_pages]],
        "figure_risks": [{"rule_key": x.get("rule_key"), "excerpt": _truncate_text(x.get("text_excerpt")), "message": _truncate_text(x.get("message"), 320)} for x in figure_items[: config.max_pages]],
        "reference_risks": {
            "reference_272_non_matched": ref272_entries[: config.max_pages],
            "reference_273_manual_review": ref273_entries[: config.max_pages],
        },
    }


def _extract_chat_content(payload: dict) -> str:
    choices = payload.get("choices")
    if not isinstance(choices, list) or not choices:
        return ""
    first = choices[0] if isinstance(choices[0], dict) else {}
    msg = first.get("message") if isinstance(first, dict) else {}
    if not isinstance(msg, dict):
        return ""
    return str(msg.get("content") or "")


def _build_llm_user_prompt(candidates: dict[str, object], *, fallback_used: bool) -> str:
    schema = {
        "llm_review_version": "5.0",
        "provider": "dashscope",
        "model": "qwen-long",
        "read_check": {
            "spec_file_read": True,
            "target_file_read": True,
            "title": "",
            "keywords": "",
            "toc_evidence": "",
            "reference_numbering_observation": "",
        },
        "basis": {
            "spec_file_read": True,
            "target_file_read": True,
            "uses_spec_docx": True,
            "uses_target_docx": True,
            "target_docx_source": "fixed|original",
            "review_mode": "fileid_docx",
            "document_upload_status": "ok|fallback_text|failed",
            "fallback_used": False,
            "fallback_reason": "",
            "model": "qwen-long",
        },
        "overall_result": "pass|needs_manual_review|needs_fix",
        "overall_risk": "low|medium|high",
        "issues": [],
        "manual_review_items": [],
        "uncertain_items": [],
        "safe_edit_plan": [],
        "rejected_claims": [],
        "validation_warnings": [],
        "notes": [],
    }
    fixed_structure = candidates.get("fixed_docx_structure", {})
    if isinstance(fixed_structure, dict):
        fixed_structure = {
            "headings": fixed_structure.get("headings", [])[:80],
            "figure_captions": fixed_structure.get("figure_captions", [])[:80],
            "table_captions": fixed_structure.get("table_captions", [])[:80],
            "references": fixed_structure.get("references", [])[:50],
            "table_count": fixed_structure.get("table_count", 0),
            "image_count": fixed_structure.get("image_count", 0),
            "paragraphs": fixed_structure.get("paragraphs", [])[:80],
        }
    return (
        "你是学士论文格式审查助手。请执行证据优先审查：仅当你能给出明确证据时，才可写入 issues。"
        " 每个 issues 项必须包含 evidence_exact_quote，且必须是你从论文中可直接引用的短句。"
        " 无法确认的内容不得进入 issues，必须进入 manual_review_items 或 uncertain_items。"
        " 禁止编造目录号 1.a/1.b、日期、URL、DOI、页码。"
        " 对 TOC 域、PAGE 域、分节符、页眉页脚继承、图片位置、表格边框等高风险项，只给人工建议，不给自动执行方案。"
        " 你必须只输出 JSON。output only JSON.\n\n"
        f"返回 JSON schema: {json.dumps(schema, ensure_ascii=False)}\n\n"
        f"fallback_used={fallback_used}\n"
        f"spec_docx_path={candidates.get('spec_docx_path', '')}\n"
        f"target_docx_path={candidates.get('summary', {}).get('target_docx', candidates.get('summary', {}).get('fixed_docx', ''))}\n"
        f"spec_text={str(candidates.get('spec_text', ''))[:28000]}\n"
        f"target_docx_structure={json.dumps(fixed_structure, ensure_ascii=False)}\n"
        f"deterministic_summary={json.dumps(candidates.get('summary', {}), ensure_ascii=False)}\n"
        f"risk_summary={json.dumps({'table_risks': candidates.get('table_risks', []), 'figure_risks': candidates.get('figure_risks', []), 'blank_or_page_risks': candidates.get('blank_or_page_risks', [])}, ensure_ascii=False)}"
    )



def _build_evidence_first_fileid_prompt() -> str:
    return json.dumps({
        "llm_review_version": "6.0",
        "provider": "dashscope",
        "model": "qwen-long",
        "read_check": {
            "spec_file_read": True,
            "target_file_read": True,
            "title": "",
            "keywords": "",
            "toc_evidence": "",
            "reference_numbering_observation": "",
        },
        "basis": {
            "spec_file_read": True,
            "target_file_read": True,
            "review_mode": "qwen_long_fileid_docx",
            "document_upload_status": "ok",
            "fallback_used": False,
            "fallback_reason": "",
        },
        "overall_result": "pass|needs_manual_review|needs_fix",
        "overall_risk": "low|medium|high",
        "issues": [],
        "manual_review_items": [],
        "uncertain_items": [],
        "safe_edit_plan": [],
        "rejected_claims": [],
        "validation_warnings": [],
        "notes": [],
    }, ensure_ascii=False) + ("\n\n"
        "请严格依据第一份《成都信息工程大学学士学位论文规范》审查第二份论文文档。\n\n"
        "【极其重要的证据规则】\n"
        "1. 每一个“明确不符合规范的问题”都必须给出第二份论文中的逐字原文证据。\n"
        "2. 如果你不能在第二份论文中找到对应原文，不得声称该问题存在。\n"
        "3. 对 Word 自动编号、目录域、页码域、参考文献自动编号、页眉页脚、分节符、表格边框、图片位置等内容，如果无法从文档文本中可靠判断，必须放入 manual_review_items，不得放入 issues。\n"
        "4. 不得编造文件中不存在的标题编号，例如不得在没有证据时声称出现“1.a”“1.b”。\n"
        "5. 不得编造专利公告日期、参考文献页码、访问日期、URL 或 DOI。\n"
        "6. 对参考文献编号要特别谨慎：Word 自动编号可能不会作为段落正文被抽取。如果你无法确认编号显示样式，只能列为人工复核；不得直接断言“未使用[1][2]编号”。\n"
        "7. evidence_found=false 的项目禁止进入 issues，只能进入 manual_review_items 或 uncertain_items。\n"
        "8. 如果发现问题，请说明你是从“论文原文证据”判断，还是从“版式/域/自动编号风险”推断。\n\n"
        "只输出 JSON，不要输出 Markdown，不要包裹 `json 代码块。"
    )

def _docx_plain_text_for_evidence(path: Path) -> str:
    if not path.exists() or not path.is_file() or path.suffix.lower() != ".docx":
        return ""
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        text = paragraph_text(p).strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def _postprocess_llm_evidence(review: dict[str, object], local_text: str) -> dict[str, object]:
    if not isinstance(review, dict):
        return review
    warnings = review.get("validation_warnings") if isinstance(review.get("validation_warnings"), list) else []
    manual = review.get("manual_review_items") if isinstance(review.get("manual_review_items"), list) else []
    issues = review.get("issues") if isinstance(review.get("issues"), list) else []
    safe_plan = review.get("safe_edit_plan") if isinstance(review.get("safe_edit_plan"), list) else []
    kept_issues: list[object] = []
    lowered = local_text or ""
    for item in issues:
        if not isinstance(item, dict):
            manual.append({"reason": str(item), "suggestion": "请人工复核。"})
            warnings.append("issue 非结构化，已转入 manual_review_items。")
            continue
        evidence_found = item.get("evidence_found", True)
        quote = str(item.get("evidence_exact_quote", "") or "").strip()
        if evidence_found is not True:
            manual.append({"reason": f"evidence_found!=true: {item.get('id', '')}", "suggestion": str(item.get("suggestion", "请人工复核。")), "related_evidence": quote})
            warnings.append(f"issue {item.get('id', '')} evidence_found!=true，已转人工复核。")
            continue
        if not quote:
            manual.append({"reason": f"evidence_exact_quote 为空: {item.get('id', '')}", "suggestion": str(item.get("suggestion", "请人工复核。"))})
            warnings.append(f"issue {item.get('id', '')} 无 evidence_exact_quote，已转人工复核。")
            continue
        if quote not in lowered:
            manual.append({"reason": f"本地未命中证据引用: {item.get('id', '')}", "suggestion": str(item.get("suggestion", "请人工复核。")), "related_evidence": quote})
            warnings.append(f"issue {item.get('id', '')} 的 evidence_exact_quote 未在本地文档中找到，已转人工复核。")
            continue
        kept_issues.append(item)

    for pattern in ["1.a", "1.b", "2.a", "2.b", "2022-05-10"]:
        if pattern in json.dumps(review, ensure_ascii=False) and pattern not in lowered:
            warnings.append(f"检测到高风险可疑模式 {pattern}，但本地文档无证据命中。")

    normalized_plan: list[dict[str, object]] = []
    for item in safe_plan:
        if not isinstance(item, dict):
            warnings.append("safe_edit_plan 含非结构化项，已忽略。")
            continue
        risk = str(item.get("risk", "low")).lower()
        quote = str(item.get("evidence_exact_quote", "") or "")
        if risk != "low":
            item["can_auto_fix"] = False
            item["blocked_reason"] = "risk_not_low"
        elif quote and quote not in lowered:
            item["can_auto_fix"] = False
            item["blocked_reason"] = "evidence_not_found_locally"
            warnings.append(f"safe_edit_plan {item.get('id', '')} 本地证据未命中，标记为不可自动执行。")
        else:
            item["can_auto_fix"] = False
            item["blocked_reason"] = "llm_plan_not_auto_applied"
        normalized_plan.append(item)

    review["issues"] = kept_issues
    review["manual_review_items"] = manual
    review["safe_edit_plan"] = normalized_plan
    review["validation_warnings"] = warnings
    return review


def probe_dashscope_document_upload_support(config: LLMReviewConfig, api_key: str, spec_docx: Path, fixed_docx: Path) -> dict[str, object]:
    try:
        if not spec_docx.exists() or not fixed_docx.exists():
            return {"supported": False, "method": "fallback_text_extraction", "reason": "missing_docx_file"}
        client = get_dashscope_client(api_key=api_key, base_url=config.base_url)
        spec_file_id = resolve_spec_file_id(client, config.spec_file_id, spec_docx, timeout_seconds=config.timeout)
        fixed_file_id = upload_target_docx(client, fixed_docx, timeout_seconds=config.timeout)
        if spec_file_id and fixed_file_id:
            return {
                "supported": True,
                "method": "file_upload",
                "file_ids": {"spec": str(spec_file_id), "target": str(fixed_file_id)},
            }
        return {"supported": False, "method": "fallback_text_extraction", "reason": "missing_file_id"}
    except Exception as exc:
        return {"supported": False, "method": "fallback_text_extraction", "reason": f"upload_error:{exc.__class__.__name__}"}


def run_qwen_docx_review(*, config: LLMReviewConfig, api_key: str, candidates: dict[str, object]) -> dict[str, object]:
    spec_docx = Path(str(candidates.get("spec_docx_path") or ""))
    fixed_docx = Path(str(candidates.get("summary", {}).get("fixed_docx", "")))
    original_docx = Path(str(candidates.get("summary", {}).get("original_docx", "")))
    target_docx = original_docx if config.doc_target == "original" and original_docx.exists() else fixed_docx
    client = get_dashscope_client(api_key=api_key, base_url=config.base_url)
    probe = probe_dashscope_document_upload_support(config, api_key, spec_docx, target_docx)
    upload_ok = bool(probe.get("supported"))
    fallback_used = not upload_ok
    fallback_reason = "" if upload_ok else str(probe.get("reason", "upload_not_supported"))
    doc_model_error_type = ""
    doc_model_error_message = ""
    raw_text = ""
    if upload_ok:
        file_ids = probe.get("file_ids", {})
        target_id = str(file_ids.get("target") or file_ids.get("fixed", ""))
        fileid_prompt = _build_evidence_first_fileid_prompt()
        try:
            raw_text = run_qwen_long_docx_review(
                client=client,
                spec_file_id=str(file_ids.get("spec", "")),
                target_file_id=target_id,
                model=config.doc_model,
                user_prompt=fileid_prompt,
                timeout_seconds=config.timeout,
            )
            if config.delete_uploaded_target_after_review and target_id:
                try:
                    dashscope_delete_uploaded_file(client, target_id)
                except Exception:
                    pass
        except Exception as exc:
            upload_ok = False
            fallback_used = True
            fallback_reason = f"doc_model_error:{exc.__class__.__name__}"
            doc_model_error_type = exc.__class__.__name__
            doc_model_error_message = str(exc)[:1000]
    if not raw_text:
        fallback_prompt = _build_llm_user_prompt(candidates, fallback_used=True)
        payload = call_openai_compatible_chat(
            base_url=config.base_url,
            api_key=api_key,
            model=config.review_model,
            messages=[
                {"role": "system", "content": LLM_REVIEW_SYSTEM_PROMPT},
                {"role": "user", "content": fallback_prompt},
            ],
            timeout=config.timeout,
        )
        raw_text = _extract_chat_content(payload)
    parsed = parse_llm_review_response(raw_text)
    if parsed.get("status") == "failed":
        parsed = {
            "llm_review_version": "4.0",
            "overall_result": "needs_manual_review",
            "overall_risk": "medium",
            "issues": [],
            "manual_review_items": [],
            "safe_edit_plan": [],
            "notes": [],
            "raw_text": raw_text,
        }

    basis = parsed.get("basis")
    if not isinstance(basis, dict):
        basis = {}
        parsed["basis"] = basis
    basis["uses_spec_docx"] = bool(candidates.get("spec_text")) or bool(probe.get("supported"))
    basis["uses_target_docx"] = True
    basis["uses_fixed_docx"] = config.doc_target == "fixed"
    basis["target_docx_source"] = config.doc_target
    basis.setdefault("uses_rendered_images", False)
    basis["document_upload_status"] = "ok" if upload_ok else "fallback_text"
    basis["fallback_used"] = fallback_used
    basis["fallback_reason"] = fallback_reason
    basis["doc_model_attempted"] = config.doc_model
    basis["fallback_model"] = config.review_model if fallback_used else ""
    basis["doc_model_error_type"] = doc_model_error_type
    basis["doc_model_error_message"] = doc_model_error_message
    basis["doc_model_fileid_attempted"] = bool(upload_ok or fallback_reason.startswith("doc_model_error:"))
    basis["doc_model_fixed_docx_name"] = target_docx.name if target_docx else ""
    basis["target_file_read"] = upload_ok or fallback_used
    basis["spec_file_read"] = bool(config.spec_file_id or (spec_docx and spec_docx.exists()))
    basis["review_mode"] = "qwen_long_fileid_docx" if upload_ok and not fallback_used else "fallback_summary"
    basis["model"] = config.doc_model
    basis.setdefault("spec_name", spec_docx.name if spec_docx else "")
    basis.setdefault("fixed_docx_name", target_docx.name if target_docx else "")

    parsed.setdefault("llm_review_version", "4.0")
    parsed.setdefault("provider", config.provider)
    parsed.setdefault("model", config.doc_model if upload_ok else config.review_model)
    parsed.setdefault("overall_result", "needs_manual_review")
    parsed.setdefault("overall_risk", "medium")
    parsed.setdefault("issues", [])
    parsed.setdefault("manual_review_items", [])
    parsed.setdefault("safe_edit_plan", [])
    parsed.setdefault("notes", [])
    actual_review_model = config.review_model if fallback_used else config.doc_model
    parsed["actual_review_model"] = actual_review_model
    basis["actual_review_model"] = actual_review_model
    return {"status": "ok", "review": parsed, "probe": probe}


def run_llm_review(config: LLMReviewConfig, candidates: dict[str, object], dotenv_values: dict[str, str] | None = None) -> dict[str, object]:
    if not config.enabled:
        return {"enabled": False}
    dotenv_values = dotenv_values or {}
    api_key = os.environ.get(config.api_key_env) or dotenv_values.get(config.api_key_env, "")
    if not api_key:
        return {
            "enabled": True,
            "provider": config.provider,
            "model": config.doc_model,
            "mode": config.mode,
            "base_url": config.base_url,
            "status": "skipped",
            "reason": "missing_api_key",
        }
    try:
        result = run_qwen_docx_review(config=config, api_key=api_key, candidates=candidates)
        if result.get("status") != "ok":
            return {
                "enabled": True,
                "provider": config.provider,
                "model": config.doc_model,
                "mode": config.mode,
                "base_url": config.base_url,
                "status": "failed",
                "reason": f"api_error:{result.get('reason', 'unknown')}",
                "doc_model_attempted": config.doc_model,
                "fallback_model": config.review_model,
                "fallback_used": False,
                "fallback_reason": str(result.get("reason", "unknown")),
                "document_upload_status": "failed",
                "uses_spec_docx": bool(candidates.get("spec_text")) or bool(candidates.get("spec_docx_path")),
                "uses_fixed_docx": bool(candidates.get("summary", {}).get("fixed_docx")),
                "raw_error": str(result.get("reason", "unknown")),
            }
        review = result.get("review", {})
        if isinstance(review, dict):
            target_source = str(review.get("basis", {}).get("target_docx_source", config.doc_target))
            target_path = Path(str(candidates.get("summary", {}).get("fixed_docx", "")))
            if target_source == "original":
                target_path = Path(str(candidates.get("summary", {}).get("original_docx", "")))
            local_text = _docx_plain_text_for_evidence(target_path) if target_path else ""
            review = _postprocess_llm_evidence(review, local_text)
            if isinstance(review.get("basis"), dict):
                review["basis"]["target_docx_source"] = target_source
        response = {
            "enabled": True,
            "provider": config.provider,
            "model": config.doc_model,
            "mode": config.mode,
            "base_url": config.base_url,
            "status": "ok",
            "review": review,
            "document_upload_status": review.get("basis", {}).get("document_upload_status", "fallback_text"),
            "fallback_used": review.get("basis", {}).get("fallback_used", True),
            "fallback_reason": review.get("basis", {}).get("fallback_reason", ""),
            "uses_spec_docx": review.get("basis", {}).get("uses_spec_docx", False),
            "uses_fixed_docx": review.get("basis", {}).get("uses_fixed_docx", False),
            "doc_model_attempted": review.get("basis", {}).get("doc_model_attempted", config.doc_model),
            "fallback_model": review.get("basis", {}).get("fallback_model", ""),
            "actual_review_model": review.get(
                "actual_review_model",
                review.get("basis", {}).get("actual_review_model", config.review_model if review.get("basis", {}).get("fallback_used", True) else config.doc_model),
            ),
            "image_upload_mode": "docx_upload" if review.get("basis", {}).get("document_upload_status") == "ok" else "fallback_text_extraction",
        }
        if config.output_path is not None:
            config.output_path.parent.mkdir(parents=True, exist_ok=True)
            config.output_path.write_text(json.dumps(response, ensure_ascii=False, indent=2), encoding="utf-8", newline="\n")
        return response
    except Exception as exc:
        return {
            "enabled": True,
            "provider": config.provider,
            "model": config.model,
            "mode": config.mode,
            "base_url": config.base_url,
            "status": "failed",
            "reason": f"api_error:{exc.__class__.__name__}",
            "doc_model_attempted": config.doc_model,
            "fallback_model": config.review_model,
            "fallback_used": False,
            "fallback_reason": f"api_error:{exc.__class__.__name__}",
            "document_upload_status": "failed",
            "uses_spec_docx": bool(candidates.get("spec_text")) or bool(candidates.get("spec_docx_path")),
            "uses_fixed_docx": bool(candidates.get("summary", {}).get("fixed_docx")),
            "raw_error": f"{exc.__class__.__name__}: {exc}",
        }

def run(
    input_path: Path,
    output_dir: Path,
    renderer: str,
    screenshots: str,
    allow_ooxml_layout_fixes: bool = False,
    run_timestamp: str | None = None,
    llm_review_config: LLMReviewConfig | None = None,
    dotenv_values: dict[str, str] | None = None,
) -> dict[str, object]:
    require_docx_dependencies()
    llm_review_config = llm_review_config or LLMReviewConfig()
    dotenv_values = dotenv_values or {}
    if not input_path.exists():
        raise FileNotFoundError(input_path)
    if not is_docx(input_path):
        raise ValueError("Input must be a .docx file.")
    output_dir.mkdir(parents=True, exist_ok=True)
    json_output_dir = output_dir / "json"
    json_output_dir.mkdir(parents=True, exist_ok=True)

    run_timestamp = run_timestamp or build_run_timestamp()
    comments_path = append_timestamp_to_name(output_dir / f"{input_path.stem}_format_comments.docx", run_timestamp)
    fixed_path = append_timestamp_to_name(output_dir / f"{input_path.stem}_format_fixed.docx", run_timestamp)
    report_path = append_timestamp_to_name(json_output_dir / f"{input_path.stem}_format_report.json", run_timestamp)
    html_report_path = append_timestamp_to_name(output_dir / f"{input_path.stem}_format_report.html", run_timestamp)
    llm_review_json_path = append_timestamp_to_name(json_output_dir / f"{input_path.stem}_llm_review.json", run_timestamp)
    llm_review_html_path = append_timestamp_to_name(output_dir / f"{input_path.stem}_llm_review.html", run_timestamp)
    llm_raw_json_path = append_timestamp_to_name(json_output_dir / f"{input_path.stem}_llm_raw.json", run_timestamp)
    llm_messages_json_path = append_timestamp_to_name(json_output_dir / f"{input_path.stem}_llm_messages.json", run_timestamp)

    source_doc = Document(str(input_path))
    source_image_count = count_document_images(source_doc)
    source_figure_caption_count = count_figure_captions(source_doc)
    normalize_document_structure(source_doc)
    with TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir) / f"{input_path.stem}_normalized.docx"
        source_doc.save(str(temp_path))

        normalized_doc = Document(str(temp_path))
        source_texts = [paragraph_text(paragraph) for paragraph in normalized_doc.paragraphs]
        structure_analysis = analyze_section_sequence(source_texts, has_toc_field=document_has_toc_field(normalized_doc))
        source_regions = structure_analysis["regions"]
        ref_entries_for_summary, _ref_entry_issues = _reference_entries_from_regions(normalized_doc.paragraphs, source_texts, source_regions)
        reference_272_check_summary = _reference_272_check_summary(ref_entries_for_summary)
        reference_273_check_summary = _reference_273_check_summary(ref_entries_for_summary)
        issues = collect_issues(normalized_doc)
        create_annotated_docx(temp_path, comments_path, issues)

        planned_renderer = probe_com_renderer(renderer)
        allow_layout_fixes = planned_renderer in {"office", "wps"} or allow_ooxml_layout_fixes
        fixed_doc = Document(str(temp_path))
        apply_supported_rules(fixed_doc, allow_layout_fixes=allow_layout_fixes)
        fixed_doc.save(str(fixed_path))
    fixed_renderer = try_com_save(fixed_path, renderer)
    comments_renderer = try_com_save(comments_path, renderer)
    # Some Word/WPS save paths can reintroduce boundary placeholders (for TOC fields/section metadata).
    post_layout_doc = Document(str(fixed_path))
    _post_layout_issues, post_layout_cleanup_changed = cleanup_visible_blank_paragraphs_after_layout(post_layout_doc)
    if post_layout_cleanup_changed:
        post_layout_doc.save(str(fixed_path))
        fixed_renderer = try_com_save(fixed_path, renderer)
    fixed_doc_for_counts = Document(str(fixed_path))
    fixed_image_count = count_document_images(fixed_doc_for_counts)
    fixed_figure_caption_count = count_figure_captions(fixed_doc_for_counts)
    if fixed_image_count < source_image_count:
        issues.append(
            Issue(
                paragraph_index=-1,
                rule_key="image_lost_after_fix",
                text_type="image integrity",
                text_excerpt="document image inventory",
                current=f"修复后图片数量 {fixed_image_count} 少于源文档 {source_image_count}。",
                expected="修复后图片数量不应少于源文档。",
                message="修复后图片数量少于源文档，可能存在图片丢失。",
                category="image",
                location="document images",
            )
        )
    attach_after_formats(fixed_path, issues)
    screenshot_status, render_qa = attach_screenshots(
        input_path=input_path,
        fixed_path=fixed_path,
        issues=issues,
        output_dir=output_dir,
        renderer_used=fixed_renderer,
        mode=screenshots,
    )

    _header_start, header_reason = find_header_start_section(normalized_doc)
    advisory = [
        "Semantic correctness of cover fields is not checked.",
        "Bibliographic content validity is not fully checked.",
        "Dynamic page number fields are preserved where present; this script does not rebuild complex field codes.",
    ]
    if header_reason == "fallback-front-matter":
        advisory.append(
            "No declaration page was detected. Header insertion starts from the first recognized abstract/TOC/body section, so verify the required declaration page and header start manually."
        )
    if header_reason == "not-found":
        advisory.append(
            "No reliable header start page was detected. The script cannot safely decide where the declaration-page header should begin."
        )
    if fixed_renderer == "ooxml":
        advisory.append(
            "The fixed DOCX was produced through OOXML fallback. Render QA images are approximate previews and may not match Microsoft Word or WPS exactly; use Word/WPS for final visual confirmation when possible."
        )
        advisory.append("由于未使用 Word/WPS COM，页码与分页检查为有限检查/报告型检查。")
        advisory.append("纯 OOXML 模式不能完全确认渲染后的封面第一页边界；当前按文档开头至后续模块标记之前作为封面逻辑区域，请用 Word/WPS 复核。")
    if fixed_renderer == "ooxml" and not allow_ooxml_layout_fixes:
        advisory.append(
            "OOXML mode did not automatically modify high-risk layout features: headers, footers, page numbers, header/footer distance, or section-start behavior. Add --allow-ooxml-layout-fixes only if the user explicitly accepts possible pagination drift."
        )
    if fixed_renderer == "ooxml" and allow_ooxml_layout_fixes:
        advisory.append(
            "OOXML high-risk layout fixes were explicitly enabled. Verify headers, footers, page numbers, section breaks, and pagination in Word/WPS before using the fixed document."
        )

    resolved_issue_count = sum(
        1
        for issue in issues
        if issue.after is not None
        and "疑似不符合项：" not in issue.after
        and "无法可靠定位修改后状态" not in issue.after
    )
    remaining_issue_count = sum(
        1
        for issue in issues
        if issue.after is None
        or "疑似不符合项：" in (issue.after or "")
        or "无法可靠定位修改后状态" in (issue.after or "")
    )
    manual_review_count = sum(
        1 for issue in issues if "人工确认" in (issue.message or "") or "人工确认" in (issue.current or "")
    )

    report = {
        "input": str(input_path.resolve()),
        "run_timestamp": run_timestamp,
        "annotated_docx": str(comments_path.resolve()),
        "fixed_docx": str(fixed_path.resolve()),
        "json_report": str(report_path.resolve()),
        "html_report": str(html_report_path.resolve()),
        "issue_count": len(issues),
        "modification_count": len(issues),
        "resolved_issue_count": resolved_issue_count,
        "remaining_issue_count": remaining_issue_count,
        "manual_review_count": manual_review_count,
        "issue_count_note": "issue_count 表示发现并尝试处理的问题数，不等同于修复后剩余问题数。",
        "issue_summary_by_category": summarize_issues(issues),
        "renderer_for_fixed": fixed_renderer,
        "renderer_for_comments": comments_renderer,
        "planned_renderer_for_layout_fixes": planned_renderer,
        "high_risk_layout_fixes_applied": allow_layout_fixes,
        "allow_ooxml_layout_fixes": allow_ooxml_layout_fixes,
        "post_layout_cleanup_changed": post_layout_cleanup_changed,
        "source_image_count": source_image_count,
        "fixed_image_count": fixed_image_count,
        "source_figure_caption_count": source_figure_caption_count,
        "fixed_figure_caption_count": fixed_figure_caption_count,
        "screenshot_status": screenshot_status,
        "render_qa": render_qa,
        "structure_analysis": structure_analysis,
        "reference_272_check_summary": reference_272_check_summary,
        "reference_273_check_summary": reference_273_check_summary,
        "layout_fix_policy": {
            "toc_page_numbering": "report_only",
            "toc_start_new_page": "report_only",
            "reason": "自动目录 TOC 域、section break、页脚 PAGE 域与 Word/WPS 字段刷新高度耦合；自动修改曾造成空白页回归，因此当前仅报告并提示人工复核。",
        },
        "structure_warnings": structure_analysis.get("warnings", []) + compute_section_boundary_findings(source_texts, structure_analysis["regions"])[0] + compute_section_boundary_findings(source_texts, structure_analysis["regions"])[1],
        "issues": [issue_dict(issue) for issue in issues],
        "unsupported_or_advisory": advisory,
    }
    candidates = collect_llm_review_candidates(report, issues, fixed_path, output_dir, output_dir / "screenshots", llm_review_config)
    report["llm_review"] = run_llm_review(llm_review_config, candidates, dotenv_values=dotenv_values)
    llm_review_obj = report["llm_review"] if isinstance(report["llm_review"], dict) else {}
    llm_review_payload = llm_review_obj.get("review", {}) if isinstance(llm_review_obj.get("review"), dict) else {}
    report["llm_review_json"] = str(llm_review_json_path.resolve())
    report["llm_review_html"] = str(llm_review_html_path.resolve())
    report["json_dir"] = str(json_output_dir.resolve())
    if isinstance(report.get("llm_review"), dict):
        report["llm_review"]["json_dir"] = report["json_dir"]
    llm_review_json_path.write_text(json.dumps(report["llm_review"], ensure_ascii=False, indent=2), encoding="utf-8", newline="\n")
    llm_raw_json_path.write_text(json.dumps(llm_review_payload, ensure_ascii=False, indent=2), encoding="utf-8", newline="\n")
    llm_messages_json_path.write_text(
        json.dumps(
            {
                "provider": llm_review_obj.get("provider", ""),
                "model": llm_review_obj.get("model", ""),
                "actual_review_model": llm_review_obj.get("actual_review_model", ""),
                "basis": llm_review_payload.get("basis", {}) if isinstance(llm_review_payload, dict) else {},
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
        newline="\n",
    )
    if llm_review_config.separate_html:
        write_llm_review_separate_html(report["llm_review"], llm_review_html_path)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    write_html_report(report, html_report_path)
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description="Check and fix bachelor thesis DOCX formatting.")
    parser.add_argument("docx", help="Input .docx thesis path")
    parser.add_argument("--output-dir", default="./format_results", help="Output directory")
    parser.add_argument("--renderer", choices=["auto", "office", "wps", "ooxml"], default="auto")
    parser.add_argument(
        "--screenshots",
        choices=["auto", "never", "require"],
        default="auto",
        help="Generate before/after page screenshots when Office/WPS rendering is available.",
    )
    parser.add_argument(
        "--allow-ooxml-layout-fixes",
        action="store_true",
        help=(
            "Explicitly allow OOXML-only changes to high-risk layout features "
            "(headers, footers, page numbers, header/footer distance, and section behavior)."
        ),
    )
    parser.add_argument("--run-timestamp", default=None, help="Optional fixed run timestamp, e.g. 20260513_155821")
    parser.add_argument("--llm-review", action="store_true", help="Enable LLM high-risk format review.")
    parser.add_argument("--llm-provider", default=None, help="LLM provider name, default dashscope.")
    parser.add_argument("--llm-model", default=None, help="LLM model name, default qwen3.6-plus.")
    parser.add_argument("--llm-review-model", default=None, help="LLM review model name, default qwen3.6-plus.")
    parser.add_argument("--llm-doc-model", default=None, help="LLM doc understanding model, default qwen-long.")
    parser.add_argument("--llm-api-key-env", default="DASHSCOPE_API_KEY", help="Environment variable name for API key.")
    parser.add_argument("--llm-base-url", default=None, help="OpenAI-compatible base URL.")
    parser.add_argument("--llm-review-mode", choices=["auto", "text", "vision"], default="auto", help="LLM review mode.")
    parser.add_argument("--llm-review-max-pages", type=int, default=8, help="Max candidate item count for review.")
    parser.add_argument("--llm-review-spec-docx", default=None, help="Spec DOCX path for LLM review.")
    parser.add_argument("--llm-review-spec-file-id", default=None, help="Spec DOCX file-id for LLM review.")
    parser.add_argument("--llm-review-doc-target", choices=["fixed", "original"], default="fixed", help="Use fixed or original docx as LLM review target.")
    parser.add_argument("--llm-review-upload-target-docx", action="store_true", default=True, help="Upload target docx for doc review.")
    parser.add_argument("--llm-review-delete-uploaded-target-after-review", action="store_true", help="Delete uploaded target docx after review.")
    parser.add_argument("--llm-review-all-pages", action="store_true", help="Use all rendered after pages for LLM review.")
    parser.add_argument("--llm-review-batch-size", type=int, default=6, help="Rendered image batch size for LLM review.")
    parser.add_argument("--llm-review-image-upload", action="store_true", default=True, help="Enable direct image upload for LLM review.")
    parser.add_argument("--llm-review-image-transport", choices=["data_uri"], default="data_uri", help="Image transport for LLM vision.")
    parser.add_argument("--llm-review-no-separate-html", action="store_true", help="Disable separate LLM review HTML output.")
    parser.add_argument("--llm-review-output", default=None, help="Optional path to save raw LLM review JSON.")
    parser.add_argument("--llm-review-timeout", type=int, default=600, help="LLM request timeout seconds.")
    args = parser.parse_args()
    dotenv_values = load_dotenv_file(DEFAULT_DOTENV_PATH)
    llm_config = resolve_llm_review_config(args, dotenv_values)
    report = run(
        Path(args.docx),
        Path(args.output_dir),
        args.renderer,
        args.screenshots,
        allow_ooxml_layout_fixes=args.allow_ooxml_layout_fixes,
        run_timestamp=args.run_timestamp,
        llm_review_config=llm_config,
        dotenv_values=dotenv_values,
    )
    print(stdout_json(report))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
