#!/usr/bin/env python3
"""Build a DOCX review map for CUIT thesis format rule locations."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
SPEC_DOCX = ROOT / "成都信息工程大学学士学位论文规范.docx"
OUTPUT_DOCX = ROOT / "cuit-thesis-rule-position-review.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size, color in [
        ("Title", 20, "1F4E79"),
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 12, "365F91"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for idx, (text, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[idx]
        cell.width = Cm(width)
        cell.text = text
        set_cell_shading(cell, "D9EAF7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, (text, width) in enumerate(zip(row, widths)):
            cell = cells[idx]
            cell.width = Cm(width)
            cell.text = text
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def main() -> None:
    if not SPEC_DOCX.exists():
        raise FileNotFoundError(SPEC_DOCX)

    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("成都信息工程大学学士学位论文格式规则位置审核稿").bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"来源文档：{SPEC_DOCX.name}").italic = True

    note = doc.add_paragraph()
    note.add_run("审核重点：").bold = True
    note.add_run(
        "本稿只说明各格式规则在规范中的适用位置和来源依据，供人工审核后再更新 checker 规则。"
        "当前已确认“论文提交日期”仅属于封面项目，不应在正文中按封面日期规则处理。"
    )

    doc.add_heading("一、结构位置边界", level=1)
    structure_rows = [
        ["封面", "1.1；3.2 表“封面”", "中图分类号、学校代码、题目、作者、专业、导师、论文提交日期等封面项目。", "只在统一封面页内生效。"],
        ["封二", "1.2；2.3.3", "原创性说明、版权使用授权书、密级等。", "不是正文；密级与封面/封二位置绑定。"],
        ["摘要和关键词", "1.3；2.5；3.3 表“摘要和关键词”", "中文摘要、英文摘要、关键词、摘要页眉页码。", "关键词规则只在摘要区关键词行生效。"],
        ["目录", "1.4；2.4；3.4 表“目录”", "目录标题、论文总页数、各级目录项、目录页眉页码。", "目录项规则不应用于正文标题。"],
        ["正文", "1.6；2.6；3.5 表“正文”", "引言/绪论、章节标题、正文段落、图表、公式、正文页眉页码。", "正文中的日期、机构名、技术名词均按正文段落或对应标题/图表规则处理。"],
        ["参考文献", "1.7；2.7；3.6 表“其它”", "参考文献标题和著录条目。", "参考文献条目规则仅在参考文献区生效。"],
        ["附录/作者简历/致谢", "1.8；1.9；3.6 表“其它”", "附录标题与正文、作者简历、致谢。", "按后置部分规则处理，不套正文普通段落规则。"],
    ]
    add_table(doc, ["位置", "规范来源", "包含内容", "分类边界"], structure_rows, [2.2, 4.2, 8.0, 8.4])

    doc.add_heading("二、格式规则来源清单", level=1)
    rule_rows = [
        ["页面设置", "全篇", "3.1 表“纸张要求和页面设置”", "A4；上/下 2.5cm，左/右 3cm。", "可全篇检查。"],
        ["封面题目", "封面", "3.2 表“封面”行“论文题目”", "中文宋体三号加粗居中；英文 Times New Roman 三号加粗居中。", "只作用于封面题目。"],
        ["封面身份字段", "封面", "3.2 表“封面”姓名、学号、学院、专业、研究方向、指导教师、辅助指导教师", "中文仿宋_GB2312 四号；英文 Times New Roman 四号加粗。", "只作用于封面字段，不作用于正文中的同名词。"],
        ["论文提交日期", "封面", "1.1.12；2.3.2；3.2 表“封面”行“论文提交日期”", "答辩通过后按最终版实际提交时间填写；封面日期用阿拉伯数字；Times New Roman 四号。", "仅封面图二/统一封面中“论文提交日期”字段。正文中出现 2020 年等日期必须按正文段落处理。"],
        ["中文摘要标题", "摘要", "2.5.1；3.3 表“摘要和关键词”行“标题”", "“摘 要”二字中间空一格，宋体三号加粗居中，段前/后 0.5 行，固定 20 磅。", "只在中文摘要标题生效。"],
        ["英文摘要标题", "摘要", "2.5.1；3.3 表“摘要和关键词”行“标题”", "ABSTRACT 全大写，Times New Roman 四号加粗居中，段前/后 0.5 行。", "只在英文摘要标题生效。"],
        ["摘要正文", "摘要", "1.3；2.5.1；3.3 表“摘要和关键词”行“段落文字”", "中文宋体小四固定 20 磅；英文 Times New Roman 小四固定 20 磅。", "只在摘要正文区生效。"],
        ["关键词", "摘要", "1.3；2.5.2；3.3 表“摘要和关键词”行“关键词”", "关键词另起一行，左顶格，关键词间用分号；中文宋体小四，“关键词”加粗；英文 key words 加粗。", "只在摘要区关键词行生效。"],
        ["目录标题", "目录", "3.4 表“目录”行“标题”", "“目 录”宋体三号加粗居中，段前/后 0.5 行。", "只在目录页标题生效。"],
        ["目录项", "目录", "2.4；3.4 表“目录”各级目录行", "一级顶格；二级缩进 2 字符；三级缩进 4 字符；宋体小四固定 20 磅，页码右对齐。", "只作用于目录页，不作用于正文标题。"],
        ["正文章标题", "正文", "2.6.1；2.6.2.1；3.5 表“正文”行“各章标题”", "宋体三号加粗居中，固定 20 磅，段前/后 0.5 行，章序号与章名空一格。", "正文每章另起页。"],
        ["正文二级标题", "正文", "2.6.2.1；3.5 表“正文”行“二级标题”", "宋体四号加粗左对齐，固定 20 磅，段前/后 0.5 行，序号与题名空一格。", "正文标题，不属于目录。"],
        ["正文三级标题", "正文", "2.6.2.1；3.5 表“正文”行“三级标题”", "宋体小四加粗左对齐，固定 20 磅，段前/后 0.5 行，序号与题名空一格。", "正文标题，不建议四级标题。"],
        ["正文段落文字", "正文", "1.6；3.5 表“正文”行“段落文字（正文）”", "中文宋体小四；英文 Times New Roman；左端对齐；首行缩进 2 个汉字符；段前后 0 磅；固定 20 磅。", "正文里的普通叙述、年份日期、机构名、技术词均按此规则。"],
        ["图题/图注", "正文", "2.6.6.1；3.5 表“正文”行“图序、图名、图注”", "按章编号，置于图下方，宋体五号居中，固定 20 磅，图序与图名空一格。", "只作用于图题/图注。"],
        ["表题/表注", "正文", "2.6.6.2；3.5 表“正文”行“表序、表名、表注”", "按章编号，置于表上方，宋体五号居中，固定 20 磅，表序与表名空一格；续表头右顶格。", "只作用于表题/表注。"],
        ["公式", "正文", "2.6.7；3.5 表“正文”行“表达式（公式）”", "公式居中；编号加圆括号，宋体五号，右顶格。", "只作用于公式编号/公式段。"],
        ["页眉", "声明页至最后页", "2.6.3；3.3/3.4/3.5 表“页眉”", "标注“成都信息工程大学学士学位论文”，宋体小五居中；英文/数字 Times New Roman 小五。", "封面及声明页之前不应添加。"],
        ["页码", "摘要/目录/正文", "2.6.2.3；3.3/3.4/3.5 表“页码”", "摘要至目录用小写罗马数字；正文至致谢用“第*页 共*页”，居中，小五。", "按文档分区处理。"],
        ["符号说明", "符号说明", "1.5；3.6 表“其它”行“符号说明”", "标题字号等同正文；说明部分宋体五号，固定 20 磅。", "仅符号说明区。"],
        ["参考文献", "参考文献", "1.7；2.7；3.6 表“其它”行“参考文献”", "标题同章标题；条目宋体小四，英文 Times New Roman，固定 20 磅；续行缩进两个字符。", "条目还需符合 GB/T 7714-2015。"],
        ["附录", "附录", "1.8；3.6 表“其它”行“附录”", "标题同参考文献；正文宋体小四，英文 Times New Roman，两端对齐，首行缩进 2 字符，固定 20 磅。", "仅附录区。"],
        ["作者简历", "作者简历", "3.6 表“其它”行“作者简历及攻读学位期间发表的学术论文与研究成果”", "标题同章标题；正文中文宋体小四，英文/数字 Times New Roman，固定 20 磅。", "仅作者简历区。"],
        ["致谢", "致谢", "1.9；3.6 表“其它”行“致谢”", "标题同章标题；正文宋体小四，固定 20 磅，标题中间空一字符。", "仅致谢区。"],
    ]
    add_table(doc, ["规则", "适用位置", "规范来源", "规则摘要", "边界/备注"], rule_rows, [2.6, 2.2, 5.2, 9.6, 6.0])

    doc.add_heading("三、当前误判说明与拟调整点", level=1)
    bug_rows = [
        ["截图中的正文日期被判为“论文提交日期”", "正文 1.2 普通段落中出现“2020 年 1 月 22 日”。", "该位置属于正文段落文字，规范来源应为 3.5 表“段落文字（正文）”。", "分类时只有在封面区域且靠近“论文提交日期”字段/封面表单结构时，才允许使用 cover submission date 规则。"],
        ["封面提交日期规则位置", "封面图二/统一封面字段“论文提交日期”。", "来源：1.1.12、2.3.2、3.2 表“封面”行“论文提交日期”。", "日期格式检查不应全局匹配所有 yyyy 年 m 月 d 日。"],
        ["正文普通段落规则位置", "引言/绪论、论文主体、结论中的普通叙述段落。", "来源：1.6、3.5 表“正文”行“段落文字（正文）”。", "正文中所有日期、机构名、英文技术名词应进入正文段落/run-format 检查。"],
    ]
    add_table(doc, ["问题", "当前证据", "应使用的规范位置", "建议调整"], bug_rows, [5.0, 6.2, 7.2, 7.2])

    doc.add_heading("四、待人工审核的问题", level=1)
    checklist = [
        "封面字段是否只按统一封面实际版式识别，而不是按关键词全篇匹配？",
        "正文开始位置是否以“第一章/引言/绪论”等结构边界为准，而不是简单按页码或目录状态推断？",
        "摘要关键词中英文规则是否需要拆成中文关键词与英文 Key words 两条？",
        "页眉页码是否只报告风险，还是允许在 Office/WPS 可用时自动修复？",
        "参考文献条目格式是否只做基础结构检查，GB/T 7714 详细著录是否保留人工审核？",
    ]
    for item in checklist:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["Normal"]
        p.add_run("□ ").bold = True
        p.add_run(item)

    doc.add_paragraph()
    final_note = doc.add_paragraph()
    final_note.add_run("下一步建议：").bold = True
    final_note.add_run(
        "请先在本审核稿中确认每条规则的位置边界。确认后，再更新 rules.json 与分类逻辑，"
        "重点修复 cover_date 只能出现在封面字段的问题。"
    )

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
